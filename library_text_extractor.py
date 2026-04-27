"""Text extraction from uploaded Intelligence Library files (PDF, DOCX, TXT/MD)."""

import logging
import os

logger = logging.getLogger(__name__)

_TEXT_CAP = 5_000_000
_TRUNCATION_NOTICE = "\n\n[Text truncated — original file exceeded 5,000,000 characters]"

_SUPPORTED_PDF = {".pdf"}
_SUPPORTED_DOCX = {".docx"}
_SUPPORTED_TEXT = {".txt", ".md"}


def extract_text_from_file(file_path: str, mime_type_or_extension: str) -> str | None:
    """
    Extract plain text from a file on disk.

    Returns the extracted text (possibly truncated to 5 M chars), or None if the
    file type is unsupported or extraction fails.  Never raises.
    """
    # Normalise to a dotted extension
    ext = mime_type_or_extension.lower()
    if "/" in ext:
        # Caller passed a MIME type; derive extension from the file name instead
        ext = os.path.splitext(file_path)[1].lower()
    elif ext and not ext.startswith("."):
        ext = "." + ext

    try:
        if ext in _SUPPORTED_PDF:
            return _extract_pdf(file_path)
        if ext in _SUPPORTED_DOCX:
            return _extract_docx(file_path)
        if ext in _SUPPORTED_TEXT:
            return _extract_text(file_path)
        logger.warning("unsupported file type for text extraction: %s", ext)
        return None
    except Exception as exc:
        logger.warning("Text extraction failed for %s: %s", file_path, exc)
        return None


def backfill_library_text_content(library_folder: str) -> None:
    """
    Populate text_content for all file-type IntelligenceItem rows where it is NULL
    and the file exists on disk.  Idempotent — re-running does nothing when all
    eligible rows already have text_content.

    Logs a summary on completion.
    """
    from models import db, IntelligenceItem  # deferred to avoid circular import

    eligible = (
        IntelligenceItem.query
        .filter(
            IntelligenceItem.entry_type == "file",
            IntelligenceItem.text_content.is_(None),
            IntelligenceItem.file_path.isnot(None),
        )
        .all()
    )

    total = len(eligible)
    filled = 0
    failed = 0

    for item in eligible:
        abs_path = os.path.join(library_folder, item.file_path)
        if not os.path.isfile(abs_path):
            continue

        ext = os.path.splitext(item.file_path)[1].lower()
        text = extract_text_from_file(abs_path, ext)

        if text is not None:
            item.text_content = text
            filled += 1
        else:
            failed += 1

    db.session.commit()
    logger.info(
        "Backfilled %d of %d eligible items; %d failed extraction",
        filled, total, failed,
    )


# ── private helpers ────────────────────────────────────────────────────────────

def _extract_pdf(file_path: str) -> str | None:
    import pypdf  # noqa: PLC0415

    reader = pypdf.PdfReader(file_path)
    pages = []
    for page in reader.pages:
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(text)

    if not pages:
        return None
    return _cap("\n\n".join(pages))


def _extract_docx(file_path: str) -> str | None:
    from docx import Document  # noqa: PLC0415

    doc = Document(file_path)
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)

    if not parts:
        return None
    return _cap("\n".join(parts))


def _extract_text(file_path: str) -> str | None:
    try:
        with open(file_path, encoding="utf-8") as f:
            text = f.read().strip()
    except UnicodeDecodeError:
        with open(file_path, encoding="latin-1") as f:
            text = f.read().strip()

    if not text:
        return None
    return _cap(text)


def _cap(text: str) -> str:
    if len(text) <= _TEXT_CAP:
        return text
    return text[:_TEXT_CAP] + _TRUNCATION_NOTICE
