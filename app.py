import csv as _csv
import io
import json
import os
import re
import secrets
import uuid
from datetime import date, datetime, timezone

from flask import (
    Flask, render_template, request, redirect, url_for, flash,
    send_from_directory, send_file, jsonify, abort, session,
)
from werkzeug.utils import secure_filename
from flask_login import (
    LoginManager, login_user, logout_user, login_required, current_user,
)

from config import Config
from models import (
    db, User, Company, Project, Drawing, DrawingPage, SearchHistory, Report,
    LaborRate, InsuranceRate, PasswordResetToken, LoginEvent,
    IntelligenceTag, IntelligenceItem, QuoteBatch,
    ComparisonSummary, QuoteComparisonExport, Takeoff, WorkspaceMessage, WorkspaceThread,
    ROLE_SUPERADMIN, ROLE_ADMIN, ROLE_USER, ROLES,
    DOC_TYPES, DEFAULT_DOC_TYPE,
    PROJECT_STATUSES, TAKEOFF_STATUSES,
)
from email_notify import send_password_reset_email_async
from apscheduler.schedulers.background import BackgroundScheduler
from pipeline import start_worker
from reports import REPORT_TEMPLATES, enqueue_report, start_report_worker
from search import search_drawings
from calculations import calculate_painting_quantities, has_any_inputs
from xlsm_importer import parse_estimate_workbook
from library_text_extractor import extract_text_from_file, backfill_library_text_content
from synonyms import expand_query_terms, STOPWORDS


CCC_ADMIN_SEEDS = [
    ("orgon@muehlhan.com", "orgon"),
    ("j.brockman@muehlhan.com", "j.brockman"),
    ("lasater@muehlhan.com", "lasater"),
    ("moore@muehlhan.com", "moore"),
]
CCC_ADMIN_TEMP_PASSWORD = "Temp?Access123"
CCC_COMPANY_ID = 1

WORK_SCOPE_OPTIONS = [
    "Coating & Painting",
    "Lead Abatement",
    "Blast Cleaning",
    "High Pressure Water Washing",
    "Bridge Work",
    "Marine Vessels",
    "Industrial Tanks",
    "Scaffolding",
    "Confined Space",
    "Traffic Control",
    "Environmental Compliance",
    "SWPPP",
    "Encroachment Permit Work",
    "Other",
]

_LEAD_TRIGGER_SCOPES = {"Coating & Painting", "Blast Cleaning", "High Pressure Water Washing"}
_BULK_UPLOAD_EXTS = {".pdf", ".docx", ".txt", ".md"}


def _seed_ccc_admins():
    """Idempotently ensure the four CCC admin accounts exist with a forced password reset.

    Skipped silently if the target company does not yet exist, and per-user if an
    account with that email already exists (we do NOT reset an existing user's
    password here — that would be a footgun for subsequent deploys).
    """
    company = db.session.get(Company, CCC_COMPANY_ID)
    if not company:
        print(f"[powerscan] CCC seed: company id={CCC_COMPANY_ID} not found, skipping admin seed", flush=True)
        return

    created = 0
    for email, username in CCC_ADMIN_SEEDS:
        if User.query.filter_by(email=email).first():
            continue
        if User.query.filter_by(username=username).first():
            print(f"[powerscan] CCC seed: username '{username}' taken, skipping {email}", flush=True)
            continue
        user = User(
            username=username,
            email=email,
            role=ROLE_ADMIN,
            company_id=CCC_COMPANY_ID,
            must_change_password=True,
        )
        user.set_password(CCC_ADMIN_TEMP_PASSWORD)
        db.session.add(user)
        created += 1
    if created:
        db.session.commit()
        print(f"[powerscan] CCC seed: created {created} admin account(s)", flush=True)


def _run_migrations(database):
    """Add any missing columns to existing tables via ALTER TABLE."""
    conn = database.engine.raw_connection()
    cursor = conn.cursor()
    # Each entry: (table, column, column_def)
    migrations = [
        ("drawing", "total_pages", "INTEGER DEFAULT 0"),
        ("drawing", "pages_processed", "INTEGER DEFAULT 0"),
        ("drawing", "doc_type", "VARCHAR(40) DEFAULT 'Drawing'"),
        ("report", "file_path", "VARCHAR(300)"),
        ("user", "must_change_password", "BOOLEAN DEFAULT 0 NOT NULL"),

        ("project", "work_scope", "TEXT"),
        ("project", "scope_details", "TEXT"),
        ("intelligence_item", "pricing_items_json", "TEXT"),
        ("intelligence_item", "conditions_text", "TEXT"),
        ("intelligence_item", "flags_json", "TEXT"),
        ("intelligence_item", "raw_text_excerpt", "TEXT"),
        ("intelligence_item", "extraction_status", "VARCHAR(20) DEFAULT 'manual'"),
        ("intelligence_item", "vendor_name", "VARCHAR(200)"),
        ("intelligence_item", "vendor_contact", "TEXT"),
        ("intelligence_item", "quote_date", "DATE"),
        ("intelligence_item", "expiration_date", "DATE"),
        ("intelligence_item", "shortlisted", "BOOLEAN DEFAULT 0 NOT NULL"),
        ("intelligence_item", "shortlist_notes", "TEXT"),
        ("intelligence_item", "shortlisted_at", "DATETIME"),
        ("intelligence_item", "shortlisted_by", "INTEGER"),
        ("intelligence_item", "shortlisted_bid_id", "INTEGER"),
        ("intelligence_item", "shortlisted_scope_option", "VARCHAR(100)"),
        ("comparison_summary", "skippy_recommendation", "TEXT"),
        ("comparison_summary", "takeoff_id", "INTEGER"),
        ("project", "status", "VARCHAR(32) NOT NULL DEFAULT 'Active'"),
        ("project", "archived_at", "DATETIME"),
        ("project", "bid_date", "DATE"),
        ("quote_batch", "takeoff_id", "INTEGER"),
        ("report", "takeoff_id", "INTEGER"),
        ("takeoff", "scopes", "TEXT"),
        ("intelligence_item", "content_hash", "VARCHAR(64)"),
        # Takeoff inputs — project parameters
        ("takeoff", "deck_area_sf", "NUMERIC"),
        ("takeoff", "blast_level", "VARCHAR(32)"),
        ("takeoff", "abrasive_type", "VARCHAR(64)"),
        ("takeoff", "abrasive_lb_per_sf", "NUMERIC"),
        # Takeoff inputs — materials
        ("takeoff", "primer_vol_pct", "NUMERIC"),
        ("takeoff", "primer_mils", "NUMERIC"),
        ("takeoff", "primer_gal", "NUMERIC"),
        ("takeoff", "second_primer_vol_pct", "NUMERIC"),
        ("takeoff", "second_primer_mils", "NUMERIC"),
        ("takeoff", "second_primer_gal", "NUMERIC"),
        ("takeoff", "stripe_prime_vol_pct", "NUMERIC"),
        ("takeoff", "stripe_prime_mils", "NUMERIC"),
        ("takeoff", "stripe_prime_gal", "NUMERIC"),
        ("takeoff", "stripe_intermediate_vol_pct", "NUMERIC"),
        ("takeoff", "stripe_intermediate_mils", "NUMERIC"),
        ("takeoff", "stripe_intermediate_gal", "NUMERIC"),
        ("takeoff", "intermediate_vol_pct", "NUMERIC"),
        ("takeoff", "intermediate_mils", "NUMERIC"),
        ("takeoff", "intermediate_gal", "NUMERIC"),
        ("takeoff", "finish_vol_pct", "NUMERIC"),
        ("takeoff", "finish_mils", "NUMERIC"),
        ("takeoff", "finish_gal", "NUMERIC"),
        # Takeoff inputs — labor SF/HR + workers/nozzle
        ("takeoff", "mobilize_sf_per_hr", "NUMERIC"),
        ("takeoff", "mobilize_workers_per_nozzle", "NUMERIC"),
        ("takeoff", "equip_setup_sf_per_hr", "NUMERIC"),
        ("takeoff", "equip_setup_workers_per_nozzle", "NUMERIC"),
        ("takeoff", "scaffold_sf_per_hr", "NUMERIC"),
        ("takeoff", "scaffold_workers_per_nozzle", "NUMERIC"),
        ("takeoff", "containment_sf_per_hr", "NUMERIC"),
        ("takeoff", "containment_workers_per_nozzle", "NUMERIC"),
        ("takeoff", "masking_sf_per_hr", "NUMERIC"),
        ("takeoff", "masking_workers_per_nozzle", "NUMERIC"),
        ("takeoff", "pressure_wash_sf_per_hr", "NUMERIC"),
        ("takeoff", "pressure_wash_workers_per_nozzle", "NUMERIC"),
        ("takeoff", "caulking_sf_per_hr", "NUMERIC"),
        ("takeoff", "caulking_workers_per_nozzle", "NUMERIC"),
        ("takeoff", "blast_sf_per_hr", "NUMERIC"),
        ("takeoff", "blast_workers_per_nozzle", "NUMERIC"),
        ("takeoff", "primer_labor_sf_per_hr", "NUMERIC"),
        ("takeoff", "primer_labor_workers_per_nozzle", "NUMERIC"),
        ("takeoff", "second_primer_labor_sf_per_hr", "NUMERIC"),
        ("takeoff", "second_primer_labor_workers_per_nozzle", "NUMERIC"),
        ("takeoff", "stripe_prime_labor_sf_per_hr", "NUMERIC"),
        ("takeoff", "stripe_prime_labor_workers_per_nozzle", "NUMERIC"),
        ("takeoff", "stripe_intermediate_labor_sf_per_hr", "NUMERIC"),
        ("takeoff", "stripe_intermediate_labor_workers_per_nozzle", "NUMERIC"),
        ("takeoff", "intermediate_labor_sf_per_hr", "NUMERIC"),
        ("takeoff", "intermediate_labor_workers_per_nozzle", "NUMERIC"),
        ("takeoff", "finish_labor_sf_per_hr", "NUMERIC"),
        ("takeoff", "finish_labor_workers_per_nozzle", "NUMERIC"),
        ("takeoff", "traffic_control_sf_per_hr", "NUMERIC"),
        ("takeoff", "traffic_control_workers_per_nozzle", "NUMERIC"),
        ("takeoff", "inspection_touchup_sf_per_hr", "NUMERIC"),
        ("takeoff", "inspection_touchup_workers_per_nozzle", "NUMERIC"),
        ("takeoff", "osha_training_sf_per_hr", "NUMERIC"),
        ("takeoff", "osha_training_workers_per_nozzle", "NUMERIC"),
        # Takeoff inputs — time-based labor
        ("takeoff", "mobilize_hrs_per_day", "NUMERIC"),
        ("takeoff", "mobilize_days", "NUMERIC"),
        ("takeoff", "equip_setup_hrs_per_day", "NUMERIC"),
        ("takeoff", "equip_setup_days", "NUMERIC"),
        ("takeoff", "scaffold_hrs_per_day", "NUMERIC"),
        ("takeoff", "scaffold_days", "NUMERIC"),
        ("takeoff", "traffic_control_hrs_per_day", "NUMERIC"),
        ("takeoff", "traffic_control_days", "NUMERIC"),
        ("takeoff", "inspection_touchup_hrs_per_day", "NUMERIC"),
        ("takeoff", "inspection_touchup_days", "NUMERIC"),
        ("takeoff", "osha_training_hrs_per_day", "NUMERIC"),
        ("takeoff", "osha_training_days", "NUMERIC"),
        # Takeoff inputs — shift structure
        ("takeoff", "shift_hours_per_day", "NUMERIC"),
        ("takeoff", "shift_days_total", "NUMERIC"),
        ("takeoff", "crew_size", "INTEGER"),
        ("takeoff", "shifts_per_day", "INTEGER"),
        # Workspace threads
        ("workspace_message", "thread_id", "INTEGER"),
        # Phase 3C Part A: drawings linked to Library
        ("intelligence_item", "drawing_id", "INTEGER"),
    ]
    for table, column, col_def in migrations:
        try:
            cursor.execute(f"ALTER TABLE {table} ADD COLUMN {column} {col_def}")
        except Exception:
            pass  # Column already exists
    # Indexes that may be missing on existing databases
    try:
        cursor.execute(
            "CREATE INDEX IF NOT EXISTS idx_intelligence_item_content_hash "
            "ON intelligence_item(content_hash)"
        )
    except Exception:
        pass
    try:
        cursor.execute(
            "CREATE INDEX IF NOT EXISTS idx_intelligence_item_drawing_id "
            "ON intelligence_item(drawing_id)"
        )
    except Exception:
        pass
    conn.commit()
    conn.close()


def _run_workspace_thread_migration(database):
    """Create one WorkspaceThread per user+project pair for any orphaned WorkspaceMessages."""
    conn = database.engine.raw_connection()
    cursor = conn.cursor()
    try:
        cursor.execute(
            "SELECT DISTINCT project_id, user_id FROM workspace_message WHERE thread_id IS NULL"
        )
        pairs = cursor.fetchall()
        for project_id, user_id in pairs:
            cursor.execute(
                "SELECT MIN(created_at), MAX(created_at) FROM workspace_message "
                "WHERE project_id=? AND user_id=? AND thread_id IS NULL",
                (project_id, user_id),
            )
            row = cursor.fetchone()
            min_dt, max_dt = (row[0], row[1]) if row else (None, None)
            cursor.execute(
                "INSERT INTO workspace_thread (project_id, user_id, title, created_at, updated_at) "
                "VALUES (?, ?, 'Imported conversation', ?, ?)",
                (project_id, user_id, min_dt, max_dt),
            )
            thread_id = cursor.lastrowid
            cursor.execute(
                "UPDATE workspace_message SET thread_id=? "
                "WHERE project_id=? AND user_id=? AND thread_id IS NULL",
                (thread_id, project_id, user_id),
            )
        conn.commit()
    except Exception as exc:
        print(f"[powerscan] workspace thread migration error: {exc}", flush=True)
    finally:
        conn.close()


def _auto_title(text: str) -> str:
    text = text.strip()
    if len(text) <= 50:
        return text
    return text[:50].rstrip() + "…"


_QUOTE_EXTRACTION_PROMPT = """You are an expert construction estimating assistant. Extract structured information from this vendor quote document.

Return ONLY a valid JSON object with exactly this structure (no markdown, no explanation, just JSON):
{
  "vendor_name": "string or null",
  "vendor_contact": "string or null — combine name, phone, email into one string",
  "quote_date": "YYYY-MM-DD or null",
  "expiration_date": "YYYY-MM-DD or null",
  "pricing_items": [
    {"label": "string", "amount": "string", "unit": "string or empty string", "notes": "string or null"}
  ],
  "conditions_text": "string or null — verbatim conditions, exclusions, validity period, scope limitations",
  "flags": ["array of applicable flag strings from the allowed list only"],
  "suggested_tags": ["array of short descriptive tag strings, lowercase"],
  "suggested_title": "short title for a library entry, 5-10 words",
  "suggested_description": "1-2 sentence description of what this quote covers and its scope"
}

Allowed flag values (use only these exact strings):
no-pricing-submitted, different-project-referenced, prevailing-wage-required, partial-scope-only, expires-soon, volume-tier-applies, needs-verification, capability-statement-only

Rules:
- Extract ALL pricing items including hourly rates, daily rates, monthly rates, overtime rates, mobilization, premiums, surcharges, and volume tiers
- If no pricing is provided, return empty pricing_items array and include "no-pricing-submitted" and/or "capability-statement-only" in flags
- Set "expires-soon" if the quote expires within 90 days from today or if it appears already expired
- Set "prevailing-wage-required" if the document mentions prevailing wage, Davis-Bacon, or certified payroll
- Set "needs-verification" if any data looks ambiguous, unclear, or inconsistent
- Preserve verbatim conditions, exclusions, and validity periods in conditions_text
- suggested_tags should reflect the type of work (e.g. "traffic control", "scaffolding", "crane mats")
"""

_ALLOWED_FLAGS = [
    "no-pricing-submitted",
    "different-project-referenced",
    "prevailing-wage-required",
    "partial-scope-only",
    "expires-soon",
    "volume-tier-applies",
    "needs-verification",
    "capability-statement-only",
]

QUOTE_BATCH_MAX_FILES = 15

_SKIPPY_SYSTEM_PROMPT = (
    "You are Skippy — a straight-talking construction estimating assistant built into PowerScan. "
    "Your job: look at the vendor quotes for a category and give a clear, opinionated recommendation. "
    "No waffling. No 'it depends' without specifics. Pick one vendor and back it up.\n\n"
    "Respond ONLY with a valid JSON object — no markdown, no explanation, just JSON — "
    "with exactly these keys:\n"
    "{\n"
    '  "pick": "vendor name + one punchy sentence on why they are the top pick",\n'
    '  "watch_outs": ["concern #1 (≤ 15 words)", "concern #2 (≤ 15 words)"],\n'
    '  "runner_up": "vendor name + one sentence on why they are close but not top",\n'
    '  "why_pick": "2-3 sentences — pricing, scope coverage, flags, reliability",\n'
    '  "why_not_runner_up": "1-2 sentences on what held the runner-up back",\n'
    '  "why_others_lower": "1-2 sentences on what knocked other vendors down the list",\n'
    '  "confirm_before_signing": ["verify item #1 (≤ 15 words)", "verify item #2 (≤ 15 words)"]\n'
    "}\n\n"
    "Rules:\n"
    "- watch_outs: 2-4 items. confirm_before_signing: 2-4 items.\n"
    "- If only one vendor exists, still give a pick but note there is no competition.\n"
    "- If the top pick submitted no pricing, flag it prominently in watch_outs.\n"
    "- runner_up and why_not_runner_up may be null strings if only one vendor.\n"
    "- why_others_lower may be null if fewer than 3 vendors.\n"
    "- Do not reference PowerScan or yourself by name in the JSON values."
)


def _generate_skippy_recommendation(api_key, project, items, category_tag):
    """Call Claude with _SKIPPY_SYSTEM_PROMPT and return parsed JSON dict, or None on failure."""
    import anthropic

    lines = []
    for i, item in enumerate(items, start=1):
        vendor = item.vendor_name or item.title or f"Vendor {i}"
        lines.append(f"\n--- Quote {i}: {vendor} ---")
        if item.quote_date:
            lines.append(f"Quote date: {item.quote_date}")
        if item.expiration_date:
            lines.append(f"Expiration: {item.expiration_date}")
        try:
            pricing = json.loads(item.pricing_items_json) if item.pricing_items_json else []
        except Exception:
            pricing = []
        if pricing:
            lines.append("Pricing:")
            for p in pricing:
                row = f"  • {p.get('label','')}: {p.get('amount','')} {p.get('unit','')}".rstrip()
                if p.get("notes"):
                    row += f" ({p['notes']})"
                lines.append(row)
        else:
            lines.append("Pricing: (none submitted)")
        try:
            flags = json.loads(item.flags_json) if item.flags_json else []
        except Exception:
            flags = []
        if flags:
            lines.append(f"Flags: {', '.join(flags)}")
        if item.conditions_text:
            lines.append(f"Conditions: {item.conditions_text[:500]}")
        if item.shortlisted:
            lines.append("Status: SHORTLISTED")

    quote_data = "\n".join(lines)
    user_prompt = (
        f"Category: {category_tag}\n"
        f"Project: {project.name}\n\n"
        f"Here are the vendor quotes:\n{quote_data}\n\n"
        "Give me your recommendation."
    )

    try:
        client = anthropic.Anthropic(api_key=api_key)
        message = client.messages.create(
            model="claude-opus-4-7",
            max_tokens=1024,
            system=_SKIPPY_SYSTEM_PROMPT,
            messages=[{"role": "user", "content": user_prompt}],
        )
        raw = message.content[0].text.strip()
        return json.loads(raw)
    except Exception:
        return None


# ── Workspace (conversational project surface) ─────────────────────────────

_WORKSPACE_SYSTEM_PROMPT = (
    "You are Skippy — a straight-talking construction estimating assistant embedded in PowerScan. "
    "You're in the Workspace: a conversational surface where estimators ask questions about a "
    "specific project and get grounded answers from the actual project documents.\n\n"
    "CRITICAL — answer using ONLY the documents and project information provided in this message's "
    "context block. The context is curated specifically for this question. If you cannot find the "
    "answer in what you've been given, say so plainly: name what you checked and that the answer "
    "wasn't there. Do not guess, invent, or reference documents that weren't provided to you.\n\n"
    "Read the specifications, library references, and engineering drawings provided and give "
    "direct, useful answers. No waffling.\n\n"
    "Citation format: cite inline immediately after the sentence, using [filename.pdf p.47] for "
    "page-specific text references or [filename.pdf] for general document references. "
    "For drawing images shown above, cite using [filename p.N] format (e.g. [bridge-plans.pdf p.3]). "
    "Put citations right after the sentence they support, not at the end of a paragraph.\n\n"
    "The user may ask questions across sessions; treat conversation history as context they expect "
    "you to remember. Be concise. Use plain language. These are working professionals — clarity "
    "and directness over hedging."
)

# ── Context assembly constants ────────────────────────────────────────────────

_CTX_MAX_TEXT_ITEMS = 10       # max Library items included in text context
_CTX_MAX_SMART_IMAGES = 10     # max drawing pages for query-targeted inclusion
_CTX_MAX_FALLBACK_IMAGES = 6   # max drawing pages for fallback (untargeted)
_CTX_SNIPPET_CHARS_TOP = 8000  # total chars across all snippets for top-ranked items
_CTX_SNIPPET_CHARS_LOW = 1500  # total chars for lower-priority items
_CTX_FALLBACK_THRESHOLD = 8    # min score to bypass fallback; below this → generic load


def _score_item(item, direct_terms: set[str], synonym_terms: set[str]) -> int:
    """Score an IntelligenceItem for relevance to the query terms."""
    score = 0

    def _hits(text: str, terms: set[str], weight: int) -> int:
        t = text.lower()
        return sum(weight for term in terms if term in t)

    # Title
    score += _hits(item.title, direct_terms, 10)
    score += _hits(item.title, synonym_terms, 4)

    # Tags
    for tag in item.tags:
        score += _hits(tag.name, direct_terms, 8)
        score += _hits(tag.name, synonym_terms, 3)

    # Description
    if item.description:
        score += _hits(item.description, direct_terms, 5)
        score += _hits(item.description, synonym_terms, 2)

    # Text content — multi-word terms score higher (specificity > volume)
    #   single-word direct: 3 pts/hit, cap 50 (150 pts max)
    #   multi-word direct:  8 pts/hit, cap 10  (80 pts max)
    #   single-word synonym: 1 pt/hit,  cap 50  (50 pts max)
    #   multi-word synonym:  4 pts/hit, cap 10  (40 pts max)
    if item.text_content:
        tl = item.text_content.lower()
        distinct_body_hits = 0
        for term in direct_terms:
            pts, cap = (8, 10) if " " in term else (3, 50)
            count = min(tl.count(term), cap)
            if count:
                score += count * pts
                distinct_body_hits += 1
                if term in tl[:1000]:
                    score += 2   # position bonus for front-of-doc appearance
        for term in synonym_terms:
            pts, cap = (4, 10) if " " in term else (1, 50)
            count = min(tl.count(term), cap)
            if count:
                score += count * pts
                if term in tl[:1000]:
                    score += 1

        # Coverage bonus: a doc that contains 3+ distinct query terms in its body
        # shows substantive treatment of the topic, not just incidental mentions
        if distinct_body_hits >= 4:
            score += 40
        elif distinct_body_hits == 3:
            score += 20

    return score


_SECTION_HEADER_RE = re.compile(
    r"^(?:§\s*\d[\d\-\.]*\b"          # § 3-1.05 / §3.1
    r"|[A-Z][A-Z0-9 \t\-—]{4,49}$"    # ALL-CAPS heading, 5-50 chars
    r"|(?:SECTION|Section)\s+\d"       # Section 3 / SECTION 3
    r"|\d+(?:\.\d+)+\s+\w"            # 3.05 Bonding
    r")",
    re.MULTILINE,
)


def _find_section_start(text: str, pos: int, lookback: int = 1500) -> int:
    """Walk backwards from pos looking for a recognisable section header.

    Returns the character offset of the header line start if found within
    lookback chars, otherwise returns max(0, pos - 500).
    """
    search_start = max(0, pos - lookback)
    region = text[search_start:pos]
    best_offset = None
    for m in _SECTION_HEADER_RE.finditer(region):
        # Keep the last (closest-to-pos) header match
        best_offset = m.start()
    if best_offset is not None:
        # Snap to line start
        line_start = region.rfind("\n", 0, best_offset)
        char_offset = search_start + (best_offset if line_start == -1 else line_start + 1)
        return char_offset
    return max(0, pos - 500)


def _extract_snippet(text: str, terms: set[str], budget: int = 2000) -> str:
    """Return up to 3 cluster-ranked snippets totalling at most budget chars.

    Algorithm:
      1. Collect all match positions for every term.
      2. Cluster positions within 1500 chars of each other.
      3. Score each cluster by density × distinct-term-count.
      4. Take top clusters; give a single dominant cluster more room.
      5. For each cluster start, walk back to the nearest section header.
      6. Concatenate snippets with a separator so Claude sees non-contiguous excerpts.
    """
    if not text:
        return ""

    tl = text.lower()

    # 1. Collect all hit positions with weights (multi-word terms count as 3 hits each
    #    so specific phrases like "payment bond" outweigh generic words like "bond")
    positions: list[tuple[int, int]] = []  # (char_offset, weight)
    for term in terms:
        w = 3 if " " in term else 1
        start = 0
        while True:
            idx = tl.find(term, start)
            if idx == -1:
                break
            positions.append((idx, w))
            start = idx + 1
    if not positions:
        # No keyword match — return beginning of doc
        return text[:budget] + ("…" if len(text) > budget else "")

    positions.sort()

    # 2. Cluster positions within 1500 chars (gap check uses char offset only)
    CLUSTER_GAP = 1500
    clusters: list[list[tuple[int, int]]] = []
    current: list[tuple[int, int]] = [positions[0]]
    for p, w in positions[1:]:
        if p - current[-1][0] <= CLUSTER_GAP:
            current.append((p, w))
        else:
            clusters.append(current)
            current = [(p, w)]
    clusters.append(current)

    # 3. Score each cluster: weighted_density × distinct-term-count
    def _cluster_score(cluster: list[tuple[int, int]]) -> float:
        # Floor span at 500 chars — without this, single-match clusters get
        # density = 1/(1/1000) = 1000, which beats every multi-match dense section
        span = max(cluster[-1][0] - cluster[0][0], 500)
        weighted_hits = sum(w for _, w in cluster)
        density = weighted_hits / (span / 1000)
        distinct = len({tl[p:p + 30] for p, _ in cluster})  # rough distinct-term proxy
        return density * distinct

    clusters.sort(key=_cluster_score, reverse=True)

    # 4. Decide how many snippets and how to divide the budget
    #
    # Large reference docs (>500K chars) always get 3 snippets — bypasses the
    # winner-takes-all check so multiple distinct sections surface even when one
    # cluster happens to hold >60% of hits for a specific search term.
    # Small docs keep the original behaviour: single dominant cluster gets full budget.
    large_doc = len(text) > 500_000
    total_hits = len(positions)
    top_hits = len(clusters[0])
    if large_doc:
        n = min(3, len(clusters))
        selected = clusters[:n]
        per_snippet_budget = budget // n
        context_pad = 500
    elif top_hits / total_hits >= 0.6 or len(clusters) == 1:
        selected = clusters[:1]
        per_snippet_budget = budget
        context_pad = 1000
    else:
        selected = clusters[:3]
        per_snippet_budget = budget // len(selected)
        context_pad = 500

    # 5 & 6. Build each snippet with section-header-aware start
    parts: list[str] = []
    for cluster in selected:
        cluster_start = cluster[0][0]
        cluster_end = cluster[-1][0]
        snippet_start = _find_section_start(text, cluster_start, lookback=1500)
        snippet_end = min(len(text), max(cluster_end + context_pad, snippet_start + per_snippet_budget))
        # Don't exceed per-snippet budget from snippet_start
        snippet_end = min(snippet_end, snippet_start + per_snippet_budget)
        piece = text[snippet_start:snippet_end]
        if snippet_start > 0:
            piece = "…" + piece
        if snippet_end < len(text):
            piece += "…"
        parts.append(piece)

    return "\n\n[...later in document...]\n\n".join(parts)


def build_workspace_context(project, query: str, processed_folder: str) -> dict:
    """Assemble context for a Workspace Claude call driven by the user's query.

    Returns:
        content_blocks  — list of image dicts ready to prepend to the Claude user message
        index_map       — {idx: {drawing_id, filename, page}} for citation extraction
        text_context    — assembled text block describing loaded documents
        used_fallback   — True if no strong matches found and generic context was used
    """
    from search import _load_and_shrink

    # 1. Synonym expansion + stopword filter
    direct_terms, synonym_terms = expand_query_terms(query)
    direct_terms = {t for t in direct_terms if t not in STOPWORDS}
    synonym_terms = {t for t in synonym_terms if t not in STOPWORDS}

    # If the query contained only stopwords there are no real keywords to match on
    if not direct_terms:
        # Skip scoring; go straight to fallback path below
        _empty_query_fallback = True
    else:
        _empty_query_fallback = False

    # 2. Fetch all candidate items (project-scoped + global)
    candidates = (
        IntelligenceItem.query
        .filter(
            db.or_(
                IntelligenceItem.project_id.is_(None),
                IntelligenceItem.project_id == project.id,
            )
        )
        .all()
    )

    # 3. Score and rank
    scored = sorted(
        [(item, _score_item(item, direct_terms, synonym_terms)) for item in candidates],
        key=lambda x: -x[1],
    )
    top_scored = [(item, s) for item, s in scored if s > 0][:_CTX_MAX_TEXT_ITEMS]
    max_score = top_scored[0][1] if top_scored else 0

    # 4. Drawing inclusion heuristic
    drawing_matches = [(item, s) for item, s in top_scored if item.drawing_id]
    text_matches    = [(item, s) for item, s in top_scored if not item.drawing_id]
    used_fallback = _empty_query_fallback or max_score < _CTX_FALLBACK_THRESHOLD

    content_blocks: list[dict] = []
    index_map: dict[int, dict] = {}
    idx = 0

    if used_fallback:
        # No strong signal — load first N generic drawing pages + recent library items
        pages = (
            db.session.query(DrawingPage, Drawing)
            .join(Drawing, DrawingPage.drawing_id == Drawing.id)
            .filter(Drawing.project_id == project.id, Drawing.status == "ready")
            .order_by(Drawing.original_filename, DrawingPage.page_number)
            .limit(_CTX_MAX_FALLBACK_IMAGES)
            .all()
        )
        for page, drawing in pages:
            abs_path = os.path.join(processed_folder, page.image_path)
            try:
                data = _load_and_shrink(abs_path)
            except Exception:
                continue
            idx += 1
            label = f"{drawing.original_filename} — page {page.page_number}"
            index_map[idx] = {"drawing_id": drawing.id, "filename": drawing.original_filename, "page": page.page_number}
            content_blocks.append({"type": "text", "text": label})
            content_blocks.append({"type": "image", "source": {"type": "base64", "media_type": "image/jpeg", "data": data}})
        # Use all scored items for text context when falling back (top 6, any score)
        fallback_text_items = [item for item, _ in scored[:6]]
    else:
        fallback_text_items = []
        # Load JPEGs for drawing-linked matches
        if drawing_matches:
            target_ids = {item.drawing_id for item, _ in drawing_matches if item.drawing_id}
            draw_pages = (
                db.session.query(DrawingPage, Drawing)
                .join(Drawing, DrawingPage.drawing_id == Drawing.id)
                .filter(Drawing.id.in_(target_ids), Drawing.status == "ready")
                .order_by(Drawing.original_filename, DrawingPage.page_number)
                .all()
            )
            for page, drawing in draw_pages:
                if idx >= _CTX_MAX_SMART_IMAGES:
                    break
                abs_path = os.path.join(processed_folder, page.image_path)
                try:
                    data = _load_and_shrink(abs_path)
                except Exception:
                    continue
                idx += 1
                label = f"{drawing.original_filename} — page {page.page_number}"
                index_map[idx] = {"drawing_id": drawing.id, "filename": drawing.original_filename, "page": page.page_number}
                content_blocks.append({"type": "text", "text": label})
                content_blocks.append({"type": "image", "source": {"type": "base64", "media_type": "image/jpeg", "data": data}})

    images_count = idx

    # 5. Build text context block
    scope_items = project.work_scope_list
    scope_line = ""
    if scope_items:
        prompt_scope = list(scope_items)
        if _LEAD_TRIGGER_SCOPES.intersection(set(scope_items)) and "Lead Abatement" not in scope_items:
            prompt_scope.append("Lead Abatement")
        scope_line = "Work scope: " + ", ".join(prompt_scope)

    ctx_lines = [f"Project: {project.name}"]
    if scope_line:
        ctx_lines.append(scope_line)
    if project.scope_details:
        ctx_lines.append(f"Scope details: {project.scope_details}")

    search_terms_str = ", ".join(sorted(direct_terms)[:8])
    if not used_fallback:
        ctx_lines.append(f"\nCONTEXT — loaded for query terms: {search_terms_str}")
    else:
        ctx_lines.append(f"\nCONTEXT — generic load (no strong keyword match for: {search_terms_str})")

    text_items_for_ctx = [item for item, _ in top_scored] if not used_fallback else fallback_text_items
    if text_items_for_ctx:
        ctx_lines.append("")
        all_terms = direct_terms | synonym_terms
        for rank, item in enumerate(text_items_for_ctx):
            scope_tag = "GLOBAL" if item.project_id is None else "PROJECT"
            header = f"[{scope_tag}] {item.original_filename or item.title}"
            if item.drawing_id and index_map:
                # Find the pages loaded for this drawing and build citation hints
                pages_for_drawing = sorted(
                    v["page"] for v in index_map.values() if v["drawing_id"] == item.drawing_id
                )
                if pages_for_drawing:
                    fname = item.original_filename or item.title
                    cite_refs = ", ".join(f"[{fname} p.{p}]" for p in pages_for_drawing)
                    header += f" — images loaded: {cite_refs}"
            ctx_lines.append(header)
            if item.text_content:
                max_chars = _CTX_SNIPPET_CHARS_TOP if rank < 5 else _CTX_SNIPPET_CHARS_LOW
                snippet = _extract_snippet(item.text_content, all_terms, max_chars)
                ctx_lines.append(snippet)
            elif item.description:
                ctx_lines.append(item.description[:300])
            ctx_lines.append("")
    else:
        ctx_lines.append("\n(No matching documents found in this project's library.)")

    if index_map:
        ctx_lines.append(
            f"{images_count} drawing page(s) shown as images above. "
            "Cite drawings using [filename p.N] format."
        )
    else:
        ctx_lines.append("No drawing images loaded for this question.")

    text_context = "\n".join(ctx_lines)

    # 6. Log
    items_count = len(text_items_for_ctx)
    print(
        f"[workspace_ctx] project={project.name!r} query={query[:60]!r} → "
        f"{items_count} text items + {images_count} images (fallback={used_fallback})",
        flush=True,
    )

    return {
        "content_blocks": content_blocks,
        "index_map": index_map,
        "text_context": text_context,
        "used_fallback": used_fallback,
    }


def _call_workspace(project, history, user_message, api_key, processed_folder):
    """Call Claude for a Workspace chat turn.

    Returns dict with 'answer' (str) and 'sources' (list of dicts).
    history is a list of WorkspaceMessage objects (already alternating user/assistant).
    """
    import anthropic
    from search import CLAUDE_MODEL, _extract_sources

    client = anthropic.Anthropic(api_key=api_key)

    # Build context driven by the user's query
    ctx = build_workspace_context(project, user_message, processed_folder)

    # User message content: images first (if any), then text context + question
    content = list(ctx["content_blocks"])
    content.append({"type": "text", "text": ctx["text_context"] + f"\n\nUser question: {user_message}"})

    # Build messages array: clean history (alternating, start with user) + new turn
    clean = list(history)
    while clean and clean[-1].role == "user":
        clean.pop()
    while clean and clean[0].role == "assistant":
        clean.pop(0)
    clean = clean[-20:]

    api_messages = [{"role": m.role, "content": m.content} for m in clean]
    api_messages.append({"role": "user", "content": content})

    try:
        response = client.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=2048,
            system=_WORKSPACE_SYSTEM_PROMPT,
            messages=api_messages,
        )
        answer = response.content[0].text
        sources = _extract_sources(answer, ctx["index_map"])
        return {"answer": answer, "sources": sources}
    except Exception as exc:
        print(f"[powerscan] workspace API error: {exc}", flush=True)
        return {"answer": "I ran into a problem generating a response. Please try again.", "sources": []}


def _extract_quote_file(api_key, file_path, original_filename):
    """Extract structured quote data from a PDF or image file using Claude Vision.

    Returns a dict with keys: original_filename, file_path, result (dict) or error (str).
    """
    import base64
    import io as _io
    import anthropic
    from PIL import Image
    from search import CLAUDE_MODEL, MAX_IMAGE_WIDTH

    ext = os.path.splitext(original_filename)[1].lower()
    is_image = ext in (".png", ".jpg", ".jpeg")

    content = []

    try:
        if is_image:
            with Image.open(file_path) as img:
                if img.mode not in ("RGB", "L"):
                    img = img.convert("RGB")
                if img.width > MAX_IMAGE_WIDTH:
                    new_h = round(img.height * MAX_IMAGE_WIDTH / img.width)
                    img = img.resize((MAX_IMAGE_WIDTH, new_h), Image.LANCZOS)
                buf = _io.BytesIO()
                img.save(buf, format="JPEG", quality=85, optimize=True)
            data = base64.standard_b64encode(buf.getvalue()).decode("utf-8")
            content.append({
                "type": "image",
                "source": {"type": "base64", "media_type": "image/jpeg", "data": data},
            })
        else:
            # PDF — convert pages to images (cap at 10 pages)
            from pdf2image import convert_from_path
            images = convert_from_path(file_path, dpi=200, first_page=1, last_page=10)
            for i, img in enumerate(images):
                if img.width > MAX_IMAGE_WIDTH:
                    new_h = round(img.height * MAX_IMAGE_WIDTH / img.width)
                    img = img.resize((MAX_IMAGE_WIDTH, new_h), Image.LANCZOS)
                buf = _io.BytesIO()
                img.save(buf, format="JPEG", quality=85, optimize=True)
                data = base64.standard_b64encode(buf.getvalue()).decode("utf-8")
                content.append({"type": "text", "text": f"Page {i + 1}:"})
                content.append({
                    "type": "image",
                    "source": {"type": "base64", "media_type": "image/jpeg", "data": data},
                })

        content.append({"type": "text", "text": _QUOTE_EXTRACTION_PROMPT})

        client = anthropic.Anthropic(api_key=api_key)
        message = client.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=2048,
            messages=[{"role": "user", "content": content}],
        )
        raw = message.content[0].text.strip()

        # Strip markdown code fences if Claude wrapped the JSON
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
            raw = raw.strip()

        result = json.loads(raw)

        # Sanitise flags to allowed values only
        raw_flags = result.get("flags") or []
        result["flags"] = [f for f in raw_flags if f in _ALLOWED_FLAGS]

        return {
            "original_filename": original_filename,
            "file_path": file_path,
            "result": result,
            "error": None,
        }
    except Exception as exc:
        return {
            "original_filename": original_filename,
            "file_path": file_path,
            "result": None,
            "error": str(exc),
        }


_CLASSIFY_PIPELINE_PROMPT = """Classify this construction document and return a JSON object with these exact fields:
{
  "doc_type": "Drawing|Contract|Specification|Bid Doc|Addendum|Estimation Notes|Quote|Other",
  "processing_pipeline": "text|drawing|quote",
  "is_image_only": true|false,
  "confidence": "high|medium|low"
}

Routing rules:
- doc_type "Drawing" → processing_pipeline "drawing"
- doc_type "Quote" (vendor quotes, pricing proposals, rate sheets, capability statements, material price lists) → processing_pipeline "quote"
- is_image_only true (PDF with no extractable text layer, regardless of doc_type) → processing_pipeline "drawing"
- All other doc types → processing_pipeline "text"

Estimation Notes are estimating notes, takeoff sheets, or handwritten cost notes.
Quote documents contain structured pricing from a vendor or supplier.
is_image_only is true only when the page contains NO readable text characters — only scanned images or graphics.

Reply with ONLY the JSON object, no other text."""


def _classify_file_pipeline(api_key, file_path, original_filename):
    """Classify a saved file and return its processing pipeline.

    Returns dict with keys: processing_pipeline ("text"|"drawing"|"quote"),
    doc_type, is_image_only, confidence. Falls back to "text" on any error.
    """
    import base64
    import anthropic as _anthropic
    from search import CLAUDE_MODEL

    _FALLBACK = {"processing_pipeline": "text", "doc_type": DEFAULT_DOC_TYPE,
                 "is_image_only": False, "confidence": "low"}

    ext = os.path.splitext(original_filename)[1].lower()
    if ext != ".pdf":
        return {"processing_pipeline": "text", "doc_type": "Other",
                "is_image_only": False, "confidence": "high"}

    if not api_key:
        return _FALLBACK

    try:
        from pdf2image import convert_from_path
        from PIL import Image as _Image

        images = convert_from_path(file_path, dpi=100, first_page=1, last_page=1)
        if not images:
            return _FALLBACK

        img = images[0]
        if img.width > 800:
            img = img.resize((800, int(img.height * 800 / img.width)), _Image.LANCZOS)
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=75)
        img_b64 = base64.standard_b64encode(buf.getvalue()).decode()

        client = _anthropic.Anthropic(api_key=api_key)
        resp = client.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=150,
            messages=[{"role": "user", "content": [
                {"type": "image", "source": {"type": "base64", "media_type": "image/jpeg", "data": img_b64}},
                {"type": "text", "text": _CLASSIFY_PIPELINE_PROMPT},
            ]}],
        )
        raw = resp.content[0].text.strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
            raw = raw.strip()

        result = json.loads(raw)
        pipeline = result.get("processing_pipeline", "text").lower()
        if pipeline not in ("text", "drawing", "quote"):
            print(f"[powerscan] _classify_file_pipeline: unrecognized pipeline '{pipeline}' for {original_filename}, defaulting to text", flush=True)
            pipeline = "text"
        result["processing_pipeline"] = pipeline
        return result
    except Exception as exc:
        print(f"[powerscan] _classify_file_pipeline error for {original_filename}: {exc}", flush=True)
        return _FALLBACK


def _write_category_sheet(ws, items, all_labels, per_item_pricing,
                          category_tag, summary_obj, project_name, ts):
    """Write one category's full comparison data into an openpyxl worksheet."""
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter

    HEADER_FILL = PatternFill("solid", fgColor="1F4E79")
    HEADER_FONT = Font(bold=True, color="FFFFFF", size=11)
    WARN_FONT   = Font(italic=True, color="CC0000", size=9)
    BOLD_FONT   = Font(bold=True, size=10)
    snapshot_note = (
        f"Snapshot as of {ts} — For reference only. "
        "Verify against current quotes before using."
    )

    # Mini-header rows
    ws.column_dimensions["A"].width = 22
    ws["A1"] = f"Category: {category_tag}"
    ws["A1"].font = Font(bold=True, size=13, color="1F4E79")
    ws["A2"] = f"Project: {project_name}"
    ws["A3"] = snapshot_note
    ws["A3"].font = WARN_FONT
    if summary_obj and summary_obj.summary_text:
        ws["A4"] = "AI Summary"
        ws["A4"].font = BOLD_FONT
        ws["A5"] = summary_obj.summary_text
        ws["A5"].alignment = Alignment(wrap_text=True)
        ws.row_dimensions[5].height = 80
        data_start_row = 7
    else:
        data_start_row = 5

    # Column headers
    headers = [
        "Vendor Name", "Contact", "Quote Date", "Expiration Date",
        "Source File", "Shortlisted", "Flags", "Conditions/Exclusions",
    ] + [f"Pricing: {lbl}" for lbl in all_labels]

    for col_idx, header in enumerate(headers, start=1):
        cell = ws.cell(row=data_start_row, column=col_idx, value=header)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(wrap_text=True)
        ws.column_dimensions[get_column_letter(col_idx)].width = 20

    # Data rows
    for row_idx, item in enumerate(items, start=data_start_row + 1):
        flags = []
        if item.flags_json:
            try:
                flags = json.loads(item.flags_json)
            except Exception:
                flags = []
        row_data = [
            item.vendor_name or "",
            item.vendor_contact or "",
            str(item.quote_date) if item.quote_date else "",
            str(item.expiration_date) if item.expiration_date else "",
            item.original_filename or "",
            "Yes" if item.shortlisted else "No",
            ", ".join(flags),
            item.conditions_text or "",
        ]
        for lbl in all_labels:
            pr = per_item_pricing.get(item.id, {}).get(lbl)
            if pr:
                row_data.append(f"{pr.get('amount','')} {pr.get('unit','')}".strip())
            else:
                row_data.append("—")
        for col_idx, val in enumerate(row_data, start=1):
            ws.cell(row=row_idx, column=col_idx, value=val).alignment = Alignment(wrap_text=True)


def _apply_item_tags(item, raw_tags_str):
    """Parse comma-separated tags, upsert IntelligenceTag rows, attach to item."""
    names = [t.strip() for t in raw_tags_str.split(",") if t.strip()]
    new_tags = []
    for name in names:
        tag = IntelligenceTag.query.filter_by(name=name).first()
        if not tag:
            tag = IntelligenceTag(name=name, usage_count=0)
            db.session.add(tag)
            db.session.flush()
        new_tags.append(tag)

    current_names = {t.name for t in item.tags}
    new_names = {t.name for t in new_tags}

    for tag in new_tags:
        if tag.name not in current_names:
            item.tags.append(tag)
            tag.usage_count += 1

    for tag in list(item.tags):
        if tag.name not in new_names:
            item.tags.remove(tag)
            if tag.usage_count > 0:
                tag.usage_count -= 1


def _decrement_removed_tags(old_tags, current_tags):
    """No-op: tag counts are already managed inside _apply_item_tags."""
    pass


_scheduler_started = False


def _do_archive_sweep(app=None):
    """Archive projects whose bid_date is 90+ days in the past and status != Archived.

    May be called inside a request context (app=None) or from the scheduler
    (pass app so we can push an app context).
    """
    from datetime import timedelta

    def _run():
        cutoff = date.today() - timedelta(days=90)
        projects = Project.query.filter(
            Project.bid_date != None,  # noqa: E711
            Project.bid_date < cutoff,
            Project.status != "Archived",
        ).all()
        count = 0
        for p in projects:
            p.status = "Archived"
            p.archived_at = datetime.now(timezone.utc)
            count += 1
        if count:
            db.session.commit()
        return count

    if app is not None:
        with app.app_context():
            return _run()
    return _run()


def backfill_content_hashes(library_folder):
    """Idempotent: compute SHA-256 for IntelligenceItems whose file exists but content_hash is NULL."""
    import hashlib
    items = IntelligenceItem.query.filter(
        IntelligenceItem.content_hash.is_(None),
        IntelligenceItem.file_path.isnot(None),
    ).all()
    total = len(items)
    hashed = 0
    for idx, item in enumerate(items):
        fpath = os.path.join(library_folder, item.file_path)
        if not os.path.exists(fpath):
            print(f"[backfill_content_hashes] WARNING: file missing for item {item.id} ({item.file_path})")
            continue
        with open(fpath, "rb") as fh:
            item.content_hash = hashlib.sha256(fh.read()).hexdigest()
        hashed += 1
        if hashed % 50 == 0:
            db.session.commit()
    if hashed % 50 != 0:
        db.session.commit()
    print(f"[backfill_content_hashes] hashed {hashed} of {total} items")


def backfill_drawings_to_library(upload_folder, api_key):
    """Idempotent: create a paired IntelligenceItem for every Drawing that lacks one.

    Classifies each Drawing's source PDF first (same classifier used by Library bulk
    upload) then branches:
      - TEXT pipeline  → run pypdf extraction, store text_content
      - DRAWING pipeline → create item with text_content=None (vision-only)
      - QUOTE or unexpected → skip with warning

    Per-drawing transactions ensure partial success survives errors or restarts.
    """
    if not api_key:
        print("[backfill] no ANTHROPIC_API_KEY configured — drawings backfill skipped (will retry on next startup with key set)", flush=True)
        return

    # Load orphans and immediately snapshot to plain dicts so SQLAlchemy session
    # expiry / rollback across iterations cannot cause DetachedInstanceError.
    linked_ids = (
        db.session.query(IntelligenceItem.drawing_id)
        .filter(IntelligenceItem.drawing_id.isnot(None))
        .subquery()
    )
    orphan_rows = Drawing.query.filter(Drawing.id.notin_(linked_ids)).all()
    orphans = [
        {
            "id": d.id,
            "original_filename": d.original_filename,
            "src_filename": d.filename,
            "project_id": d.project_id,
            "uploaded_by": d.uploaded_by,
        }
        for d in orphan_rows
    ]
    del orphan_rows  # free ORM objects; we work from plain dicts below

    total = len(orphans)
    print(f"[backfill] starting: {total} drawings to process", flush=True)

    text_count = 0
    drawing_count = 0
    quote_skipped = 0
    error_count = 0

    for n, d in enumerate(orphans, start=1):
        label = f"{n}/{total} {d['original_filename']}"
        try:
            file_path = os.path.join(upload_folder, d["src_filename"])
            if not os.path.isfile(file_path):
                print(f"[backfill] {label} → ERROR: source PDF not found at {file_path}", flush=True)
                error_count += 1
                continue

            classification = _classify_file_pipeline(api_key, file_path, d["original_filename"])
            pipeline = classification.get("processing_pipeline", "text")

            if pipeline == "quote":
                print(f"[backfill] {label} → QUOTE-skipped (unexpected in drawings pipeline)", flush=True)
                quote_skipped += 1
                continue

            if pipeline == "drawing":
                text_content = None
                print(f"[backfill] {label} → DRAWING → vision-only", flush=True)
                drawing_count += 1
            else:
                ext = os.path.splitext(d["src_filename"])[1].lower()
                text_content = extract_text_from_file(file_path, ext)
                chars = len(text_content) if text_content else 0
                print(f"[backfill] {label} → TEXT → extracted {chars} chars", flush=True)
                text_count += 1

            item = IntelligenceItem(
                title=os.path.splitext(d["original_filename"])[0],
                entry_type="text",
                text_content=text_content,
                original_filename=d["original_filename"],
                project_id=d["project_id"],
                auto_include_in_search=True,
                uploaded_by=d["uploaded_by"],
                drawing_id=d["id"],
            )
            db.session.add(item)
            db.session.commit()

        except Exception as exc:
            try:
                db.session.rollback()
            except Exception:
                pass
            print(f"[backfill] {label} → ERROR: {type(exc).__name__}: {exc}", flush=True)
            error_count += 1

    print(
        f"[backfill] complete: {total} drawings processed → "
        f"{text_count} TEXT (extracted) + {drawing_count} DRAWING (vision-only) + "
        f"{quote_skipped} QUOTE-skipped + {error_count} errors",
        flush=True,
    )


def _drop_feedback_table():
    """Idempotent: drop the legacy feedback table if it still exists in the database."""
    try:
        with db.engine.connect() as conn:
            conn.execute(db.text("DROP TABLE IF EXISTS feedback"))
            conn.commit()
        print("[powerscan] feedback table removed (or was already absent)", flush=True)
    except Exception as exc:
        print(f"[powerscan] could not drop feedback table: {exc}", flush=True)


def migrate_projects_to_takeoffs():
    """Idempotent: create one Draft takeoff per project that has no takeoffs yet."""
    projects = Project.query.all()
    created = 0
    for project in projects:
        if not project.takeoffs:
            existing_count = Takeoff.query.filter_by(project_id=project.id).count()
            n = existing_count + 1
            t = Takeoff(
                project_id=project.id,
                name=f"Takeoff {n}",
                status="Draft",
                revision_note="Auto-created by migration",
                scopes=project.work_scope,
            )
            db.session.add(t)
            created += 1
    if created:
        db.session.commit()
        print(f"[migrate_projects_to_takeoffs] Created {created} takeoff(s).")


def backfill_takeoff_ids():
    """Idempotent: assign the earliest takeoff per project to records with NULL takeoff_id."""
    projects = Project.query.all()
    batch_count = 0
    summary_count = 0
    report_count = 0
    has_report_project_id = hasattr(Report, "project_id")
    for project in projects:
        earliest = (
            Takeoff.query.filter_by(project_id=project.id)
            .order_by(Takeoff.created_at.asc())
            .first()
        )
        if not earliest:
            continue
        batches = QuoteBatch.query.filter_by(project_id=project.id, takeoff_id=None).all()
        for b in batches:
            b.takeoff_id = earliest.id
            batch_count += 1
        summaries = ComparisonSummary.query.filter_by(project_id=project.id, takeoff_id=None).all()
        for s in summaries:
            s.takeoff_id = earliest.id
            summary_count += 1
        if has_report_project_id:
            reports = Report.query.filter_by(project_id=project.id, takeoff_id=None).all()
            for r in reports:
                r.takeoff_id = earliest.id
                report_count += 1
        else:
            print("[backfill_takeoff_ids] Report has no project_id column — skipping reports")
    db.session.commit()
    print(f"[backfill_takeoff_ids] quote_batches: {batch_count}, comparison_summaries: {summary_count}, reports: {report_count}")


def rename_legacy_initial_takeoffs():
    """One-shot: rename 'Initial Takeoff / Final / Auto-created' records to 'Takeoff 1 / Draft'."""
    legacy = Takeoff.query.filter_by(name="Initial Takeoff", status="Final").all()
    updated = 0
    for t in legacy:
        # Rename if revision_note is NULL (old migration record) or contains the marker (new)
        if t.revision_note is None or "Auto-created" in t.revision_note:
            t.name = "Takeoff 1"
            t.status = "Draft"
            updated += 1
    if updated:
        db.session.commit()
    print(f"[takeoff_rename] updated {updated} legacy Initial Takeoff records to Takeoff 1 / Draft")


def backfill_takeoff_scopes():
    """Idempotent: copy parent Project scopes to any Takeoff where scopes IS NULL."""
    takeoffs = Takeoff.query.filter(Takeoff.scopes.is_(None)).all()
    updated = 0
    for t in takeoffs:
        t.scopes = t.project.work_scope
        updated += 1
    if updated:
        db.session.commit()
    print(f"[backfill_takeoff_scopes] updated {updated} takeoffs")


def _get_active_takeoff(project_id):
    """Return the active takeoff for a project (Final if exists, else latest Draft, else create one)."""
    final = (
        Takeoff.query.filter_by(project_id=project_id, status="Final")
        .order_by(Takeoff.created_at.desc())
        .first()
    )
    if final:
        return final
    draft = (
        Takeoff.query.filter_by(project_id=project_id, status="Draft")
        .order_by(Takeoff.created_at.desc())
        .first()
    )
    if draft:
        return draft
    project = db.session.get(Project, project_id)
    t = Takeoff(
        project_id=project_id,
        name="New Takeoff",
        status="Draft",
        scopes=project.work_scope if project else None,
    )
    db.session.add(t)
    db.session.commit()
    return t


_LAYERS = ["primer", "second_primer", "stripe_prime", "stripe_intermediate", "intermediate", "finish"]
_ALL_LABOR_TASKS = [
    "mobilize", "equip_setup", "scaffold", "containment", "masking",
    "pressure_wash", "caulking", "blast", "primer_labor",
    "second_primer_labor", "stripe_prime_labor", "stripe_intermediate_labor",
    "intermediate_labor", "finish_labor", "traffic_control",
    "inspection_touchup", "osha_training",
]
_TIME_TASKS = ["mobilize", "equip_setup", "scaffold", "traffic_control", "inspection_touchup", "osha_training"]


def _apply_takeoff_inputs(takeoff, form):
    """
    Write all takeoff input fields from a form-like mapping onto a Takeoff
    record. Does NOT commit — caller is responsible for db.session.commit().
    `form` may be flask.request.form or any dict-like object.
    """
    def _num(field, scale=2):
        v = (form.get(field) or "").strip()
        if not v:
            return None
        try:
            return round(float(v), scale)
        except (ValueError, TypeError):
            return None

    def _int(field):
        v = (form.get(field) or "").strip()
        if not v:
            return None
        try:
            return int(v)
        except (ValueError, TypeError):
            return None

    def _str(field):
        return (form.get(field) or "").strip() or None

    takeoff.deck_area_sf = _num("deck_area_sf")
    takeoff.blast_level = _str("blast_level")
    takeoff.abrasive_type = _str("abrasive_type")
    takeoff.abrasive_lb_per_sf = _num("abrasive_lb_per_sf", 3)

    for layer in _LAYERS:
        setattr(takeoff, f"{layer}_vol_pct", _num(f"{layer}_vol_pct"))
        setattr(takeoff, f"{layer}_mils", _num(f"{layer}_mils"))
        setattr(takeoff, f"{layer}_gal", _num(f"{layer}_gal"))

    for task in _ALL_LABOR_TASKS:
        setattr(takeoff, f"{task}_sf_per_hr", _num(f"{task}_sf_per_hr"))
        setattr(takeoff, f"{task}_workers_per_nozzle", _num(f"{task}_workers_per_nozzle"))

    for task in _TIME_TASKS:
        setattr(takeoff, f"{task}_hrs_per_day", _num(f"{task}_hrs_per_day"))
        setattr(takeoff, f"{task}_days", _num(f"{task}_days"))

    takeoff.shift_hours_per_day = _num("shift_hours_per_day")
    takeoff.shift_days_total = _num("shift_days_total")
    takeoff.crew_size = _int("crew_size")
    takeoff.shifts_per_day = _int("shifts_per_day")


def create_app():
    app = Flask(__name__)
    app.config.from_object(Config)

    os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)
    os.makedirs(app.config["PROCESSED_FOLDER"], exist_ok=True)
    os.makedirs(app.config["REPORTS_FOLDER"], exist_ok=True)
    os.makedirs(os.path.join(app.instance_path), exist_ok=True)
    library_folder = os.path.join(os.path.dirname(app.config["UPLOAD_FOLDER"]), "library_files")
    app.config["LIBRARY_FOLDER"] = library_folder
    os.makedirs(library_folder, exist_ok=True)

    db.init_app(app)

    login_manager = LoginManager()
    login_manager.login_view = "login"
    login_manager.init_app(app)

    @login_manager.user_loader
    def load_user(user_id):
        return db.session.get(User, int(user_id))

    with app.app_context():
        db.create_all()
        # Migrate: add columns that may be missing from older databases
        _run_migrations(db)
        # Migrate: create WorkspaceThreads for any orphaned WorkspaceMessages
        _run_workspace_thread_migration(db)
        # Create default superadmin if none exists
        if not User.query.filter_by(role=ROLE_SUPERADMIN).first():
            admin = User(
                username="admin",
                email="admin@powerscan.local",
                role=ROLE_SUPERADMIN,
            )
            admin.set_password("admin123")
            db.session.add(admin)
            db.session.commit()

        # Seed CCC admin accounts (idempotent — skipped if email already exists)
        _seed_ccc_admins()
        # Create Takeoff shell for any projects that have none
        migrate_projects_to_takeoffs()
        # Backfill takeoff_id on existing batches/summaries/reports
        backfill_takeoff_ids()
        # Rename any legacy "Initial Takeoff / Final" records to "Takeoff 1 / Draft"
        rename_legacy_initial_takeoffs()
        # Backfill scopes on any takeoffs where scopes is still NULL
        backfill_takeoff_scopes()
        # Compute SHA-256 content hashes for existing library files
        backfill_content_hashes(app.config["LIBRARY_FOLDER"])
        # Create paired Library items for any drawings that lack one (Phase 3C Part A)
        backfill_drawings_to_library(app.config["UPLOAD_FOLDER"], app.config.get("ANTHROPIC_API_KEY", ""))
        # Drop the legacy feedback table (Share Idea feature removed)
        _drop_feedback_table()

    # ── CLI commands ────────────────────────────────────────────────────────
    @app.cli.command("backfill-library-text")
    def _backfill_library_text_cmd():
        """Extract and store text_content for existing file-type Library entries."""
        backfill_library_text_content(app.config["LIBRARY_FOLDER"])

    # Start background conversion worker thread
    start_worker(app)
    start_report_worker(app)

    # Start archive-sweep scheduler (daily at 00:05 UTC)
    global _scheduler_started
    if not _scheduler_started:
        _scheduler_started = True
        _sched = BackgroundScheduler(timezone="UTC")
        _sched.add_job(
            lambda: _do_archive_sweep(app),
            "cron", hour=0, minute=5,
            id="archive_sweep", replace_existing=True,
        )
        _sched.start()

    # ── Decorators ──────────────────────────────────────────────

    def admin_required(f):
        from functools import wraps

        @wraps(f)
        def decorated(*args, **kwargs):
            if not current_user.is_authenticated or not current_user.is_admin:
                abort(403)
            return f(*args, **kwargs)
        return decorated

    def superadmin_required(f):
        from functools import wraps

        @wraps(f)
        def decorated(*args, **kwargs):
            if not current_user.is_authenticated or not current_user.is_superadmin:
                abort(403)
            return f(*args, **kwargs)
        return decorated

    # ── Auth Routes ─────────────────────────────────────────────

    @app.route("/login", methods=["GET", "POST"])
    def login():
        if current_user.is_authenticated:
            if current_user.must_change_password:
                return redirect(url_for("change_password"))
            if current_user.role != ROLE_SUPERADMIN:
                return redirect(url_for("projects", company_id=current_user.company_id))
            return redirect(url_for("dashboard"))
        if request.method == "POST":
            identifier = request.form.get("username", "").strip()
            password = request.form.get("password", "")
            user = (
                User.query.filter_by(username=identifier).first()
                or User.query.filter_by(email=identifier).first()
            )
            if user and user.check_password(password):
                login_user(user)
                try:
                    ip = (request.headers.get("X-Forwarded-For") or request.remote_addr or "").split(",")[0].strip()[:45]
                    db.session.add(LoginEvent(user_id=user.id, ip_address=ip or None))
                    db.session.commit()
                except Exception:
                    pass
                if user.must_change_password:
                    flash("Please choose a new password to finish logging in.", "info")
                    return redirect(url_for("change_password"))
                next_page = request.args.get("next")
                if not next_page and user.role != ROLE_SUPERADMIN:
                    return redirect(url_for("projects", company_id=user.company_id))
                return redirect(next_page or url_for("dashboard"))
            flash("Invalid username or password.", "danger")
        return render_template("login.html")

    @app.route("/change-password", methods=["GET", "POST"])
    @login_required
    def change_password():
        if request.method == "POST":
            new_password = request.form.get("new_password", "")
            confirm_password = request.form.get("confirm_password", "")
            if len(new_password) < 8:
                flash("Password must be at least 8 characters.", "danger")
            elif new_password != confirm_password:
                flash("Passwords do not match.", "danger")
            elif current_user.check_password(new_password):
                flash("New password must be different from your current password.", "danger")
            else:
                current_user.set_password(new_password)
                current_user.must_change_password = False
                db.session.commit()
                flash("Password updated. Welcome!", "success")
                return redirect(url_for("dashboard"))
        return render_template("change_password.html", forced=current_user.must_change_password)

    @app.route("/forgot-password", methods=["GET", "POST"])
    def forgot_password():
        if current_user.is_authenticated:
            return redirect(url_for("dashboard"))
        if request.method == "POST":
            email = (request.form.get("email") or "").strip().lower()
            user = User.query.filter(db.func.lower(User.email) == email).first()
            if user:
                token_val = secrets.token_urlsafe(32)
                tok = PasswordResetToken(user_id=user.id, token=token_val)
                db.session.add(tok)
                db.session.commit()
                send_password_reset_email_async(app, tok.id)
            flash("If that email is on file you will receive a reset link shortly.", "info")
            return redirect(url_for("login"))
        return render_template("forgot_password.html")

    @app.route("/reset-password/<token>", methods=["GET", "POST"])
    def reset_password(token):
        if current_user.is_authenticated:
            return redirect(url_for("dashboard"))
        tok = PasswordResetToken.query.filter_by(token=token, used=False).first()
        if not tok:
            flash("This reset link is invalid or has already been used.", "danger")
            return redirect(url_for("login"))
        age = (datetime.now(timezone.utc) - tok.created_at.replace(tzinfo=timezone.utc)).total_seconds()
        if age > 3600:
            flash("This reset link has expired. Please request a new one.", "danger")
            return redirect(url_for("forgot_password"))
        if request.method == "POST":
            new_pw = request.form.get("new_password", "")
            confirm_pw = request.form.get("confirm_password", "")
            if len(new_pw) < 8:
                flash("Password must be at least 8 characters.", "danger")
            elif new_pw != confirm_pw:
                flash("Passwords do not match.", "danger")
            else:
                tok.user.set_password(new_pw)
                tok.user.must_change_password = False
                tok.used = True
                db.session.commit()
                flash("Password updated. Please log in.", "success")
                return redirect(url_for("login"))
        return render_template("reset_password.html", token=token)

    # Guard: users with a forced reset flag are locked to the change-password page
    # until they pick a new password (or log out).
    @app.before_request
    def _enforce_password_reset():
        if not current_user.is_authenticated:
            return None
        if not current_user.must_change_password:
            return None
        allowed_endpoints = {"change_password", "logout", "login", "static", "forgot_password", "reset_password"}
        if request.endpoint in allowed_endpoints:
            return None
        return redirect(url_for("change_password"))

    @app.route("/logout")
    @login_required
    def logout():
        logout_user()
        return redirect(url_for("login"))

    # ── Dashboard ───────────────────────────────────────────────

    @app.route("/")
    @login_required
    def dashboard():
        if current_user.role != ROLE_SUPERADMIN:
            return redirect(url_for("projects", company_id=current_user.company_id))
        companies = Company.query.order_by(Company.name).all()
        stats = {
            "companies": Company.query.count(),
            "projects": sum(len(c.projects) for c in companies),
            "drawings": sum(len(p.drawings) for c in companies for p in c.projects),
        }
        return render_template("dashboard.html", companies=companies, stats=stats)

    # ── Company Routes ──────────────────────────────────────────

    @app.route("/companies")
    @login_required
    def companies():
        if current_user.is_superadmin:
            company_list = Company.query.order_by(Company.name).all()
        elif current_user.company_id:
            company_list = [current_user.company]
        else:
            company_list = []
        return render_template("companies.html", companies=company_list)

    @app.route("/companies/new", methods=["GET", "POST"])
    @login_required
    @admin_required
    def new_company():
        if request.method == "POST":
            name = request.form.get("name", "").strip()
            if not name:
                flash("Company name is required.", "danger")
            elif Company.query.filter_by(name=name).first():
                flash("Company already exists.", "danger")
            else:
                company = Company(name=name)
                db.session.add(company)
                db.session.commit()
                flash(f"Company '{name}' created.", "success")
                return redirect(url_for("companies"))
        return render_template("company_form.html")

    @app.route("/companies/<int:company_id>/delete", methods=["POST"])
    @login_required
    @superadmin_required
    def delete_company(company_id):
        company = db.session.get(Company, company_id) or abort(404)
        db.session.delete(company)
        db.session.commit()
        flash(f"Company '{company.name}' deleted.", "success")
        return redirect(url_for("companies"))

    # ── Project Routes ──────────────────────────────────────────

    @app.route("/companies/<int:company_id>/projects")
    @login_required
    def projects(company_id):
        company = db.session.get(Company, company_id) or abort(404)
        if not current_user.is_superadmin and current_user.company_id != company_id:
            abort(403)

        year_filter    = request.args.get("year",           "")
        search         = request.args.get("search",         "").strip()
        show_archived  = request.args.get("show_archived",  "0") == "1"

        all_projects = (
            Project.query
            .filter_by(company_id=company_id)
            .order_by(Project.name)
            .all()
        )

        years = sorted(
            {p.created_at.year for p in all_projects if p.created_at},
            reverse=True,
        )

        def _matches_search(p):
            if not search:
                return True
            sl = search.lower()
            return sl in (p.name or "").lower() or sl in (p.description or "").lower()

        def _matches_year(p):
            if not year_filter:
                return True
            try:
                return p.created_at and p.created_at.year == int(year_filter)
            except ValueError:
                return True

        active_projects = [
            p for p in all_projects
            if p.status == "Active" and _matches_year(p) and _matches_search(p)
        ]
        on_hold_projects = [
            p for p in all_projects
            if p.status == "On Hold" and _matches_year(p) and _matches_search(p)
        ]
        complete_projects = [
            p for p in all_projects
            if p.status == "Complete" and _matches_year(p) and _matches_search(p)
        ]

        archived = []
        if show_archived:
            archived = sorted(
                [p for p in all_projects
                 if p.status == "Archived" and _matches_year(p) and _matches_search(p)],
                key=lambda p: p.archived_at or p.created_at,
                reverse=True,
            )

        return render_template(
            "projects.html",
            company=company,
            active_projects=active_projects,
            on_hold_projects=on_hold_projects,
            complete_projects=complete_projects,
            archived_projects=archived,
            years=years,
            year_filter=year_filter,
            search=search,
            show_archived=show_archived,
            today=date.today(),
        )

    @app.route("/companies/<int:company_id>/projects/new", methods=["GET", "POST"])
    @login_required
    @admin_required
    def new_project(company_id):
        company = db.session.get(Company, company_id) or abort(404)
        if request.method == "POST":
            name = request.form.get("name", "").strip()
            description = request.form.get("description", "").strip()
            scope_items = request.form.getlist("work_scope")
            scope_details = request.form.get("scope_details", "").strip()
            status = request.form.get("status", "Active")
            if status not in PROJECT_STATUSES:
                status = "Active"
            bid_date_str = request.form.get("bid_date", "").strip()
            bid_date_val = None
            if bid_date_str:
                try:
                    bid_date_val = date.fromisoformat(bid_date_str)
                except ValueError:
                    pass
            if not name:
                flash("Project name is required.", "danger")
            else:
                project = Project(
                    name=name,
                    description=description,
                    company_id=company.id,
                    work_scope=json.dumps(scope_items) if scope_items else None,
                    scope_details=scope_details or None,
                    status=status,
                    bid_date=bid_date_val,
                )
                db.session.add(project)
                db.session.commit()
                flash(f"Project '{name}' created.", "success")
                return redirect(url_for("projects", company_id=company.id))
        return render_template("project_form.html", company=company, project=None,
                               scope_options=WORK_SCOPE_OPTIONS, current_scope=[],
                               project_statuses=PROJECT_STATUSES)

    @app.route("/projects/<int:project_id>/edit", methods=["GET", "POST"])
    @login_required
    @admin_required
    def edit_project(project_id):
        project = db.session.get(Project, project_id) or abort(404)
        company = db.session.get(Company, project.company_id) or abort(404)
        if not current_user.is_superadmin and current_user.company_id != project.company_id:
            abort(403)
        if request.method == "POST":
            name = request.form.get("name", "").strip()
            description = request.form.get("description", "").strip()
            scope_items = request.form.getlist("work_scope")
            scope_details = request.form.get("scope_details", "").strip()
            new_status = request.form.get("status", "Active")
            if new_status not in PROJECT_STATUSES:
                new_status = "Active"
            bid_date_str = request.form.get("bid_date", "").strip()
            bid_date_val = None
            if bid_date_str:
                try:
                    bid_date_val = date.fromisoformat(bid_date_str)
                except ValueError:
                    pass
            if not name:
                flash("Project name is required.", "danger")
            else:
                old_status = project.status
                project.name = name
                project.description = description
                project.work_scope = json.dumps(scope_items) if scope_items else None
                project.scope_details = scope_details or None
                project.status = new_status
                project.bid_date = bid_date_val
                if new_status == "Archived" and old_status != "Archived":
                    project.archived_at = datetime.now(timezone.utc)
                elif new_status != "Archived":
                    project.archived_at = None
                db.session.commit()
                flash("Project updated.", "success")
                return redirect(url_for("project_hub", project_id=project.id))
        return render_template("project_form.html", company=company, project=project,
                               scope_options=WORK_SCOPE_OPTIONS,
                               current_scope=project.work_scope_list,
                               project_statuses=PROJECT_STATUSES)

    @app.route("/companies/<int:company_id>/projects/archive-sweep", methods=["POST"])
    @login_required
    @admin_required
    def project_archive_sweep(company_id):
        company = db.session.get(Company, company_id) or abort(404)
        if not current_user.is_superadmin and current_user.company_id != company_id:
            abort(403)
        count = _do_archive_sweep()
        if count:
            flash(f"Archive sweep complete — {count} project{'s' if count != 1 else ''} archived.", "success")
        else:
            flash("Archive sweep complete — no projects needed archiving.", "info")
        return redirect(url_for("projects", company_id=company_id))

    @app.route("/projects/<int:project_id>/delete", methods=["POST"])
    @login_required
    @admin_required
    def delete_project(project_id):
        import shutil
        project = db.session.get(Project, project_id) or abort(404)
        if not current_user.is_superadmin and current_user.company_id != project.company_id:
            abort(403)

        project_name = project.name
        company_id = project.company_id

        # Collect disk paths before touching the DB
        files_to_delete = []
        dirs_to_delete = []

        drawings = Drawing.query.filter_by(project_id=project_id).all()
        for d in drawings:
            files_to_delete.append(os.path.join(app.config["UPLOAD_FOLDER"], d.filename))
            dirs_to_delete.append(os.path.join(app.config["PROCESSED_FOLDER"], str(d.id)))

        # NEVER modify records with project_id IS NULL — those are Company Global and must survive project deletion
        lib_items = IntelligenceItem.query.filter(
            IntelligenceItem.project_id == project_id
        ).all()
        for item in lib_items:
            if item.file_path:
                files_to_delete.append(os.path.join(app.config["LIBRARY_FOLDER"], item.file_path))

        batches = QuoteBatch.query.filter(QuoteBatch.project_id == project_id).all()
        for b in batches:
            dirs_to_delete.append(
                os.path.join(app.config["LIBRARY_FOLDER"], "quotes_staging", b.batch_id)
            )

        reports = Report.query.filter(Report.project_id == project_id).all()
        for r in reports:
            if r.file_path:
                files_to_delete.append(os.path.join(app.config["REPORTS_FOLDER"], r.file_path))

        # Single transaction: explicit deletes for non-ORM-cascade models first,
        # then cascade from db.session.delete(project) handles the rest.
        try:
            # NEVER modify records with project_id IS NULL — those are Company Global and must survive project deletion
            ComparisonSummary.query.filter(
                ComparisonSummary.project_id == project_id
            ).delete(synchronize_session="fetch")
            QuoteComparisonExport.query.filter(
                QuoteComparisonExport.project_id == project_id
            ).delete(synchronize_session="fetch")
            QuoteBatch.query.filter(
                QuoteBatch.project_id == project_id
            ).delete(synchronize_session="fetch")
            Takeoff.query.filter(
                Takeoff.project_id == project_id
            ).delete(synchronize_session="fetch")
            WorkspaceMessage.query.filter(
                WorkspaceMessage.project_id == project_id
            ).delete(synchronize_session="fetch")
            WorkspaceThread.query.filter(
                WorkspaceThread.project_id == project_id
            ).delete(synchronize_session="fetch")
            # Cascade handles: Drawings/DrawingPages, SearchHistory, Reports,
            # IntelligenceItems (project_id==project_id only — NULL items untouched)
            db.session.delete(project)
            db.session.commit()
        except Exception as exc:
            db.session.rollback()
            app.logger.error(f"Project delete failed for id={project_id}: {exc}")
            flash(f"Could not delete '{project_name}' — an error occurred. No data was changed.", "danger")
            return redirect(url_for("project_hub", project_id=project_id))

        # DB committed — clean up disk (non-fatal: log failures but don't crash)
        for f in files_to_delete:
            try:
                if os.path.exists(f):
                    os.remove(f)
                    app.logger.info(f"Project delete: removed file {f}")
            except Exception as exc:
                app.logger.warning(f"Project delete: could not remove file {f}: {exc}")
        for d in dirs_to_delete:
            try:
                if os.path.isdir(d):
                    shutil.rmtree(d, ignore_errors=True)
                    app.logger.info(f"Project delete: removed dir {d}")
            except Exception as exc:
                app.logger.warning(f"Project delete: could not remove dir {d}: {exc}")

        flash(f"Project '{project_name}' deleted.", "success")
        return redirect(url_for("projects", company_id=company_id))

    # ── Project Hub ─────────────────────────────────────────────

    @app.route("/projects/<int:project_id>")
    @login_required
    def project_hub(project_id):
        project = db.session.get(Project, project_id) or abort(404)
        if not current_user.is_superadmin and current_user.company_id != project.company_id:
            abort(403)

        doc_count = Drawing.query.filter_by(project_id=project_id).count()
        intel_count = IntelligenceItem.query.filter_by(project_id=project_id).count()
        saved_categories = (
            db.session.query(QuoteBatch.category_tag)
            .filter_by(project_id=project_id, status="saved")
            .distinct()
            .count()
        )
        reviewing_batches = QuoteBatch.query.filter_by(
            project_id=project_id, status="reviewing"
        ).count()
        report_count = Report.query.filter_by(project_id=project_id).count()
        takeoff_count = Takeoff.query.filter_by(project_id=project_id).count()
        previous_bids_count = 0
        bid_date = None
        days_to_bid = None

        if project.bid_date:
            from datetime import date as _date
            bid_date = project.bid_date
            days_to_bid = (bid_date - _date.today()).days

        return render_template(
            "project_hub.html",
            project=project,
            doc_count=doc_count,
            intel_count=intel_count,
            saved_categories=saved_categories,
            reviewing_batches=reviewing_batches,
            report_count=report_count,
            takeoff_count=takeoff_count,
            previous_bids_count=previous_bids_count,
            bid_date=bid_date,
            days_to_bid=days_to_bid,
        )

    @app.route("/projects/<int:project_id>/previous-bids")
    @login_required
    def project_previous_bids(project_id):
        project = db.session.get(Project, project_id) or abort(404)
        if not current_user.is_superadmin and current_user.company_id != project.company_id:
            abort(403)
        return render_template("project_previous_bids_placeholder.html", project=project)

    # ── Workspace ───────────────────────────────────────────────

    def _ws_drawings_list(project_id):
        return (
            Drawing.query
            .filter_by(project_id=project_id)
            .order_by(Drawing.created_at.desc())
            .all()
        )

    def _ws_thread_list(project_id, user_id):
        return (
            WorkspaceThread.query
            .filter_by(project_id=project_id, user_id=user_id)
            .order_by(WorkspaceThread.updated_at.desc())
            .all()
        )

    @app.route("/projects/<int:project_id>/workspace")
    @login_required
    def project_workspace(project_id):
        project = db.session.get(Project, project_id) or abort(404)
        if not current_user.is_superadmin and current_user.company_id != project.company_id:
            abort(403)
        return render_template(
            "project_workspace.html",
            project=project,
            threads=_ws_thread_list(project_id, current_user.id),
            messages=[],
            active_thread=None,
            drawings_list=_ws_drawings_list(project_id),
        )

    @app.route("/projects/<int:project_id>/workspace/thread/<int:thread_id>")
    @login_required
    def project_workspace_thread(project_id, thread_id):
        project = db.session.get(Project, project_id) or abort(404)
        if not current_user.is_superadmin and current_user.company_id != project.company_id:
            abort(403)
        thread = db.session.get(WorkspaceThread, thread_id) or abort(404)
        if thread.project_id != project_id or thread.user_id != current_user.id:
            abort(403)
        raw_messages = (
            WorkspaceMessage.query
            .filter_by(thread_id=thread_id)
            .order_by(WorkspaceMessage.created_at.asc())
            .all()
        )
        for msg in raw_messages:
            try:
                msg.sources = json.loads(msg.sources_json) if msg.sources_json else []
            except Exception:
                msg.sources = []
        return render_template(
            "project_workspace.html",
            project=project,
            threads=_ws_thread_list(project_id, current_user.id),
            messages=raw_messages,
            active_thread=thread,
            drawings_list=_ws_drawings_list(project_id),
        )

    @app.route("/projects/<int:project_id>/workspace/send", methods=["POST"])
    @login_required
    def workspace_send(project_id):
        project = db.session.get(Project, project_id) or abort(404)
        if not current_user.is_superadmin and current_user.company_id != project.company_id:
            abort(403)

        data = request.get_json(silent=True) or {}
        user_text = data.get("message", "").strip()
        if not user_text:
            return jsonify({"error": "Empty message."}), 400
        if not app.config["ANTHROPIC_API_KEY"]:
            return jsonify({"error": "Claude API key not configured."}), 500

        thread_id = data.get("thread_id")
        thread = None
        if thread_id:
            thread = db.session.get(WorkspaceThread, int(thread_id))
            if not thread or thread.project_id != project_id or thread.user_id != current_user.id:
                return jsonify({"error": "Thread not found."}), 404

        # Capture history BEFORE saving the new user message
        if thread:
            history = (
                WorkspaceMessage.query
                .filter_by(thread_id=thread.id)
                .order_by(WorkspaceMessage.created_at.asc())
                .all()
            )
        else:
            now = datetime.now(timezone.utc)
            thread = WorkspaceThread(
                project_id=project_id,
                user_id=current_user.id,
                title=_auto_title(user_text),
                created_at=now,
                updated_at=now,
            )
            db.session.add(thread)
            db.session.flush()
            history = []

        # Save user message
        user_msg = WorkspaceMessage(
            project_id=project_id,
            user_id=current_user.id,
            thread_id=thread.id,
            role="user",
            content=user_text,
        )
        db.session.add(user_msg)
        db.session.commit()

        # Call Claude
        result = _call_workspace(
            project, history, user_text,
            app.config["ANTHROPIC_API_KEY"],
            app.config["PROCESSED_FOLDER"],
        )

        answer = result.get("answer", "Sorry, I could not generate a response.")
        sources = result.get("sources", [])

        # Save assistant message and bump thread timestamp
        asst_msg = WorkspaceMessage(
            project_id=project_id,
            user_id=current_user.id,
            thread_id=thread.id,
            role="assistant",
            content=answer,
            sources_json=json.dumps(sources) if sources else None,
        )
        db.session.add(asst_msg)
        thread.updated_at = datetime.now(timezone.utc)
        db.session.commit()

        return jsonify({
            "answer": answer,
            "sources": sources,
            "thread_id": thread.id,
            "thread_title": thread.title,
            "created_at": asst_msg.created_at.strftime("%b %d, %Y %H:%M"),
        })

    @app.route("/projects/<int:project_id>/workspace/thread/<int:thread_id>/delete", methods=["POST"])
    @login_required
    def workspace_thread_delete(project_id, thread_id):
        project = db.session.get(Project, project_id) or abort(404)
        if not current_user.is_superadmin and current_user.company_id != project.company_id:
            abort(403)
        thread = db.session.get(WorkspaceThread, thread_id) or abort(404)
        if thread.project_id != project_id or thread.user_id != current_user.id:
            abort(403)
        db.session.delete(thread)
        db.session.commit()
        return jsonify({"ok": True})

    @app.route("/projects/<int:project_id>/workspace/thread/<int:thread_id>/rename", methods=["POST"])
    @login_required
    def workspace_thread_rename(project_id, thread_id):
        project = db.session.get(Project, project_id) or abort(404)
        if not current_user.is_superadmin and current_user.company_id != project.company_id:
            abort(403)
        thread = db.session.get(WorkspaceThread, thread_id) or abort(404)
        if thread.project_id != project_id or thread.user_id != current_user.id:
            abort(403)
        data = request.get_json(silent=True) or {}
        title = data.get("title", "").strip()
        if not title:
            return jsonify({"error": "Title cannot be empty."}), 400
        thread.title = title[:200]
        db.session.commit()
        return jsonify({"ok": True, "title": thread.title})

    # ── Takeoff Routes ──────────────────────────────────────────

    @app.route("/projects/<int:project_id>/takeoffs")
    @login_required
    def takeoffs_list(project_id):
        project = db.session.get(Project, project_id) or abort(404)
        if not current_user.is_superadmin and current_user.company_id != project.company_id:
            abort(403)
        takeoffs = Takeoff.query.filter_by(project_id=project_id).order_by(Takeoff.created_at.desc()).all()
        return render_template("takeoffs.html", project=project, takeoffs=takeoffs)

    @app.route("/projects/<int:project_id>/takeoffs/new", methods=["GET", "POST"])
    @login_required
    @admin_required
    def new_takeoff(project_id):
        project = db.session.get(Project, project_id) or abort(404)
        if not current_user.is_superadmin and current_user.company_id != project.company_id:
            abort(403)
        if request.method == "POST":
            name = request.form.get("name", "").strip() or "New Takeoff"
            status = request.form.get("status", "Draft")
            if status not in TAKEOFF_STATUSES:
                status = "Draft"
            revision_note = request.form.get("revision_note", "").strip() or None
            submitted_amount_raw = request.form.get("submitted_amount", "").strip()
            submitted_amount = None
            if submitted_amount_raw:
                try:
                    submitted_amount = float(submitted_amount_raw.replace(",", ""))
                except ValueError:
                    pass
            selected_scopes = request.form.getlist("scopes")
            if not selected_scopes:
                flash("At least one scope must be selected.", "danger")
                return render_template(
                    "takeoff_form.html",
                    project=project,
                    takeoff=None,
                    takeoff_statuses=TAKEOFF_STATUSES,
                    project_scopes=project.work_scope_list,
                    selected_scopes=[],
                )
            t = Takeoff(
                project_id=project_id,
                name=name,
                status=status,
                revision_note=revision_note,
                submitted_amount=submitted_amount,
                scopes=json.dumps(selected_scopes),
                created_by_user_id=current_user.id,
            )
            db.session.add(t)
            db.session.commit()
            flash(f"Takeoff '{name}' created.", "success")
            return redirect(url_for("takeoff_detail", takeoff_id=t.id))
        return render_template(
            "takeoff_form.html",
            project=project,
            takeoff=None,
            takeoff_statuses=TAKEOFF_STATUSES,
            project_scopes=project.work_scope_list,
            selected_scopes=project.work_scope_list,
        )

    @app.route("/takeoffs/<int:takeoff_id>/edit", methods=["GET", "POST"])
    @login_required
    @admin_required
    def edit_takeoff(takeoff_id):
        takeoff = db.session.get(Takeoff, takeoff_id) or abort(404)
        project = takeoff.project
        if not current_user.is_superadmin and current_user.company_id != project.company_id:
            abort(403)
        if request.method == "POST":
            takeoff.name = request.form.get("name", "").strip() or takeoff.name
            status = request.form.get("status", "Draft")
            if status in TAKEOFF_STATUSES:
                takeoff.status = status
            takeoff.revision_note = request.form.get("revision_note", "").strip() or None
            submitted_amount_raw = request.form.get("submitted_amount", "").strip()
            if submitted_amount_raw:
                try:
                    takeoff.submitted_amount = float(submitted_amount_raw.replace(",", ""))
                except ValueError:
                    pass
            else:
                takeoff.submitted_amount = None
            # Scopes are locked once Final; ignore submitted scopes if so
            if takeoff.status != "Final":
                selected_scopes = request.form.getlist("scopes")
                if not selected_scopes:
                    flash("At least one scope must be selected.", "danger")
                    return render_template(
                        "takeoff_form.html",
                        project=project,
                        takeoff=takeoff,
                        takeoff_statuses=TAKEOFF_STATUSES,
                        project_scopes=project.work_scope_list,
                        selected_scopes=takeoff.scopes_list,
                    )
                takeoff.scopes = json.dumps(selected_scopes)
            db.session.commit()
            flash("Takeoff updated.", "success")
            return redirect(url_for("takeoff_detail", takeoff_id=takeoff.id))
        return render_template(
            "takeoff_form.html",
            project=project,
            takeoff=takeoff,
            takeoff_statuses=TAKEOFF_STATUSES,
            project_scopes=project.work_scope_list,
            selected_scopes=takeoff.scopes_list,
        )

    @app.route("/takeoffs/<int:takeoff_id>")
    @login_required
    def takeoff_detail(takeoff_id):
        takeoff = db.session.get(Takeoff, takeoff_id) or abort(404)
        project = takeoff.project
        if not current_user.is_superadmin and current_user.company_id != project.company_id:
            abort(403)
        batches = (
            QuoteBatch.query.filter_by(takeoff_id=takeoff_id)
            .filter(QuoteBatch.status != "cancelled")
            .order_by(QuoteBatch.created_at.desc())
            .all()
        )
        summaries = ComparisonSummary.query.filter_by(takeoff_id=takeoff_id).order_by(ComparisonSummary.generated_at.desc()).all()
        reports = Report.query.filter_by(takeoff_id=takeoff_id).order_by(Report.created_at.desc()).all()
        quantities = calculate_painting_quantities(takeoff) if has_any_inputs(takeoff) else None
        return render_template(
            "takeoff_detail.html",
            project=project,
            takeoff=takeoff,
            batches=batches,
            summaries=summaries,
            reports=reports,
            quantities=quantities,
        )

    @app.route("/takeoffs/<int:takeoff_id>/inputs", methods=["GET", "POST"])
    @login_required
    def takeoff_inputs(takeoff_id):
        takeoff = db.session.get(Takeoff, takeoff_id) or abort(404)
        project = takeoff.project
        if not current_user.is_superadmin and current_user.company_id != project.company_id:
            abort(403)

        if request.method == "POST":
            _apply_takeoff_inputs(takeoff, request.form)
            db.session.commit()
            flash("Takeoff inputs saved.", "success")
            return redirect(url_for("takeoff_inputs", takeoff_id=takeoff_id))

        return render_template("takeoff_inputs.html", project=project, takeoff=takeoff)

    @app.route("/takeoffs/<int:takeoff_id>/import-xlsm", methods=["GET", "POST"])
    @login_required
    def import_xlsm(takeoff_id):
        takeoff = db.session.get(Takeoff, takeoff_id) or abort(404)
        project = takeoff.project
        if not current_user.is_superadmin and current_user.company_id != project.company_id:
            abort(403)

        if request.method == "GET":
            return render_template("takeoff_import_upload.html", project=project, takeoff=takeoff)

        # POST — parse the uploaded workbook
        f = request.files.get("xlsm_file")
        if not f or not f.filename:
            flash("No file selected.", "danger")
            return render_template("takeoff_import_upload.html", project=project, takeoff=takeoff)

        fname = secure_filename(f.filename)
        if not fname.lower().endswith((".xlsm", ".xlsx")):
            flash("Only .xlsm or .xlsx files are accepted.", "danger")
            return render_template("takeoff_import_upload.html", project=project, takeoff=takeoff)

        f.seek(0, 2)
        file_size = f.tell()
        f.seek(0)
        if file_size > 10 * 1024 * 1024:
            flash("File exceeds the 10 MB size limit.", "danger")
            return render_template("takeoff_import_upload.html", project=project, takeoff=takeoff)

        try:
            parsed = parse_estimate_workbook(f.stream)
        except Exception as exc:
            app.logger.error("XLSM parse error for takeoff %s: %s", takeoff_id, exc)
            flash(f"Could not parse workbook: {exc}", "danger")
            return render_template("takeoff_import_upload.html", project=project, takeoff=takeoff)

        found_count = sum(1 for v in parsed.values() if v is not None)
        return render_template(
            "takeoff_import_review.html",
            project=project,
            takeoff=takeoff,
            parsed=parsed,
            found_count=found_count,
            total_count=len(parsed),
            filename=fname,
        )

    @app.route("/takeoffs/<int:takeoff_id>/import-xlsm/save", methods=["POST"])
    @login_required
    def import_xlsm_save(takeoff_id):
        takeoff = db.session.get(Takeoff, takeoff_id) or abort(404)
        project = takeoff.project
        if not current_user.is_superadmin and current_user.company_id != project.company_id:
            abort(403)
        _apply_takeoff_inputs(takeoff, request.form)
        db.session.commit()
        flash("Takeoff inputs imported and saved.", "success")
        return redirect(url_for("takeoff_inputs", takeoff_id=takeoff_id))

    # ── Drawing Routes ──────────────────────────────────────────

    @app.route("/projects/<int:project_id>/drawings", methods=["GET", "POST"])
    @login_required
    def drawings(project_id):
        project = db.session.get(Project, project_id) or abort(404)
        if not current_user.is_superadmin and current_user.company_id != project.company_id:
            abort(403)

        # Search tab has moved to Workspace
        if request.method == "GET" and request.args.get("tab", "").strip().lower() == "search":
            flash("Search has been upgraded — meet Skippy, your new Workspace assistant.", "info")
            return redirect(url_for("project_workspace", project_id=project_id))

        filter_doc_type = request.values.get("filter_doc_type", "").strip()
        if filter_doc_type and filter_doc_type not in DOC_TYPES:
            filter_doc_type = ""

        results = None
        query = ""
        search_doc_type = ""
        if request.method == "POST":
            query = request.form.get("query", "").strip()
            search_doc_type = request.form.get("search_doc_type", "").strip()
            if search_doc_type and search_doc_type not in DOC_TYPES:
                search_doc_type = ""
            if not query:
                flash("Please enter a question.", "danger")
            elif not app.config["ANTHROPIC_API_KEY"]:
                flash("Claude API key not configured. Set ANTHROPIC_API_KEY environment variable.", "danger")
            else:
                scope_items = project.work_scope_list
                scope_context = None
                if scope_items:
                    prompt_scope = list(scope_items)
                    if _LEAD_TRIGGER_SCOPES.intersection(scope_items) and "Lead Abatement" not in scope_items:
                        prompt_scope.append("Lead Abatement")
                    scope_text = ", ".join(prompt_scope)
                    scope_context = (
                        f"This project involves the following work scope: {scope_text}."
                    )
                    if project.scope_details:
                        scope_context += f" Additional details: {project.scope_details}."
                    scope_context += (
                        " Focus your answer on provisions and requirements relevant to this scope."
                    )
                include_library = request.form.get("include_library") == "1"
                library_context = None
                if include_library:
                    lib_items = (
                        IntelligenceItem.query
                        .filter_by(auto_include_in_search=True)
                        .filter(
                            db.or_(
                                IntelligenceItem.project_id.is_(None),
                                IntelligenceItem.project_id == project.id,
                            )
                        )
                        .all()
                    )
                    if not current_user.is_superadmin:
                        lib_items = [
                            it for it in lib_items
                            if it.project_id is None or (
                                it.project and it.project.company_id == current_user.company_id
                            )
                        ]
                    if lib_items:
                        parts = []
                        for it in lib_items:
                            tags_str = ", ".join(t.name for t in it.tags)
                            body = it.text_content or f"[file: {it.original_filename}]"
                            entry = f"- {it.title}"
                            if it.description:
                                entry += f": {it.description}"
                            if tags_str:
                                entry += f" [tags: {tags_str}]"
                            if it.text_content:
                                entry += f"\n  Content: {body}"
                            parts.append(entry)
                        library_context = "Intelligence Library entries:\n" + "\n".join(parts)

                results = search_drawings(
                    query,
                    project.id,
                    app.config["ANTHROPIC_API_KEY"],
                    app.config["PROCESSED_FOLDER"],
                    doc_type=search_doc_type or None,
                    scope_context=scope_context,
                    library_context=library_context,
                )
                history = SearchHistory(
                    project_id=project.id,
                    user_id=current_user.id,
                    query=query,
                    answer=results.get("answer", "") if results else "",
                    doc_type_filter=search_doc_type or None,
                )
                db.session.add(history)
                db.session.commit()

        drawings_q = Drawing.query.filter_by(project_id=project.id)
        if filter_doc_type:
            drawings_q = drawings_q.filter_by(doc_type=filter_doc_type)
        drawings_list = drawings_q.order_by(Drawing.created_at.desc()).all()

        reports_list = (
            Report.query.filter_by(project_id=project.id)
            .order_by(Report.created_at.desc())
            .limit(20)
            .all()
        )

        history_entries = (
            db.session.query(SearchHistory)
            .filter_by(project_id=project.id)
            .order_by(SearchHistory.created_at.desc())
            .limit(50)
            .all()
        )

        active_tab = request.args.get("tab", "").strip().lower()
        if active_tab not in ("documents", "search", "reports", "history"):
            if request.method == "POST" or query:
                active_tab = "search"
            else:
                active_tab = "documents"

        has_ready_docs = Drawing.query.filter_by(project_id=project.id, status="ready").count() > 0

        reviewing_batches = QuoteBatch.query.filter_by(
            project_id=project_id, status="reviewing"
        ).order_by(QuoteBatch.created_at.desc()).all()
        reviewing_batch_counts = {}
        for b in reviewing_batches:
            try:
                reviewing_batch_counts[b.batch_id] = len(json.loads(b.entries_json)) if b.entries_json else 0
            except Exception:
                reviewing_batch_counts[b.batch_id] = 0

        return render_template(
            "drawings.html",
            project=project,
            drawings_list=drawings_list,
            results=results,
            query=query,
            doc_types=DOC_TYPES,
            filter_doc_type=filter_doc_type,
            search_doc_type=search_doc_type,
            report_templates=REPORT_TEMPLATES,
            reports_list=reports_list,
            history_entries=history_entries,
            active_tab=active_tab,
            has_ready_docs=has_ready_docs,
            reviewing_batches=reviewing_batches,
            reviewing_batch_counts=reviewing_batch_counts,
        )

    @app.route("/notes")
    @login_required
    @admin_required
    def notes_library():
        return redirect(url_for("intelligence_library"))

    # ── Intelligence Library ───────────────────────────────────────────────────

    @app.route("/library")
    @login_required
    def intelligence_library():
        tag_filter = request.args.get("tag", "").strip()
        scope_filter = request.args.get("scope", "").strip()
        project_filter = request.args.get("project_id", "").strip()
        search_q = request.args.get("q", "").strip()

        q = IntelligenceItem.query
        if not current_user.is_superadmin:
            q = q.filter(
                db.or_(
                    IntelligenceItem.project_id.is_(None),
                    IntelligenceItem.project_id.in_(
                        db.session.query(Project.id).filter_by(company_id=current_user.company_id)
                    ),
                )
            )
        if tag_filter:
            q = q.join(IntelligenceItem.tags).filter(IntelligenceTag.name == tag_filter)
        if scope_filter:
            q = q.filter(
                db.or_(
                    IntelligenceItem.work_scope_json.is_(None),
                    IntelligenceItem.work_scope_json.contains(scope_filter),
                )
            )
        if project_filter:
            q = q.filter(IntelligenceItem.project_id == int(project_filter))
        if search_q:
            like = f"%{search_q}%"
            q = q.filter(
                db.or_(
                    IntelligenceItem.title.ilike(like),
                    IntelligenceItem.description.ilike(like),
                )
            )
        items = q.order_by(IntelligenceItem.created_at.desc()).all()

        # Legacy Estimation Notes
        lq = db.session.query(Drawing).filter_by(doc_type="Estimation Notes")
        if not current_user.is_superadmin:
            lq = lq.join(Project, Drawing.project_id == Project.id).filter(
                Project.company_id == current_user.company_id
            )
        if not tag_filter and not scope_filter:
            legacy_notes = lq.order_by(Drawing.created_at.desc()).all()
        else:
            legacy_notes = []

        tag_cloud = (
            IntelligenceTag.query
            .filter(IntelligenceTag.usage_count > 0)
            .order_by(IntelligenceTag.usage_count.desc())
            .limit(40)
            .all()
        )

        # Projects for filter dropdown (scoped per user)
        if current_user.is_superadmin:
            all_projects = Project.query.order_by(Project.name).all()
        else:
            all_projects = (
                Project.query
                .filter_by(company_id=current_user.company_id)
                .order_by(Project.name)
                .all()
            )

        filtered_project = None
        if project_filter:
            try:
                filtered_project = db.session.get(Project, int(project_filter))
            except (ValueError, TypeError):
                pass

        return render_template(
            "library.html",
            items=items,
            legacy_notes=legacy_notes,
            tag_cloud=tag_cloud,
            all_projects=all_projects,
            scope_options=WORK_SCOPE_OPTIONS,
            tag_filter=tag_filter,
            scope_filter=scope_filter,
            project_filter=project_filter,
            search_q=search_q,
            filtered_project=filtered_project,
        )

    @app.route("/library/tags")
    @login_required
    @admin_required
    def library_tags():
        """Autocomplete endpoint — returns JSON list of tag names sorted by usage desc."""
        q = request.args.get("q", "").strip()
        query = IntelligenceTag.query
        if q:
            query = query.filter(IntelligenceTag.name.ilike(f"%{q}%"))
        tags = query.order_by(IntelligenceTag.usage_count.desc(), IntelligenceTag.name).limit(20).all()
        return jsonify([t.name for t in tags])

    def _library_projects_for_user():
        """Return (all_projects, include_global) scoped to the current user's role."""
        if current_user.role == ROLE_SUPERADMIN:
            return Project.query.order_by(Project.name).all(), True
        if current_user.role == ROLE_ADMIN:
            return (
                Project.query
                .filter_by(company_id=current_user.company_id)
                .order_by(Project.name)
                .all(),
                True,
            )
        return (
            Project.query
            .filter_by(company_id=current_user.company_id)
            .order_by(Project.name)
            .all(),
            False,
        )

    def _resolve_library_scope(form, project_id):
        """Resolve final project_id given URL project context and scope radio."""
        if current_user.role == ROLE_USER:
            return project_id
        return None if form.get("scope", "project") == "global" else project_id

    @app.route("/library/add", methods=["GET", "POST"])
    @login_required
    def library_add():
        if request.method == "POST":
            title = request.form.get("title", "").strip()
            if not title:
                flash("Title is required.", "danger")
                return redirect(url_for("library_add"))

            description = request.form.get("description", "").strip() or None
            auto_include = request.form.get("auto_include_in_search") == "1"
            work_scope_selected = request.form.getlist("work_scope")
            raw_tags = request.form.get("tags_input", "")

            new_file = request.files.get("library_file")
            if new_file and new_file.filename:
                ext = os.path.splitext(new_file.filename)[1].lower()
                safe_name = uuid.uuid4().hex + ext
                dest = os.path.join(app.config["LIBRARY_FOLDER"], safe_name)
                new_file.save(dest)
                extracted = extract_text_from_file(dest, ext)
                item = IntelligenceItem(
                    title=title,
                    description=description,
                    entry_type="file",
                    text_content=extracted,
                    file_path=safe_name,
                    original_filename=new_file.filename,
                    file_mime=new_file.content_type,
                    project_id=None,
                    work_scope_json=json.dumps(work_scope_selected) if work_scope_selected else None,
                    auto_include_in_search=auto_include,
                    uploaded_by=current_user.id,
                )
            else:
                text_content = request.form.get("text_content", "").strip() or None
                item = IntelligenceItem(
                    title=title,
                    description=description,
                    entry_type="text",
                    text_content=text_content,
                    project_id=None,
                    work_scope_json=json.dumps(work_scope_selected) if work_scope_selected else None,
                    auto_include_in_search=auto_include,
                    uploaded_by=current_user.id,
                )

            db.session.add(item)
            db.session.flush()
            _apply_item_tags(item, raw_tags)
            db.session.commit()
            flash("Global Library entry added.", "success")
            return redirect(url_for("intelligence_library"))

        return render_template(
            "library_item_form.html",
            item=None,
            scope_options=WORK_SCOPE_OPTIONS,
            current_tags="",
            current_scope=[],
            lock_to_global=True,
        )

    @app.route("/projects/<int:project_id>/library/add", methods=["GET", "POST"])
    @login_required
    def library_add_for_project(project_id):
        project = db.session.get(Project, project_id)
        if not project:
            abort(404)
        if not current_user.is_superadmin and project.company_id != current_user.company_id:
            abort(403)

        if request.method == "POST":
            mode = request.form.get("mode", "text")

            if mode == "upload":
                # ── Bulk file upload with classifier routing ─────────────────
                import hashlib
                files = request.files.getlist("library_files")
                raw_tags = request.form.get("tags_input", "")
                work_scope_selected = request.form.getlist("work_scope")
                resolved_project_id = _resolve_library_scope(request.form, project_id)
                api_key = app.config.get("ANTHROPIC_API_KEY") or ""

                library_saved_ids = []
                drawing_ids = []
                quote_files_for_batch = []  # [{dest, safe_name, original_filename}]
                skipped = []

                for f in files:
                    if not f or not f.filename:
                        continue
                    ext = os.path.splitext(f.filename)[1].lower()
                    if ext not in _BULK_UPLOAD_EXTS:
                        skipped.append(f"{f.filename}: unsupported file type")
                        continue

                    # Save to library folder first (temp landing spot)
                    safe_name = uuid.uuid4().hex + ext
                    dest = os.path.join(app.config["LIBRARY_FOLDER"], safe_name)
                    f.save(dest)

                    # Classify (PDFs only; non-PDF always → text)
                    classification = _classify_file_pipeline(api_key, dest, f.filename)
                    pipeline = classification.get("processing_pipeline", "text")
                    doc_type = classification.get("doc_type", DEFAULT_DOC_TYPE)
                    if doc_type not in DOC_TYPES:
                        doc_type = DEFAULT_DOC_TYPE

                    if pipeline == "drawing":
                        # Move to UPLOAD_FOLDER; drawing pipeline worker picks it up
                        drawing_dest = os.path.join(app.config["UPLOAD_FOLDER"], safe_name)
                        os.rename(dest, drawing_dest)
                        drawing = Drawing(
                            filename=safe_name,
                            original_filename=f.filename,
                            project_id=project_id,
                            uploaded_by=current_user.id,
                            doc_type=doc_type,
                            status="pending",
                        )
                        db.session.add(drawing)
                        db.session.flush()
                        drawing_ids.append(drawing.id)
                        # Silently create a paired Library item for text search.
                        # Classifier already identified this as a drawing so pypdf
                        # would yield nothing useful — store vision-only (text_content=None).
                        # Not added to library_saved_ids so it doesn't appear in bulk review.
                        _lib_item = IntelligenceItem(
                            title=os.path.splitext(f.filename)[0],
                            entry_type="text",
                            text_content=None,
                            original_filename=f.filename,
                            project_id=project_id,
                            auto_include_in_search=True,
                            uploaded_by=current_user.id,
                            drawing_id=drawing.id,
                        )
                        db.session.add(_lib_item)
                        db.session.flush()

                    elif pipeline == "quote":
                        # Collect; will be batched below after the loop
                        quote_files_for_batch.append({
                            "dest": dest,
                            "safe_name": safe_name,
                            "original_filename": f.filename,
                        })

                    else:
                        # text pipeline — stays in LIBRARY_FOLDER
                        extracted = extract_text_from_file(dest, ext)
                        item = IntelligenceItem(
                            title=os.path.splitext(f.filename)[0],
                            entry_type="file",
                            text_content=extracted,
                            file_path=safe_name,
                            original_filename=f.filename,
                            file_mime=f.content_type,
                            project_id=resolved_project_id,
                            work_scope_json=json.dumps(work_scope_selected) if (resolved_project_id is None and work_scope_selected) else None,
                            auto_include_in_search=True,
                            uploaded_by=current_user.id,
                        )
                        db.session.add(item)
                        db.session.flush()
                        _apply_item_tags(item, raw_tags)
                        library_saved_ids.append(item.id)

                # ── Process quote files as a single QuoteBatch ────────────────
                quote_batch_ref = None
                if quote_files_for_batch:
                    quote_batch_id = str(uuid.uuid4())
                    staging_dir = os.path.join(
                        app.config["LIBRARY_FOLDER"], "quotes_staging", quote_batch_id
                    )
                    os.makedirs(staging_dir, exist_ok=True)
                    entries = []
                    for qf in quote_files_for_batch:
                        staging_dest = os.path.join(staging_dir, qf["safe_name"])
                        os.rename(qf["dest"], staging_dest)
                        with open(staging_dest, "rb") as fh:
                            content_hash = hashlib.sha256(fh.read()).hexdigest()
                        entry = _extract_quote_file(api_key, staging_dest, qf["original_filename"])
                        entry["staged_filename"] = qf["safe_name"]
                        entry["content_hash"] = content_hash
                        existing = IntelligenceItem.query.filter_by(
                            project_id=project_id, content_hash=content_hash,
                        ).first()
                        if existing:
                            entry["is_duplicate_of_existing"] = True
                            entry["existing_item_id"] = existing.id
                            entry["existing_item_title"] = existing.title
                            entry["existing_item_uploaded_date"] = existing.created_at.strftime("%b %d, %Y")
                        entries.append(entry)
                    active_takeoff = _get_active_takeoff(project_id)
                    quote_batch = QuoteBatch(
                        batch_id=quote_batch_id,
                        project_id=project_id,
                        user_id=current_user.id,
                        status="reviewing",
                        category_tag=None,
                        entries_json=json.dumps({"duplicates_within_session": 0, "entries": entries}),
                        takeoff_id=active_takeoff.id if active_takeoff else None,
                    )
                    db.session.add(quote_batch)
                    db.session.flush()
                    quote_batch_ref = {
                        "batch_id": quote_batch_id,
                        "filenames": [qf["original_filename"] for qf in quote_files_for_batch],
                    }

                db.session.commit()

                for msg in skipped:
                    flash(msg, "warning")

                if not library_saved_ids and not drawing_ids and not quote_batch_ref:
                    flash("No valid files were uploaded.", "danger")
                    return redirect(url_for("library_add_for_project", project_id=project_id))

                batch_id = uuid.uuid4().hex
                session[f"lib_batch_{batch_id}"] = {
                    "item_ids": library_saved_ids,
                    "drawing_ids": drawing_ids,
                    "quote_batches": [quote_batch_ref] if quote_batch_ref else [],
                    "project_id": project_id,
                    "default_scope": request.form.get("scope", "project"),
                }
                return redirect(url_for("library_bulk_review", batch_id=batch_id))

            else:
                # ── Typed entry ───────────────────────────────────────────────
                title = request.form.get("title", "").strip()
                if not title:
                    flash("Title is required.", "danger")
                    return redirect(url_for("library_add_for_project", project_id=project_id))

                description = request.form.get("description", "").strip() or None
                text_content = request.form.get("text_content", "").strip() or None
                work_scope_selected = request.form.getlist("work_scope")
                auto_include = request.form.get("auto_include_in_search") == "1"
                resolved_project_id = _resolve_library_scope(request.form, project_id)
                raw_tags = request.form.get("tags_input", "")

                item = IntelligenceItem(
                    title=title,
                    description=description,
                    entry_type="text",
                    text_content=text_content,
                    project_id=resolved_project_id,
                    work_scope_json=json.dumps(work_scope_selected) if (resolved_project_id is None and work_scope_selected) else None,
                    auto_include_in_search=auto_include,
                    uploaded_by=current_user.id,
                )
                db.session.add(item)
                db.session.flush()
                _apply_item_tags(item, raw_tags)
                db.session.commit()
                flash("Intelligence Library item added.", "success")
                return redirect(url_for("intelligence_library"))

        # GET
        return render_template(
            "library_add.html",
            current_project=project,
            scope_options=WORK_SCOPE_OPTIONS,
            lock_to_project=True,
        )

    def _batch_item_ids(batch_data):
        """Extract library item_ids from session batch (supports old list and new dict format)."""
        if isinstance(batch_data, dict):
            return batch_data.get("item_ids", [])
        return batch_data or []

    def _auto_accept_quote_batch(batch, project_id):
        """Create IntelligenceItems from a QuoteBatch without form data, then mark saved."""
        import shutil as _shutil

        def _parse_d(val):
            if not val:
                return None
            try:
                from datetime import date as _date
                return _date.fromisoformat(str(val)[:10])
            except Exception:
                return None

        staging_dir = os.path.join(app.config["LIBRARY_FOLDER"], "quotes_staging", batch.batch_id)
        raw = json.loads(batch.entries_json or "[]")
        entries = raw.get("entries", raw) if isinstance(raw, dict) else raw
        saved = 0
        for entry in entries:
            if entry.get("error"):
                continue
            staged_fn = entry.get("staged_filename", "")
            original_fn = entry.get("original_filename", staged_fn)
            title = (entry.get("title") or "").strip() or original_fn or "Untitled Quote"
            new_file_path = None
            file_mime = None
            if staged_fn:
                src = os.path.join(staging_dir, staged_fn)
                if os.path.exists(src):
                    ext = os.path.splitext(staged_fn)[1]
                    dest_name = uuid.uuid4().hex + ext
                    dest = os.path.join(app.config["LIBRARY_FOLDER"], dest_name)
                    os.rename(src, dest)
                    new_file_path = dest_name
                    ext_lower = ext.lower()
                    if ext_lower == ".pdf":
                        file_mime = "application/pdf"
                    elif ext_lower == ".png":
                        file_mime = "image/png"
                    elif ext_lower in (".jpg", ".jpeg"):
                        file_mime = "image/jpeg"
            pricing_rows = entry.get("pricing_items") or []
            valid_flags = [f for f in (entry.get("flags") or []) if f in _ALLOWED_FLAGS]
            intel_item = IntelligenceItem(
                title=title,
                description=entry.get("description") or None,
                entry_type="file",
                file_path=new_file_path,
                original_filename=original_fn,
                file_mime=file_mime,
                project_id=project_id,
                auto_include_in_search=True,
                uploaded_by=current_user.id,
                vendor_name=entry.get("vendor_name") or None,
                quote_date=_parse_d(entry.get("quote_date")),
                expiration_date=_parse_d(entry.get("expiration_date")),
                conditions_text=entry.get("conditions_text") or None,
                pricing_items_json=json.dumps(pricing_rows) if pricing_rows else None,
                flags_json=json.dumps(valid_flags) if valid_flags else None,
                extraction_status="auto-extracted",
                content_hash=entry.get("content_hash"),
            )
            db.session.add(intel_item)
            db.session.flush()
            if batch.category_tag:
                _apply_item_tags(intel_item, batch.category_tag)
            saved += 1
        batch.status = "saved"
        db.session.commit()
        if os.path.isdir(staging_dir):
            _shutil.rmtree(staging_dir, ignore_errors=True)
        return saved

    @app.route("/library/bulk-review/<batch_id>")
    @login_required
    def library_bulk_review(batch_id):
        batch_data = session.get(f"lib_batch_{batch_id}")
        item_ids = _batch_item_ids(batch_data)
        drawing_ids = batch_data.get("drawing_ids", []) if isinstance(batch_data, dict) else []
        quote_batch_refs = batch_data.get("quote_batches", []) if isinstance(batch_data, dict) else []

        if not item_ids and not drawing_ids and not quote_batch_refs:
            flash("Upload session not found or expired. Please upload again.", "warning")
            return redirect(url_for("intelligence_library"))

        items = (
            IntelligenceItem.query
            .filter(
                IntelligenceItem.id.in_(item_ids),
                IntelligenceItem.uploaded_by == current_user.id,
            )
            .all()
        ) if item_ids else []

        file_sizes = {}
        for item in items:
            if item.file_path:
                full = os.path.join(app.config["LIBRARY_FOLDER"], item.file_path)
                try:
                    file_sizes[item.id] = os.path.getsize(full)
                except OSError:
                    file_sizes[item.id] = None

        drawings = [d for did in drawing_ids for d in [db.session.get(Drawing, did)] if d]

        context_project_id = batch_data.get("project_id") if isinstance(batch_data, dict) else None
        context_project = db.session.get(Project, context_project_id) if context_project_id else None

        project_categories = _project_quote_categories(context_project_id) if context_project_id else []
        existing_cat_names = {c["name"] for c in project_categories}
        scope_suggestions = [s for s in WORK_SCOPE_OPTIONS if s not in existing_cat_names]

        enriched_quotes = []
        for qbr in quote_batch_refs:
            qb = QuoteBatch.query.filter_by(batch_id=qbr["batch_id"]).first()
            qb_raw = json.loads(qb.entries_json or "[]") if qb else {}
            qb_entries = qb_raw.get("entries", qb_raw) if isinstance(qb_raw, dict) else qb_raw
            # Derive a suggested category from the first entry's suggested_tags
            first_entry = qb_entries[0] if qb_entries else {}
            ai_suggestions = first_entry.get("suggested_tags") or []
            suggested_cat = ai_suggestions[0].title() if ai_suggestions else None
            enriched_quotes.append({
                "batch_id": qbr["batch_id"],
                "filenames": qbr["filenames"],
                "status": qb.status if qb else "unknown",
                "entries": qb_entries,
                "category_tag": qb.category_tag if qb else None,
                "suggested_category": suggested_cat,
            })

        default_scope = batch_data.get("default_scope", "project") if isinstance(batch_data, dict) else "project"

        return render_template(
            "library_bulk_review.html",
            batch_id=batch_id,
            items=items,
            file_sizes=file_sizes,
            drawings=drawings,
            quote_batch_refs=enriched_quotes,
            context_project_id=context_project_id,
            scope_options=WORK_SCOPE_OPTIONS,
            context_project=context_project,
            project_categories=project_categories,
            scope_suggestions=scope_suggestions,
            default_scope=default_scope,
        )

    @app.route("/library/bulk-review/<batch_id>/save", methods=["POST"])
    @login_required
    def library_bulk_review_save(batch_id):
        batch_data = session.get(f"lib_batch_{batch_id}")
        item_ids = _batch_item_ids(batch_data)
        context_project_id = batch_data.get("project_id") if isinstance(batch_data, dict) else None

        items = (
            IntelligenceItem.query
            .filter(
                IntelligenceItem.id.in_(item_ids),
                IntelligenceItem.uploaded_by == current_user.id,
            )
            .all()
        ) if item_ids else []

        for item in items:
            sid = str(item.id)
            title = request.form.get(f"title_{sid}", "").strip()
            if title:
                item.title = title
            item.description = request.form.get(f"description_{sid}", "").strip() or None
            scope_val = request.form.get(f"scope_{sid}", "project")
            if current_user.role == ROLE_USER:
                item.project_id = context_project_id
            else:
                item.project_id = None if scope_val == "global" else context_project_id
            work_scope = request.form.getlist(f"work_scope_{sid}")
            item.work_scope_json = json.dumps(work_scope) if (item.project_id is None and work_scope) else None
            raw_tags = request.form.get(f"tags_{sid}", "")
            _apply_item_tags(item, raw_tags)

        if items:
            db.session.commit()

        session.pop(f"lib_batch_{batch_id}", None)
        n = len(items)
        if n:
            flash(f"{n} item{'s' if n != 1 else ''} saved to the Intelligence Library.", "success")
        else:
            flash("Files dispatched to their pipelines.", "success")
        return redirect(url_for("intelligence_library"))

    @app.route("/library/bulk-review/<batch_id>/discard", methods=["POST"])
    @login_required
    def library_bulk_review_discard(batch_id):
        batch_data = session.get(f"lib_batch_{batch_id}")
        item_ids = _batch_item_ids(batch_data)
        drawing_ids = batch_data.get("drawing_ids", []) if isinstance(batch_data, dict) else []
        quote_batch_refs = batch_data.get("quote_batches", []) if isinstance(batch_data, dict) else []

        if item_ids:
            items = (
                IntelligenceItem.query
                .filter(
                    IntelligenceItem.id.in_(item_ids),
                    IntelligenceItem.uploaded_by == current_user.id,
                )
                .all()
            )
            for item in items:
                if item.file_path:
                    fpath = os.path.join(app.config["LIBRARY_FOLDER"], item.file_path)
                    if os.path.exists(fpath):
                        os.remove(fpath)
                _apply_item_tags(item, "")
                db.session.delete(item)
            db.session.commit()

        session.pop(f"lib_batch_{batch_id}", None)
        if drawing_ids or quote_batch_refs:
            flash(
                "Library entries discarded. "
                "Drawings and quotes already dispatched cannot be discarded from this screen — "
                "manage them from their respective cards.",
                "info",
            )
        else:
            flash("Upload discarded. Files have been removed.", "info")
        return redirect(url_for("intelligence_library"))

    @app.route("/library/<int:item_id>/edit", methods=["GET", "POST"])
    @login_required
    @admin_required
    def library_edit(item_id):
        item = db.session.get(IntelligenceItem, item_id) or abort(404)
        if item.project_id is None and current_user.role not in (ROLE_ADMIN, ROLE_SUPERADMIN):
            abort(403)
        if not current_user.is_superadmin and item.project_id:
            proj = db.session.get(Project, item.project_id)
            if proj and proj.company_id != current_user.company_id:
                abort(403)

        if request.method == "POST":
            title = request.form.get("title", "").strip()
            if not title:
                flash("Title is required.", "danger")
                return redirect(url_for("library_edit", item_id=item_id))

            item.title = title
            item.description = request.form.get("description", "").strip() or None
            item.auto_include_in_search = request.form.get("auto_include_in_search") == "1"
            if current_user.role != ROLE_USER:
                scope_val = request.form.get("scope", "project")
                current_proj_id_raw = request.form.get("current_project_id", "").strip()
                if scope_val == "global":
                    item.project_id = None
                elif current_proj_id_raw.isdigit():
                    item.project_id = int(current_proj_id_raw)
            work_scope_selected = request.form.getlist("work_scope")
            item.work_scope_json = json.dumps(work_scope_selected) if (item.project_id is None and work_scope_selected) else None

            if item.entry_type != "text":
                new_file = request.files.get("library_file")
                if new_file and new_file.filename:
                    if item.file_path:
                        old = os.path.join(app.config["LIBRARY_FOLDER"], item.file_path)
                        if os.path.exists(old):
                            os.remove(old)
                    ext = os.path.splitext(new_file.filename)[1]
                    safe_name = uuid.uuid4().hex + ext
                    dest = os.path.join(app.config["LIBRARY_FOLDER"], safe_name)
                    new_file.save(dest)
                    item.file_path = safe_name
                    item.original_filename = new_file.filename
                    item.file_mime = new_file.content_type

            old_tags = list(item.tags)
            raw_tags = request.form.get("tags_input", "")
            _apply_item_tags(item, raw_tags)
            _decrement_removed_tags(old_tags, item.tags)
            db.session.commit()
            flash("Item updated.", "success")
            return redirect(url_for("intelligence_library"))

        return render_template(
            "library_item_form.html",
            item=item,
            scope_options=WORK_SCOPE_OPTIONS,
            current_tags=", ".join(t.name for t in item.tags),
            current_scope=item.work_scope_list,
        )

    @app.route("/library/<int:item_id>/delete", methods=["POST"])
    @login_required
    @admin_required
    def library_delete(item_id):
        item = db.session.get(IntelligenceItem, item_id) or abort(404)
        if item.project_id is None and current_user.role not in (ROLE_ADMIN, ROLE_SUPERADMIN):
            abort(403)
        if not current_user.is_superadmin and item.project_id:
            proj = db.session.get(Project, item.project_id)
            if proj and proj.company_id != current_user.company_id:
                abort(403)
        if item.file_path:
            fpath = os.path.join(app.config["LIBRARY_FOLDER"], item.file_path)
            if os.path.exists(fpath):
                os.remove(fpath)
        for tag in list(item.tags):
            if tag.usage_count > 0:
                tag.usage_count -= 1
        db.session.delete(item)
        db.session.commit()
        flash("Item deleted.", "success")
        return redirect(url_for("intelligence_library"))

    @app.route("/library/<int:item_id>/download")
    @login_required
    def library_download(item_id):
        item = db.session.get(IntelligenceItem, item_id) or abort(404)
        if not item.file_path:
            abort(404)
        if not current_user.is_superadmin and item.project_id:
            proj = db.session.get(Project, item.project_id)
            if proj and proj.company_id != current_user.company_id:
                abort(403)
        return send_from_directory(
            app.config["LIBRARY_FOLDER"],
            item.file_path,
            as_attachment=True,
            download_name=item.original_filename or item.file_path,
        )

    # ── Bulk Quote Intake ─────────────────────────────────────────────────────

    @app.route("/projects/<int:project_id>/quotes/bulk-intake", methods=["GET"])
    @login_required
    @admin_required
    def bulk_quote_intake(project_id):
        project = db.session.get(Project, project_id) or abort(404)
        if not current_user.is_superadmin and current_user.company_id != project.company_id:
            abort(403)
        flash(
            "The standalone Quotes uploader has moved. Drop quote files into the Library — "
            "PowerScan will route them automatically.",
            "info",
        )
        return redirect(url_for("library_add_for_project", project_id=project_id))

    @app.route("/projects/<int:project_id>/quotes/bulk-intake", methods=["POST"])
    @login_required
    @admin_required
    def bulk_quote_intake_post(project_id):
        return redirect(url_for("bulk_quote_intake", project_id=project_id))

    @app.route("/quotes/bulk-upload")
    @login_required
    def quotes_bulk_upload_redirect():
        flash(
            "The standalone Quotes uploader has moved. Drop quote files into the Library — "
            "PowerScan will route them automatically.",
            "info",
        )
        return redirect(url_for("intelligence_library"))

    @app.route("/projects/<int:project_id>/quotes/bulk-intake/review/<batch_id>", methods=["GET"])
    @login_required
    @admin_required
    def bulk_quote_review(project_id, batch_id):
        project = db.session.get(Project, project_id) or abort(404)
        if not current_user.is_superadmin and current_user.company_id != project.company_id:
            abort(403)
        batch = QuoteBatch.query.filter_by(batch_id=batch_id, project_id=project_id).first() or abort(404)
        if batch.user_id != current_user.id and not current_user.is_superadmin:
            abort(403)
        raw = json.loads(batch.entries_json or "[]")
        if isinstance(raw, list):
            entries = raw
            duplicates_within_session = 0
        else:
            entries = raw.get("entries", [])
            duplicates_within_session = raw.get("duplicates_within_session", 0)
        return render_template(
            "bulk_quote_review.html",
            project=project,
            batch=batch,
            entries=entries,
            allowed_flags=_ALLOWED_FLAGS,
            entry_count=len(entries),
            duplicates_within_session=duplicates_within_session,
        )

    @app.route("/projects/<int:project_id>/quotes/bulk-intake/review/<batch_id>/save", methods=["POST"])
    @login_required
    @admin_required
    def bulk_quote_save(project_id, batch_id):
        project = db.session.get(Project, project_id) or abort(404)
        if not current_user.is_superadmin and current_user.company_id != project.company_id:
            abort(403)
        batch = QuoteBatch.query.filter_by(batch_id=batch_id, project_id=project_id).first() or abort(404)
        if batch.user_id != current_user.id and not current_user.is_superadmin:
            abort(403)

        staging_dir = os.path.join(app.config["LIBRARY_FOLDER"], "quotes_staging", batch_id)
        raw = json.loads(batch.entries_json or "[]")
        if isinstance(raw, list):
            entries = raw
        else:
            entries = raw.get("entries", [])
        entry_count = int(request.form.get("entry_count", len(entries)))
        saved = 0
        skipped_dupes = 0
        replaced_dupes = 0

        # Server-side validation: every non-removed duplicate must have an action
        missing_actions = []
        for i in range(entry_count):
            entry = entries[i] if i < len(entries) else {}
            if request.form.get(f"entry_{i}_remove"):
                continue
            if entry.get("is_duplicate_of_existing") and not request.form.get(f"duplicate_action_{i}"):
                missing_actions.append(i + 1)
        if missing_actions:
            flash(
                f"Please choose an action for duplicate file(s) at position(s): {', '.join(str(n) for n in missing_actions)}.",
                "danger",
            )
            return redirect(url_for("bulk_quote_review", project_id=project_id, batch_id=batch_id))

        for i in range(entry_count):
            prefix = f"entry_{i}_"
            entry = entries[i] if i < len(entries) else {}

            if request.form.get(prefix + "remove"):
                continue

            # Handle duplicate decision before the rest of the save logic
            if entry.get("is_duplicate_of_existing"):
                dup_action = request.form.get(f"duplicate_action_{i}", "skip")
                if dup_action == "skip":
                    staged_fn = request.form.get(prefix + "staged_filename", "")
                    if staged_fn:
                        staged_path = os.path.join(staging_dir, staged_fn)
                        if os.path.exists(staged_path):
                            os.remove(staged_path)
                    skipped_dupes += 1
                    continue
                if dup_action == "replace":
                    old_id = entry.get("existing_item_id")
                    if old_id:
                        old_item = db.session.get(IntelligenceItem, int(old_id))
                        if old_item:
                            if old_item.file_path:
                                old_fpath = os.path.join(app.config["LIBRARY_FOLDER"], old_item.file_path)
                                if os.path.exists(old_fpath):
                                    os.remove(old_fpath)
                            for tag in list(old_item.tags):
                                if tag.usage_count > 0:
                                    tag.usage_count -= 1
                            db.session.delete(old_item)
                            db.session.flush()
                    replaced_dupes += 1
                # "upload" or post-replace: fall through to normal save

            staged_filename = request.form.get(prefix + "staged_filename", "")
            original_filename = request.form.get(prefix + "original_filename", staged_filename)
            title = request.form.get(prefix + "title", "").strip()
            if not title:
                continue

            # Move file from staging to main library folder
            new_file_path = None
            file_mime = None
            if staged_filename:
                src = os.path.join(staging_dir, staged_filename)
                if os.path.exists(src):
                    ext = os.path.splitext(staged_filename)[1]
                    dest_name = uuid.uuid4().hex + ext
                    dest = os.path.join(app.config["LIBRARY_FOLDER"], dest_name)
                    os.rename(src, dest)
                    new_file_path = dest_name
                    ext_lower = ext.lower()
                    if ext_lower == ".pdf":
                        file_mime = "application/pdf"
                    elif ext_lower in (".png",):
                        file_mime = "image/png"
                    elif ext_lower in (".jpg", ".jpeg"):
                        file_mime = "image/jpeg"

            description = request.form.get(prefix + "description", "").strip() or None
            vendor_name = request.form.get(prefix + "vendor_name", "").strip() or None
            vendor_contact = request.form.get(prefix + "vendor_contact", "").strip() or None
            conditions = request.form.get(prefix + "conditions_text", "").strip() or None
            extraction_status = request.form.get(prefix + "extraction_status", "auto-extracted")

            def _parse_date(val):
                val = (val or "").strip()
                if not val:
                    return None
                try:
                    return date.fromisoformat(val)
                except ValueError:
                    return None

            qdate = _parse_date(request.form.get(prefix + "quote_date"))
            expdate = _parse_date(request.form.get(prefix + "expiration_date"))

            # Pricing items — collect indexed rows
            pricing_rows = []
            j = 0
            while True:
                lbl = request.form.get(f"{prefix}pricing_{j}_label", "")
                amt = request.form.get(f"{prefix}pricing_{j}_amount", "")
                unit = request.form.get(f"{prefix}pricing_{j}_unit", "")
                notes = request.form.get(f"{prefix}pricing_{j}_notes", "")
                if not any([lbl, amt, unit, notes]):
                    break
                pricing_rows.append({
                    "label": lbl.strip(),
                    "amount": amt.strip(),
                    "unit": unit.strip(),
                    "notes": notes.strip() or None,
                })
                j += 1
                if j > 100:
                    break

            selected_flags = request.form.getlist(prefix + "flags")
            valid_flags = [f for f in selected_flags if f in _ALLOWED_FLAGS]

            item = IntelligenceItem(
                title=title,
                description=description,
                entry_type="file",
                file_path=new_file_path,
                original_filename=original_filename,
                file_mime=file_mime,
                project_id=project_id,
                auto_include_in_search=True,
                uploaded_by=current_user.id,
                vendor_name=vendor_name,
                vendor_contact=vendor_contact,
                quote_date=qdate,
                expiration_date=expdate,
                conditions_text=conditions,
                pricing_items_json=json.dumps(pricing_rows) if pricing_rows else None,
                flags_json=json.dumps(valid_flags) if valid_flags else None,
                extraction_status=extraction_status,
                content_hash=entry.get("content_hash"),
            )
            db.session.add(item)
            db.session.flush()

            # Category tag + additional tags
            all_tags_parts = []
            if batch.category_tag:
                category_tag_val = request.form.get(prefix + "category_tag", batch.category_tag).strip()
                if category_tag_val:
                    all_tags_parts.append(category_tag_val)
            extra_tags = request.form.get(prefix + "extra_tags", "").strip()
            if extra_tags:
                all_tags_parts.extend([t.strip() for t in extra_tags.split(",") if t.strip()])
            if all_tags_parts:
                _apply_item_tags(item, ", ".join(all_tags_parts))

            saved += 1

        batch.status = "saved"
        db.session.commit()

        # Clean up any leftover staging files
        import shutil
        if os.path.isdir(staging_dir):
            shutil.rmtree(staging_dir, ignore_errors=True)

        parts = [f"{saved} quote{'s' if saved != 1 else ''} saved to the Intelligence Library."]
        if skipped_dupes:
            parts.append(f"Skipped {skipped_dupes} duplicate{'s' if skipped_dupes != 1 else ''}.")
        if replaced_dupes:
            parts.append(f"Replaced {replaced_dupes} existing {'entries' if replaced_dupes != 1 else 'entry'}.")
        flash(" ".join(parts), "success")
        return redirect(url_for("intelligence_library") + f"?project_id={project_id}")

    @app.route("/projects/<int:project_id>/quotes/bulk-intake/review/<batch_id>/cancel", methods=["POST"])
    @login_required
    @admin_required
    def bulk_quote_cancel(project_id, batch_id):
        project = db.session.get(Project, project_id) or abort(404)
        if not current_user.is_superadmin and current_user.company_id != project.company_id:
            abort(403)
        batch = QuoteBatch.query.filter_by(batch_id=batch_id, project_id=project_id).first() or abort(404)
        if batch.user_id != current_user.id and not current_user.is_superadmin:
            abort(403)

        import shutil
        staging_dir = os.path.join(app.config["LIBRARY_FOLDER"], "quotes_staging", batch_id)
        if os.path.isdir(staging_dir):
            shutil.rmtree(staging_dir, ignore_errors=True)

        batch.status = "cancelled"
        db.session.commit()
        flash("Batch cancelled.", "info")
        return redirect(url_for("project_hub", project_id=project_id))

    # ── Library bulk-review: inline quote accept/discard ──────────────────────

    @app.route("/library/bulk-review/<lib_batch_id>/quote/<quote_batch_id>/accept", methods=["POST"])
    @login_required
    def library_bulk_review_quote_accept(lib_batch_id, quote_batch_id):
        batch = QuoteBatch.query.filter_by(batch_id=quote_batch_id).first() or abort(404)
        if not current_user.is_superadmin:
            proj = db.session.get(Project, batch.project_id)
            if not proj or proj.company_id != current_user.company_id:
                abort(403)
        category_tag = request.form.get("category_tag", "").strip()
        if not category_tag:
            flash("Please select a category before accepting.", "danger")
            return redirect(url_for("library_bulk_review", batch_id=lib_batch_id))
        if batch.status == "reviewing":
            batch.category_tag = category_tag
            n = _auto_accept_quote_batch(batch, batch.project_id)
            flash(f"{n} quote{'s' if n != 1 else ''} accepted under '{category_tag}'.", "success")
        return redirect(url_for("library_bulk_review", batch_id=lib_batch_id))

    @app.route("/library/bulk-review/<lib_batch_id>/quote/<quote_batch_id>/discard", methods=["POST"])
    @login_required
    def library_bulk_review_quote_discard(lib_batch_id, quote_batch_id):
        batch = QuoteBatch.query.filter_by(batch_id=quote_batch_id).first() or abort(404)
        if not current_user.is_superadmin:
            proj = db.session.get(Project, batch.project_id)
            if not proj or proj.company_id != current_user.company_id:
                abort(403)
        if batch.status == "reviewing":
            import shutil
            staging_dir = os.path.join(app.config["LIBRARY_FOLDER"], "quotes_staging", quote_batch_id)
            if os.path.isdir(staging_dir):
                shutil.rmtree(staging_dir, ignore_errors=True)
            batch.status = "cancelled"
            db.session.commit()
            flash("Quote batch discarded.", "info")
        return redirect(url_for("library_bulk_review", batch_id=lib_batch_id))

    @app.route("/library/bulk-review/<lib_batch_id>/quotes/accept-all", methods=["POST"])
    @login_required
    def library_bulk_review_quotes_accept_all(lib_batch_id):
        batch_data = session.get(f"lib_batch_{lib_batch_id}")
        if not batch_data:
            flash("Session expired.", "warning")
            return redirect(url_for("intelligence_library"))
        total_saved = 0
        skipped_no_cat = 0
        for qbr in batch_data.get("quote_batches", []):
            qb = QuoteBatch.query.filter_by(batch_id=qbr["batch_id"]).first()
            if not qb or qb.status != "reviewing":
                continue
            if not current_user.is_superadmin:
                proj = db.session.get(Project, qb.project_id)
                if not proj or proj.company_id != current_user.company_id:
                    continue
            cat = request.form.get(f"category_{qbr['batch_id']}", "").strip()
            if not cat:
                skipped_no_cat += 1
                continue
            qb.category_tag = cat
            total_saved += _auto_accept_quote_batch(qb, qb.project_id)
        parts = []
        if total_saved:
            parts.append(f"{total_saved} quote{'s' if total_saved != 1 else ''} accepted and saved to the Intelligence Library.")
        if skipped_no_cat:
            parts.append(f"{skipped_no_cat} batch{'es' if skipped_no_cat != 1 else ''} skipped — no category selected.")
        flash(" ".join(parts) if parts else "Nothing to accept.", "success" if total_saved else "warning")
        return redirect(url_for("library_bulk_review", batch_id=lib_batch_id))

    # ── Quote category management ─────────────────────────────────────────────

    @app.route("/projects/<int:project_id>/quotes/item/<int:item_id>/category", methods=["POST"])
    @login_required
    @admin_required
    def quote_item_set_category(project_id, item_id):
        _auth_project(project_id)
        item = db.session.get(IntelligenceItem, item_id) or abort(404)
        if item.project_id != project_id:
            abort(403)
        new_cat = request.form.get("category", "").strip()
        if not new_cat:
            flash("Category is required.", "danger")
            return redirect(url_for("quote_compare_select", project_id=project_id))
        _update_item_category(item, project_id, new_cat)
        db.session.commit()
        flash(f"Quote moved to '{new_cat}'.", "success")
        return redirect(url_for("quote_compare_select", project_id=project_id))

    @app.route("/projects/<int:project_id>/quotes/categories/merge", methods=["POST"])
    @login_required
    @admin_required
    def quote_categories_merge(project_id):
        _auth_project(project_id)
        source = request.form.get("source_category", "").strip()
        target = request.form.get("target_category", "").strip()
        if not source or not target or source == target:
            flash("Select two different categories to merge.", "danger")
            return redirect(url_for("quote_compare_select", project_id=project_id))
        items = _items_for_category(project_id, source)
        n = len(items)
        for item in items:
            _update_item_category(item, project_id, target)
        for b in QuoteBatch.query.filter_by(project_id=project_id, category_tag=source).all():
            b.category_tag = target
        db.session.commit()
        flash(f"Merged {n} quote{'s' if n != 1 else ''} from '{source}' into '{target}'.", "success")
        return redirect(url_for("quote_compare_select", project_id=project_id))

    # ── Quote Comparison ──────────────────────────────────────────────────────

    def _auth_project(project_id):
        """Return project or abort. Enforce company scope for non-superadmin."""
        project = db.session.get(Project, project_id) or abort(404)
        if not current_user.is_superadmin and current_user.company_id != project.company_id:
            abort(403)
        return project

    def _items_for_category(project_id, category_tag):
        """Return IntelligenceItems for this project that have the given tag."""
        return (
            IntelligenceItem.query
            .filter_by(project_id=project_id)
            .join(IntelligenceItem.tags)
            .filter(IntelligenceTag.name == category_tag)
            .all()
        )

    def _project_quote_categories(project_id):
        """Return [{name, count}] for all saved quote categories in this project, sorted alpha."""
        batches = QuoteBatch.query.filter_by(project_id=project_id, status="saved").all()
        seen = set()
        names = []
        for b in batches:
            tag = (b.category_tag or "").strip()
            if tag and tag not in seen:
                seen.add(tag)
                names.append(tag)
        result = []
        for name in sorted(names, key=str.lower):
            count = len(_items_for_category(project_id, name))
            result.append({"name": name, "count": count})
        return result

    def _update_item_category(item, project_id, new_category_name):
        """Replace the category tag on an IntelligenceItem without touching other tags."""
        saved_cats = {
            b.category_tag for b in
            QuoteBatch.query.filter_by(project_id=project_id, status="saved").all()
            if b.category_tag
        }
        for tag in list(item.tags):
            if tag.name in saved_cats:
                item.tags.remove(tag)
                if tag.usage_count > 0:
                    tag.usage_count -= 1
        new_tag = IntelligenceTag.query.filter_by(name=new_category_name).first()
        if not new_tag:
            new_tag = IntelligenceTag(name=new_category_name, usage_count=0)
            db.session.add(new_tag)
            db.session.flush()
        if new_tag not in item.tags:
            item.tags.append(new_tag)
            new_tag.usage_count += 1

    def _build_pricing_matrix(items):
        """Return (all_labels, per_item_pricing) for the comparison table.

        all_labels  — ordered list of every unique pricing label across all items
        per_item_pricing — dict mapping item.id → dict(label → {amount,unit,notes})
        """
        label_order = []
        seen_labels = set()
        per_item = {}
        for item in items:
            rows = []
            if item.pricing_items_json:
                try:
                    rows = json.loads(item.pricing_items_json)
                except Exception:
                    rows = []
            per_item[item.id] = {}
            for r in rows:
                lbl = (r.get("label") or "").strip()
                if not lbl:
                    continue
                per_item[item.id][lbl] = r
                if lbl not in seen_labels:
                    seen_labels.add(lbl)
                    label_order.append(lbl)
        return label_order, per_item

    @app.route("/projects/<int:project_id>/quotes/compare")
    @login_required
    @admin_required
    def quote_compare_select(project_id):
        project = _auth_project(project_id)

        # Distinct saved category_tags from QuoteBatch — the clean label the
        # user chose at intake time, not free-form tags from IntelligenceItem.
        saved_batches = (
            QuoteBatch.query
            .filter_by(project_id=project_id, status="saved")
            .all()
        )
        seen = set()
        ordered_tags = []
        for b in saved_batches:
            tag = (b.category_tag or "").strip()
            if tag and tag not in seen:
                seen.add(tag)
                ordered_tags.append(tag)

        def _headline_price(item):
            if not item.pricing_items_json:
                return "no pricing submitted"
            try:
                rows = json.loads(item.pricing_items_json)
            except Exception:
                return "no pricing submitted"
            if not rows:
                return "no pricing submitted"
            first = rows[0]
            amt = (first.get("amount") or "").strip().lstrip("$")
            unit = (first.get("unit") or "").strip()
            if not amt:
                return "no pricing submitted"
            return f"${amt} {unit}".strip() if unit else f"${amt}"

        def _parse_flags(item):
            if not item.flags_json:
                return []
            try:
                return json.loads(item.flags_json) or []
            except Exception:
                return []

        # Index batches by category_tag for delete button (take the earliest batch per tag)
        batch_by_tag = {}
        for b in saved_batches:
            tag = (b.category_tag or "").strip()
            if tag and tag not in batch_by_tag:
                batch_by_tag[tag] = b.id

        categories = []
        for tag in ordered_tags:
            items = _items_for_category(project_id, tag)
            cat_items = []
            for it in items:
                cat_items.append({
                    "id": it.id,
                    "vendor": it.vendor_name or it.title,
                    "headline_price": _headline_price(it),
                    "flags": _parse_flags(it),
                    "file_path": it.file_path or None,
                })
            categories.append({
                "name": tag,
                "count": len(items),
                "quotes": cat_items,
                "batch_id": batch_by_tag.get(tag),
            })

        # Sort: most vendors first, then alpha by name
        categories.sort(key=lambda c: (-c["count"], c["name"].lower()))

        project_categories = _project_quote_categories(project_id)
        existing_cat_names = {c["name"] for c in project_categories}
        scope_suggestions = [s for s in WORK_SCOPE_OPTIONS if s not in existing_cat_names]

        return render_template(
            "quote_compare_select.html",
            project=project,
            categories=categories,
            project_categories=project_categories,
            scope_suggestions=scope_suggestions,
        )

    @app.route("/quote-batch/<int:batch_id>/delete", methods=["POST"])
    @login_required
    @admin_required
    def delete_quote_batch(batch_id):
        import shutil
        batch = db.session.get(QuoteBatch, batch_id) or abort(404)
        project = db.session.get(Project, batch.project_id) or abort(404)
        if not current_user.is_superadmin and current_user.company_id != project.company_id:
            abort(403)

        project_id = batch.project_id
        category_tag = batch.category_tag or ""

        # Delete all IntelligenceItems tagged with this category for this project
        items = _items_for_category(project_id, category_tag)
        item_count = len(items)
        for item in items:
            if item.file_path:
                fpath = os.path.join(app.config["LIBRARY_FOLDER"], item.file_path)
                if os.path.exists(fpath):
                    os.remove(fpath)
            for tag in list(item.tags):
                if tag.usage_count > 0:
                    tag.usage_count -= 1
            db.session.delete(item)

        # Delete the ComparisonSummary for this category
        ComparisonSummary.query.filter_by(
            project_id=project_id, category_tag=category_tag
        ).delete(synchronize_session=False)

        # Delete all saved batches for this (project, category_tag) and their staging dirs
        all_batches = QuoteBatch.query.filter_by(
            project_id=project_id, category_tag=category_tag
        ).all()
        for b in all_batches:
            staging_dir = os.path.join(
                app.config["LIBRARY_FOLDER"], "quotes_staging", b.batch_id
            )
            if os.path.isdir(staging_dir):
                shutil.rmtree(staging_dir, ignore_errors=True)
            db.session.delete(b)

        db.session.commit()
        label = category_tag or "(unlabeled)"
        flash(f"Deleted batch '{label}' and its {item_count} vendor quote(s).", "success")
        return redirect(url_for("quote_compare_select", project_id=project_id))

    @app.route("/quote-item/<int:item_id>/delete", methods=["POST"])
    @login_required
    @admin_required
    def delete_quote_item(item_id):
        item = db.session.get(IntelligenceItem, item_id) or abort(404)
        if not current_user.is_superadmin and item.project_id:
            proj = db.session.get(Project, item.project_id)
            if not proj or proj.company_id != current_user.company_id:
                abort(403)

        project_id = item.project_id
        category_tag = request.form.get("category_tag", "")
        vendor_name = item.vendor_name or item.title or "Unknown vendor"

        # Delete file artifact
        if item.file_path:
            fpath = os.path.join(app.config["LIBRARY_FOLDER"], item.file_path)
            if os.path.exists(fpath):
                os.remove(fpath)

        # Decrement tag usage counts and delete the item
        for tag in list(item.tags):
            if tag.usage_count > 0:
                tag.usage_count -= 1
        db.session.delete(item)
        db.session.flush()

        # Check remaining vendors in this category
        remaining = _items_for_category(project_id, category_tag) if category_tag else []

        if not remaining:
            # Last vendor — cascade: delete ComparisonSummary and all QuoteBatch rows
            import shutil
            ComparisonSummary.query.filter_by(
                project_id=project_id, category_tag=category_tag
            ).delete(synchronize_session=False)
            all_batches = QuoteBatch.query.filter_by(
                project_id=project_id, category_tag=category_tag
            ).all()
            for b in all_batches:
                staging_dir = os.path.join(
                    app.config["LIBRARY_FOLDER"], "quotes_staging", b.batch_id
                )
                if os.path.isdir(staging_dir):
                    shutil.rmtree(staging_dir, ignore_errors=True)
                db.session.delete(b)
            db.session.commit()
            label = category_tag or "(unlabeled)"
            flash(
                f"Deleted vendor quote '{vendor_name}' and removed empty batch '{label}'.",
                "success",
            )
        else:
            # Vendors remain — clear Skippy cache so it regenerates with the updated pool
            summary = ComparisonSummary.query.filter_by(
                project_id=project_id, category_tag=category_tag
            ).first()
            if summary:
                summary.skippy_recommendation = None
            db.session.commit()
            label = category_tag or "(unlabeled)"
            flash(f"Deleted vendor quote '{vendor_name}' from '{label}'.", "success")

        return redirect(url_for("quote_compare_select", project_id=project_id))

    @app.route("/projects/<int:project_id>/quotes/compare/<path:category_tag>")
    @login_required
    @admin_required
    def quote_compare(project_id, category_tag):
        project = _auth_project(project_id)
        items = _items_for_category(project_id, category_tag)

        if not items:
            flash(f"No quotes found for category '{category_tag}'.", "warning")
            return redirect(url_for("quote_compare_select", project_id=project_id))

        only_one = len(items) == 1

        # Fetch or generate Skippy recommendation (synchronous on first load)
        summary_obj = ComparisonSummary.query.filter_by(
            project_id=project_id, category_tag=category_tag
        ).first()

        skippy_data = None
        if summary_obj and summary_obj.skippy_recommendation:
            try:
                skippy_data = json.loads(summary_obj.skippy_recommendation)
            except Exception:
                skippy_data = None

        if skippy_data is None and not only_one:
            api_key = app.config.get("ANTHROPIC_API_KEY") or os.getenv("ANTHROPIC_API_KEY", "")
            if api_key:
                skippy_data = _generate_skippy_recommendation(api_key, project, items, category_tag)
                if skippy_data is not None:
                    if not summary_obj:
                        active_takeoff = _get_active_takeoff(project_id)
                        summary_obj = ComparisonSummary(
                            project_id=project_id,
                            category_tag=category_tag,
                            summary_text="",
                            takeoff_id=active_takeoff.id,
                        )
                        db.session.add(summary_obj)
                    summary_obj.skippy_recommendation = json.dumps(skippy_data)
                    db.session.commit()

        all_labels, per_item_pricing = _build_pricing_matrix(items)
        shortlisted_count = sum(1 for it in items if it.shortlisted)

        item_flags = {}
        for it in items:
            try:
                item_flags[it.id] = json.loads(it.flags_json) if it.flags_json else []
            except Exception:
                item_flags[it.id] = []

        now_date = datetime.now(timezone.utc).strftime("%Y-%m-%d")

        return render_template(
            "quote_compare.html",
            project=project,
            category_tag=category_tag,
            items=items,
            only_one=only_one,
            skippy_data=skippy_data,
            all_labels=all_labels,
            per_item_pricing=per_item_pricing,
            shortlisted_count=shortlisted_count,
            item_flags=item_flags,
            now_date=now_date,
        )

    @app.route("/projects/<int:project_id>/quotes/compare/<path:category_tag>/skippy",
               methods=["POST"])
    @login_required
    @admin_required
    def quote_skippy(project_id, category_tag):
        _auth_project(project_id)
        force = (request.get_json(force=True, silent=True) or {}).get("force", False)

        summary_obj = ComparisonSummary.query.filter_by(
            project_id=project_id, category_tag=category_tag
        ).first()

        if summary_obj and summary_obj.skippy_recommendation and not force:
            try:
                return jsonify({"ok": True, "data": json.loads(summary_obj.skippy_recommendation)})
            except Exception:
                pass

        api_key = current_app.config.get("ANTHROPIC_API_KEY")
        if not api_key:
            return jsonify({"ok": False, "error": "Claude API key not configured."}), 503

        items = _items_for_category(project_id, category_tag)
        if not items:
            return jsonify({"ok": False, "error": "No quotes found for this category."}), 404

        project = db.session.get(Project, project_id)
        data = _generate_skippy_recommendation(api_key, project, items, category_tag)
        if data is None:
            return jsonify({"ok": False, "error": "Failed to generate recommendation."}), 500

        if not summary_obj:
            active_takeoff = _get_active_takeoff(project_id)
            summary_obj = ComparisonSummary(
                project_id=project_id,
                category_tag=category_tag,
                summary_text="",
                takeoff_id=active_takeoff.id,
            )
            db.session.add(summary_obj)

        summary_obj.skippy_recommendation = json.dumps(data)
        db.session.commit()

        return jsonify({"ok": True, "data": data})

    @app.route("/projects/<int:project_id>/quotes/compare/<path:category_tag>/shortlist",
               methods=["POST"])
    @login_required
    @admin_required
    def quote_shortlist(project_id, category_tag):
        _auth_project(project_id)
        data = request.get_json(force=True, silent=True) or {}
        updates = data.get("items", [])
        for entry in updates:
            item_id = entry.get("id")
            if not item_id:
                continue
            item = db.session.get(IntelligenceItem, item_id)
            if not item or item.project_id != project_id:
                continue
            item.shortlisted = bool(entry.get("shortlisted", False))
            item.shortlist_notes = entry.get("notes") or None
            if item.shortlisted:
                item.shortlisted_at = datetime.now(timezone.utc)
                item.shortlisted_by = current_user.id
            else:
                item.shortlisted_at = None
                item.shortlisted_by = None
        db.session.commit()
        return jsonify({"ok": True})

    @app.route("/projects/<int:project_id>/quotes/compare/<path:category_tag>/export/excel")
    @login_required
    @admin_required
    def quote_export_excel(project_id, category_tag):
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
        from openpyxl.utils import get_column_letter

        project = _auth_project(project_id)
        items = _items_for_category(project_id, category_tag)
        if not items:
            flash("No quotes to export.", "warning")
            return redirect(url_for("quote_compare_select", project_id=project_id))

        summary_obj = ComparisonSummary.query.filter_by(
            project_id=project_id, category_tag=category_tag
        ).first()
        all_labels, per_item_pricing = _build_pricing_matrix(items)

        ts = datetime.now(timezone.utc).strftime("%B %d, %Y %H:%M UTC")
        snapshot_note = f"Snapshot as of {ts} — For reference only. Verify against current quotes before using."

        HEADER_FILL = PatternFill("solid", fgColor="1F4E79")
        HEADER_FONT = Font(bold=True, color="FFFFFF", size=11)
        SUBHEADER_FILL = PatternFill("solid", fgColor="D6E4F0")
        SUBHEADER_FONT = Font(bold=True, size=10)
        WARN_FONT = Font(italic=True, color="CC0000", size=9)

        wb = openpyxl.Workbook()

        # ── Sheet 1: Summary ──────────────────────────────────────────────────
        ws1 = wb.active
        ws1.title = "Summary"
        ws1.column_dimensions["A"].width = 100
        ws1["A1"] = f"Quote Comparison — {category_tag}"
        ws1["A1"].font = Font(bold=True, size=14, color="1F4E79")
        ws1["A2"] = f"Project: {project.name}"
        ws1["A3"] = snapshot_note
        ws1["A3"].font = WARN_FONT
        ws1["A4"] = ""
        if summary_obj:
            ws1["A5"] = "AI Narrative Summary"
            ws1["A5"].font = SUBHEADER_FONT
            ws1["A6"] = summary_obj.summary_text
            ws1["A6"].alignment = Alignment(wrap_text=True)
            gen_ts = summary_obj.generated_at.strftime("%b %d, %Y %H:%M UTC") \
                if summary_obj.generated_at else ""
            ws1["A7"] = f"Summary generated: {gen_ts}"
            ws1["A7"].font = Font(italic=True, size=9)
        ws1.row_dimensions[6].height = 120

        # ── Sheet 2: Full Comparison ──────────────────────────────────────────
        ws2 = wb.create_sheet("Full Comparison")
        _write_category_sheet(ws2, items, all_labels, per_item_pricing,
                              category_tag, summary_obj, project.name, ts)

        # ── Sheet 3: Audit Info ───────────────────────────────────────────────
        ws3 = wb.create_sheet("Audit Info")
        ws3.column_dimensions["A"].width = 30
        ws3.column_dimensions["B"].width = 60
        audit_rows = [
            ("Generated by", current_user.username),
            ("Generated at", ts),
            ("Project", project.name),
            ("Category", category_tag),
            ("Vendor count", len(items)),
            ("Shortlisted count", sum(1 for it in items if it.shortlisted)),
            ("Note", snapshot_note),
        ]
        for r, (k, v) in enumerate(audit_rows, start=1):
            ws3.cell(row=r, column=1, value=k).font = Font(bold=True)
            ws3.cell(row=r, column=2, value=str(v))

        # Log export
        log = QuoteComparisonExport(
            project_id=project_id,
            user_id=current_user.id,
            export_type="excel",
            category_tag=category_tag,
            vendor_count=len(items),
            shortlisted_count=sum(1 for it in items if it.shortlisted),
        )
        db.session.add(log)
        db.session.commit()

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        safe_cat = "".join(c if c.isalnum() or c in "-_ " else "_" for c in category_tag)[:40]
        fname = f"QuoteComparison_{safe_cat}_{datetime.now().strftime('%Y%m%d')}.xlsx"
        return send_file(
            buf,
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            as_attachment=True,
            download_name=fname,
        )

    @app.route("/projects/<int:project_id>/quotes/compare/<path:category_tag>/export/pdf")
    @login_required
    @admin_required
    def quote_export_pdf(project_id, category_tag):
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak,
        )

        project = _auth_project(project_id)
        items = _items_for_category(project_id, category_tag)
        shortlisted = [it for it in items if it.shortlisted]

        if not shortlisted:
            flash("Shortlist at least one vendor before downloading the PDF.", "warning")
            return redirect(url_for("quote_compare", project_id=project_id, category_tag=category_tag))

        summary_obj = ComparisonSummary.query.filter_by(
            project_id=project_id, category_tag=category_tag
        ).first()
        all_labels, per_item_pricing = _build_pricing_matrix(shortlisted)

        ts = datetime.now(timezone.utc).strftime("%B %d, %Y %H:%M UTC")
        snapshot_note = f"Snapshot as of {ts} — For reference only."

        buf = io.BytesIO()
        styles = getSampleStyleSheet()
        BRAND_BLUE = colors.HexColor("#1F4E79")

        header_style = ParagraphStyle(
            "Header", parent=styles["Normal"],
            fontSize=9, textColor=colors.red, fontName="Helvetica-Bold",
            spaceAfter=4,
        )
        title_style = ParagraphStyle(
            "Title", parent=styles["Normal"],
            fontSize=18, textColor=BRAND_BLUE, fontName="Helvetica-Bold",
            spaceAfter=6,
        )
        h2_style = ParagraphStyle(
            "H2", parent=styles["Normal"],
            fontSize=13, textColor=BRAND_BLUE, fontName="Helvetica-Bold",
            spaceBefore=12, spaceAfter=4,
        )
        body_style = styles["BodyText"]
        small_style = ParagraphStyle(
            "Small", parent=styles["Normal"],
            fontSize=8, textColor=colors.grey,
        )

        def _footer(canvas, doc):
            canvas.saveState()
            canvas.setFont("Helvetica", 7)
            canvas.setFillColor(colors.grey)
            canvas.drawString(0.75 * inch, 0.4 * inch,
                              f"Internal Use Only — {ts}   |   Page {doc.page}")
            canvas.restoreState()

        doc = SimpleDocTemplate(
            buf, pagesize=letter,
            leftMargin=0.75 * inch, rightMargin=0.75 * inch,
            topMargin=0.75 * inch, bottomMargin=0.6 * inch,
        )

        story = []

        # Cover page
        story.append(Paragraph("INTERNAL USE ONLY — DRAFT COMPARISON", header_style))
        story.append(Spacer(1, 0.3 * inch))
        story.append(Paragraph(f"Quote Comparison", title_style))
        story.append(Paragraph(f"Category: {category_tag}", styles["Heading2"]))
        story.append(Paragraph(f"Project: {project.name}", styles["Heading3"]))
        story.append(Spacer(1, 0.15 * inch))
        story.append(Paragraph(snapshot_note, small_style))
        story.append(Paragraph(
            f"Shortlisted vendors: {len(shortlisted)} of {len(items)}",
            body_style,
        ))
        story.append(PageBreak())

        # AI Summary
        story.append(Paragraph("INTERNAL USE ONLY — DRAFT COMPARISON", header_style))
        story.append(Paragraph("Estimator Summary", h2_style))
        if summary_obj:
            story.append(Paragraph(summary_obj.summary_text, body_style))
            gen_ts = summary_obj.generated_at.strftime("%b %d, %Y %H:%M UTC") \
                if summary_obj.generated_at else ""
            story.append(Paragraph(f"Summary generated: {gen_ts}", small_style))
        else:
            story.append(Paragraph("(No summary available)", body_style))
        story.append(Spacer(1, 0.2 * inch))

        # Shortlisted comparison table
        story.append(Paragraph("Shortlisted Vendor Comparison", h2_style))

        # Build table data: header row + one row per attribute group
        col_width = (letter[0] - 1.5 * inch) / (len(shortlisted) + 1)
        label_col_width = 1.4 * inch
        vendor_col_width = (letter[0] - 1.5 * inch - label_col_width) / max(len(shortlisted), 1)

        def _cell(text, bold=False, wrap=True):
            style_name = "Helvetica-Bold" if bold else "Helvetica"
            return Paragraph(
                str(text or ""),
                ParagraphStyle("c", fontName=style_name, fontSize=8,
                               leading=10, wordWrap="CJK" if wrap else None),
            )

        hdr_row = [_cell("", bold=True)] + [
            _cell(it.vendor_name or it.title or "—", bold=True) for it in shortlisted
        ]
        table_data = [hdr_row]

        # Vendor info rows
        info_fields = [
            ("Quote Date", lambda it: str(it.quote_date) if it.quote_date else "—"),
            ("Expiration", lambda it: str(it.expiration_date) if it.expiration_date else "—"),
            ("Contact", lambda it: (it.vendor_contact or "—")[:60]),
            ("Source File", lambda it: (it.original_filename or "—")[:40]),
        ]
        for label, getter in info_fields:
            table_data.append(
                [_cell(label, bold=True)] + [_cell(getter(it)) for it in shortlisted]
            )

        # Pricing rows
        if all_labels:
            table_data.append([_cell("— Pricing —", bold=True)] + [_cell("") for _ in shortlisted])
            for lbl in all_labels:
                row = [_cell(lbl, bold=True)]
                for it in shortlisted:
                    pr = per_item_pricing.get(it.id, {}).get(lbl)
                    if pr:
                        val = f"{pr.get('amount','')} {pr.get('unit','')}".strip()
                        if pr.get("notes"):
                            val += f"\n({pr['notes']})"
                    else:
                        val = "—"
                    row.append(_cell(val))
                table_data.append(row)

        # Conditions
        table_data.append([_cell("— Conditions —", bold=True)] + [_cell("") for _ in shortlisted])
        table_data.append(
            [_cell("Conditions /\nExclusions", bold=True)] +
            [_cell((it.conditions_text or "—")[:300]) for it in shortlisted]
        )

        # Flags
        table_data.append([_cell("— Flags —", bold=True)] + [_cell("") for _ in shortlisted])
        for it in shortlisted:
            flags = []
            if it.flags_json:
                try:
                    flags = json.loads(it.flags_json)
                except Exception:
                    flags = []
            pass
        flag_row = [_cell("Flags", bold=True)]
        for it in shortlisted:
            flags = []
            if it.flags_json:
                try:
                    flags = json.loads(it.flags_json)
                except Exception:
                    flags = []
            flag_row.append(_cell(", ".join(flags) if flags else "—"))
        table_data.append(flag_row)

        col_widths = [label_col_width] + [vendor_col_width] * len(shortlisted)
        tbl = Table(table_data, colWidths=col_widths, repeatRows=1)
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E79")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
            ("BACKGROUND", (0, 1), (0, -1), colors.HexColor("#EEF4FA")),
            ("ROWBACKGROUNDS", (1, 1), (-1, -1), [colors.white, colors.HexColor("#F8FBFF")]),
        ]))
        story.append(tbl)

        doc.build(story, onFirstPage=_footer, onLaterPages=_footer)
        buf.seek(0)

        # Log export
        log = QuoteComparisonExport(
            project_id=project_id,
            user_id=current_user.id,
            export_type="pdf",
            category_tag=category_tag,
            vendor_count=len(items),
            shortlisted_count=len(shortlisted),
        )
        db.session.add(log)
        db.session.commit()

        safe_cat = "".join(c if c.isalnum() or c in "-_ " else "_" for c in category_tag)[:40]
        fname = f"QuoteComparison_Shortlist_{safe_cat}_{datetime.now().strftime('%Y%m%d')}.pdf"
        return send_file(buf, mimetype="application/pdf", as_attachment=True, download_name=fname)

    @app.route("/projects/<int:project_id>/quotes/compare/export-all/excel")
    @login_required
    @admin_required
    def quote_export_all_excel(project_id):
        import openpyxl

        project = _auth_project(project_id)
        saved_batches = QuoteBatch.query.filter_by(
            project_id=project_id, status="saved"
        ).all()
        seen = set()
        ordered_tags = []
        for b in saved_batches:
            tag = (b.category_tag or "").strip()
            if tag and tag not in seen:
                seen.add(tag)
                ordered_tags.append(tag)

        if not ordered_tags:
            flash("No saved quote categories found for this project.", "warning")
            return redirect(url_for("quote_compare_select", project_id=project_id))

        ts = datetime.now(timezone.utc).strftime("%B %d, %Y %H:%M UTC")
        wb = openpyxl.Workbook()
        wb.remove(wb.active)  # remove default blank sheet

        total_vendors = 0
        total_shortlisted = 0

        for tag in ordered_tags:
            items = _items_for_category(project_id, tag)
            if not items:
                continue
            all_labels, per_item_pricing = _build_pricing_matrix(items)
            summary_obj = ComparisonSummary.query.filter_by(
                project_id=project_id, category_tag=tag
            ).first()
            # Sheet name: truncate to 31 chars (Excel limit)
            sheet_name = tag[:31] if len(tag) <= 31 else tag[:28] + "..."
            ws = wb.create_sheet(title=sheet_name)
            _write_category_sheet(ws, items, all_labels, per_item_pricing,
                                  tag, summary_obj, project.name, ts)
            total_vendors += len(items)
            total_shortlisted += sum(1 for it in items if it.shortlisted)

        # Audit sheet at the end
        from openpyxl.styles import Font
        ws_audit = wb.create_sheet("Audit Info")
        ws_audit.column_dimensions["A"].width = 30
        ws_audit.column_dimensions["B"].width = 60
        audit_rows = [
            ("Generated by", current_user.username),
            ("Generated at", ts),
            ("Project", project.name),
            ("Categories exported", ", ".join(ordered_tags)),
            ("Total vendor count", total_vendors),
            ("Total shortlisted", total_shortlisted),
        ]
        for r, (k, v) in enumerate(audit_rows, start=1):
            ws_audit.cell(row=r, column=1, value=k).font = Font(bold=True)
            ws_audit.cell(row=r, column=2, value=str(v))

        log = QuoteComparisonExport(
            project_id=project_id,
            user_id=current_user.id,
            export_type="excel",
            category_tag="ALL",
            vendor_count=total_vendors,
            shortlisted_count=total_shortlisted,
        )
        db.session.add(log)
        db.session.commit()

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        safe_proj = "".join(c if c.isalnum() or c in "-_ " else "_" for c in project.name)[:30]
        fname = f"QuoteComparison_AllCategories_{safe_proj}_{datetime.now().strftime('%Y%m%d')}.xlsx"
        return send_file(
            buf,
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            as_attachment=True,
            download_name=fname,
        )

    @app.route("/projects/<int:project_id>/quotes/compare/export-all/pdf")
    @login_required
    @admin_required
    def quote_export_all_pdf(project_id):
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak,
        )

        project = _auth_project(project_id)
        saved_batches = QuoteBatch.query.filter_by(
            project_id=project_id, status="saved"
        ).all()
        seen = set()
        ordered_tags = []
        for b in saved_batches:
            tag = (b.category_tag or "").strip()
            if tag and tag not in seen:
                seen.add(tag)
                ordered_tags.append(tag)

        if not ordered_tags:
            flash("No saved quote categories found for this project.", "warning")
            return redirect(url_for("quote_compare_select", project_id=project_id))

        ts = datetime.now(timezone.utc).strftime("%B %d, %Y %H:%M UTC")
        BRAND_BLUE = colors.HexColor("#1F4E79")
        styles = getSampleStyleSheet()

        header_style = ParagraphStyle(
            "Header", parent=styles["Normal"],
            fontSize=9, textColor=colors.red, fontName="Helvetica-Bold", spaceAfter=4,
        )
        title_style = ParagraphStyle(
            "Title", parent=styles["Normal"],
            fontSize=18, textColor=BRAND_BLUE, fontName="Helvetica-Bold", spaceAfter=6,
        )
        h2_style = ParagraphStyle(
            "H2", parent=styles["Normal"],
            fontSize=13, textColor=BRAND_BLUE, fontName="Helvetica-Bold",
            spaceBefore=12, spaceAfter=4,
        )
        h3_style = ParagraphStyle(
            "H3", parent=styles["Normal"],
            fontSize=11, textColor=BRAND_BLUE, fontName="Helvetica-Bold",
            spaceBefore=8, spaceAfter=3,
        )
        body_style = styles["BodyText"]
        small_style = ParagraphStyle(
            "Small", parent=styles["Normal"], fontSize=8, textColor=colors.grey,
        )

        def _footer(canvas, doc):
            canvas.saveState()
            canvas.setFont("Helvetica", 7)
            canvas.setFillColor(colors.grey)
            canvas.drawString(0.75 * inch, 0.4 * inch,
                              f"Internal Use Only — {ts}   |   Page {doc.page}")
            canvas.restoreState()

        def _cell(text, bold=False):
            fn = "Helvetica-Bold" if bold else "Helvetica"
            return Paragraph(
                str(text or ""),
                ParagraphStyle("c", fontName=fn, fontSize=8, leading=10),
            )

        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf, pagesize=letter,
            leftMargin=0.75 * inch, rightMargin=0.75 * inch,
            topMargin=0.75 * inch, bottomMargin=0.6 * inch,
        )

        story = []

        # Cover page
        story.append(Paragraph("INTERNAL USE ONLY — DRAFT COMPARISON", header_style))
        story.append(Spacer(1, 0.3 * inch))
        story.append(Paragraph("Quote Comparison — All Categories", title_style))
        story.append(Paragraph(f"Project: {project.name}", styles["Heading3"]))
        story.append(Spacer(1, 0.1 * inch))
        story.append(Paragraph(
            f"Snapshot as of {ts} — For reference only.", small_style
        ))
        story.append(Paragraph(
            f"Categories: {', '.join(ordered_tags)}", body_style
        ))
        story.append(PageBreak())

        first_cat = True
        total_vendors = 0
        total_shortlisted = 0

        for tag in ordered_tags:
            items = _items_for_category(project_id, tag)
            if not items:
                continue
            shortlisted = [it for it in items if it.shortlisted]
            all_labels, per_item_pricing = _build_pricing_matrix(shortlisted or items)
            summary_obj = ComparisonSummary.query.filter_by(
                project_id=project_id, category_tag=tag
            ).first()
            total_vendors += len(items)
            total_shortlisted += len(shortlisted)

            if not first_cat:
                story.append(PageBreak())
            first_cat = False

            story.append(Paragraph("INTERNAL USE ONLY — DRAFT COMPARISON", header_style))
            story.append(Paragraph(tag, h2_style))
            story.append(Paragraph(
                f"{len(items)} vendor(s) — "
                f"{len(shortlisted)} shortlisted",
                small_style,
            ))

            if summary_obj:
                story.append(Paragraph("Estimator Summary", h3_style))
                story.append(Paragraph(summary_obj.summary_text, body_style))
                gen_ts = summary_obj.generated_at.strftime("%b %d, %Y %H:%M UTC") \
                    if summary_obj.generated_at else ""
                story.append(Paragraph(f"Summary generated: {gen_ts}", small_style))
                story.append(Spacer(1, 0.15 * inch))

            display_items = shortlisted if shortlisted else items
            display_note = "Shortlisted vendors" if shortlisted else "All vendors (none shortlisted)"
            story.append(Paragraph(display_note, h3_style))

            label_col_width = 1.3 * inch
            vendor_col_width = (
                (letter[0] - 1.5 * inch - label_col_width) / max(len(display_items), 1)
            )
            col_widths = [label_col_width] + [vendor_col_width] * len(display_items)

            hdr_row = [_cell("", bold=True)] + [
                _cell(it.vendor_name or it.title or "—", bold=True)
                for it in display_items
            ]
            table_data = [hdr_row]

            info_fields = [
                ("Quote Date", lambda it: str(it.quote_date) if it.quote_date else "—"),
                ("Expiration",  lambda it: str(it.expiration_date) if it.expiration_date else "—"),
                ("Contact",     lambda it: (it.vendor_contact or "—")[:60]),
                ("Source File", lambda it: (it.original_filename or "—")[:40]),
            ]
            for label, getter in info_fields:
                table_data.append(
                    [_cell(label, bold=True)] + [_cell(getter(it)) for it in display_items]
                )

            if all_labels:
                table_data.append(
                    [_cell("— Pricing —", bold=True)] + [_cell("") for _ in display_items]
                )
                for lbl in all_labels:
                    row = [_cell(lbl, bold=True)]
                    for it in display_items:
                        pr = per_item_pricing.get(it.id, {}).get(lbl)
                        if pr:
                            val = f"{pr.get('amount','')} {pr.get('unit','')}".strip()
                            if pr.get("notes"):
                                val += f" ({pr['notes']})"
                        else:
                            val = "—"
                        row.append(_cell(val))
                    table_data.append(row)

            table_data.append(
                [_cell("— Conditions —", bold=True)] + [_cell("") for _ in display_items]
            )
            table_data.append(
                [_cell("Conditions", bold=True)] +
                [_cell((it.conditions_text or "—")[:300]) for it in display_items]
            )

            flag_row = [_cell("Flags", bold=True)]
            for it in display_items:
                flags = []
                if it.flags_json:
                    try:
                        flags = json.loads(it.flags_json)
                    except Exception:
                        flags = []
                flag_row.append(_cell(", ".join(flags) if flags else "—"))
            table_data.append(flag_row)

            tbl = Table(table_data, colWidths=col_widths, repeatRows=1)
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E79")),
                ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
                ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE",   (0, 0), (-1, -1), 8),
                ("VALIGN",     (0, 0), (-1, -1), "TOP"),
                ("GRID",       (0, 0), (-1, -1), 0.25, colors.lightgrey),
                ("BACKGROUND", (0, 1), (0, -1), colors.HexColor("#EEF4FA")),
                ("ROWBACKGROUNDS", (1, 1), (-1, -1),
                 [colors.white, colors.HexColor("#F8FBFF")]),
            ]))
            story.append(tbl)

        doc.build(story, onFirstPage=_footer, onLaterPages=_footer)
        buf.seek(0)

        log = QuoteComparisonExport(
            project_id=project_id,
            user_id=current_user.id,
            export_type="pdf",
            category_tag="ALL",
            vendor_count=total_vendors,
            shortlisted_count=total_shortlisted,
        )
        db.session.add(log)
        db.session.commit()

        safe_proj = "".join(c if c.isalnum() or c in "-_ " else "_" for c in project.name)[:30]
        fname = f"QuoteComparison_AllCategories_{safe_proj}_{datetime.now().strftime('%Y%m%d')}.pdf"
        return send_file(buf, mimetype="application/pdf", as_attachment=True, download_name=fname)

    @app.route("/projects/<int:project_id>/upload", methods=["GET", "POST"])
    @login_required
    def upload_drawing(project_id):
        """Legacy upload path — redirected to Library add (Phase 3C Part A)."""
        project = db.session.get(Project, project_id) or abort(404)
        if not current_user.is_superadmin and current_user.company_id != project.company_id:
            abort(403)
        return redirect(url_for("library_add_for_project", project_id=project_id))

    @app.route("/projects/<int:project_id>/upload-file", methods=["POST"])
    @login_required
    def upload_single_file(project_id):
        """Legacy AJAX endpoint — no longer active (Phase 3C Part A)."""
        return ("This upload endpoint is no longer active. Use the Intelligence Library.", 410)

    @app.route("/projects/<int:project_id>/classify-file", methods=["POST"])
    @login_required
    def classify_file(project_id):
        """Legacy classifier endpoint — no longer active (Phase 3C Part A)."""
        return ("This endpoint is no longer active.", 410)


    @app.route("/drawings/<int:drawing_id>")
    @login_required
    def drawing_detail(drawing_id):
        drawing = db.session.get(Drawing, drawing_id) or abort(404)
        if not current_user.is_superadmin and current_user.company_id != drawing.project.company_id:
            abort(403)
        pages = DrawingPage.query.filter_by(drawing_id=drawing.id).order_by(DrawingPage.page_number).all()
        return render_template("drawing_detail.html", drawing=drawing, pages=pages)

    @app.route("/drawings/<int:drawing_id>/status")
    @login_required
    def drawing_status(drawing_id):
        drawing = db.session.get(Drawing, drawing_id) or abort(404)
        return jsonify({
            "status": drawing.status,
            "total_pages": drawing.total_pages,
            "pages_processed": drawing.pages_processed,
        })

    @app.route("/drawings/<int:drawing_id>/reprocess", methods=["POST"])
    @login_required
    def reprocess_drawing(drawing_id):
        drawing = db.session.get(Drawing, drawing_id) or abort(404)
        if not current_user.is_superadmin and current_user.company_id != drawing.project.company_id:
            abort(403)
        if drawing.status == "processing":
            flash("Document is already being processed.", "warning")
            return redirect(url_for("drawing_detail", drawing_id=drawing.id))

        drawing.status = "pending"
        drawing.pages_processed = 0
        db.session.commit()
        flash("Reconverting PDF pages to images.", "success")
        return redirect(url_for("drawing_detail", drawing_id=drawing.id))

    @app.route("/drawings/<int:drawing_id>/delete", methods=["POST"])
    @login_required
    @admin_required
    def delete_drawing(drawing_id):
        drawing = db.session.get(Drawing, drawing_id) or abort(404)
        project_id = drawing.project_id
        db.session.delete(drawing)
        db.session.commit()
        flash("Document deleted.", "success")
        return redirect(url_for("drawings", project_id=project_id))

    @app.route("/processed/<path:filename>")
    @login_required
    def serve_processed(filename):
        return send_from_directory(app.config["PROCESSED_FOLDER"], filename)

    # ── Search History ──────────────────────────────────────────

    def _project_history(project_id):
        project = db.session.get(Project, project_id) or abort(404)
        if not current_user.is_superadmin and current_user.company_id != project.company_id:
            abort(403)
        entries = (
            db.session.query(SearchHistory)
            .filter_by(project_id=project.id)
            .order_by(SearchHistory.created_at.desc())
            .all()
        )
        return project, entries

    @app.route("/projects/<int:project_id>/history")
    @login_required
    def search_history(project_id):
        project, entries = _project_history(project_id)
        return render_template("search_history.html", project=project, entries=entries)

    @app.route("/projects/<int:project_id>/report/generate", methods=["POST"])
    @login_required
    def generate_project_report(project_id):
        project = db.session.get(Project, project_id) or abort(404)
        if not current_user.is_superadmin and current_user.company_id != project.company_id:
            abort(403)

        if not app.config["ANTHROPIC_API_KEY"]:
            flash("Claude API key not configured.", "danger")
            return redirect(url_for("drawings", project_id=project.id))

        template_id = request.form.get("template_id", "").strip()
        if not template_id or (template_id != "custom" and template_id not in REPORT_TEMPLATES):
            flash("Please choose a valid report template.", "danger")
            return redirect(url_for("drawings", project_id=project.id))

        if template_id == "estimating_intelligence":
            custom_prompt = json.dumps({
                "notes_scope": request.form.get("notes_scope", "project").strip(),
                "focus_areas": request.form.get("focus_areas", "").strip(),
            })
        else:
            custom_prompt = request.form.get("custom_prompt", "").strip()

        try:
            enqueue_report(project.id, current_user.id, template_id, custom_prompt)
        except ValueError as e:
            flash(str(e), "danger")
            return redirect(url_for("drawings", project_id=project.id))

        flash("Report is being generated in the background. It will appear in the Reports tab when ready.", "info")
        return redirect(url_for("drawings", project_id=project.id, tab="reports"))

    @app.route("/reports/<int:report_id>/download")
    @login_required
    def download_report(report_id):
        report = db.session.get(Report, report_id) or abort(404)
        project = db.session.get(Project, report.project_id) or abort(404)
        if not current_user.is_superadmin and current_user.company_id != project.company_id:
            abort(403)
        if report.status != "ready" or not report.file_path:
            flash("Report is not ready yet.", "warning")
            return redirect(url_for("drawings", project_id=project.id, tab="reports"))
        file_path = os.path.join(app.config["REPORTS_FOLDER"], report.file_path)
        with open(file_path, "rb") as f:
            data = f.read()
        return send_file(
            io.BytesIO(data),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=report.file_path,
        )

    @app.route("/reports/<int:report_id>/download/pdf")
    @login_required
    def download_report_pdf(report_id):
        import subprocess
        import shutil
        import tempfile

        report = db.session.get(Report, report_id) or abort(404)
        project = db.session.get(Project, report.project_id) or abort(404)
        if not current_user.is_superadmin and current_user.company_id != project.company_id:
            abort(403)
        if report.status != "ready" or not report.file_path:
            flash("Report is not ready yet.", "warning")
            return redirect(url_for("drawings", project_id=project.id, tab="reports"))

        docx_path = os.path.join(app.config["REPORTS_FOLDER"], report.file_path)
        pdf_name = os.path.splitext(report.file_path)[0] + ".pdf"

        # Try LibreOffice conversion first
        pdf_bytes = None
        tmp_dir = tempfile.mkdtemp()
        try:
            result = subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf",
                 "--outdir", tmp_dir, docx_path],
                capture_output=True,
                timeout=60,
            )
            if result.returncode == 0:
                base = os.path.splitext(os.path.basename(docx_path))[0]
                pdf_path = os.path.join(tmp_dir, base + ".pdf")
                if os.path.exists(pdf_path):
                    with open(pdf_path, "rb") as f:
                        pdf_bytes = f.read()
        except Exception:
            pass
        finally:
            shutil.rmtree(tmp_dir, ignore_errors=True)

        # Fall back to reportlab if LibreOffice unavailable or failed
        if pdf_bytes is None:
            from docx import Document as DocxDocument
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from xml.sax.saxutils import escape as xml_escape

            buf = io.BytesIO()
            rdoc = SimpleDocTemplate(
                buf, pagesize=letter,
                leftMargin=0.75 * inch, rightMargin=0.75 * inch,
                topMargin=0.75 * inch, bottomMargin=0.75 * inch,
            )
            styles = getSampleStyleSheet()
            story = []
            try:
                docx = DocxDocument(docx_path)
                for para in docx.paragraphs:
                    text = para.text.strip()
                    if not text:
                        story.append(Spacer(1, 6))
                        continue
                    style_name = para.style.name if para.style else "Normal"
                    if "Heading 1" in style_name:
                        story.append(Paragraph(xml_escape(text), styles["Heading1"]))
                    elif "Heading 2" in style_name:
                        story.append(Paragraph(xml_escape(text), styles["Heading2"]))
                    elif "Heading 3" in style_name:
                        story.append(Paragraph(xml_escape(text), styles["Heading3"]))
                    else:
                        story.append(Paragraph(xml_escape(text), styles["BodyText"]))
                for table in docx.tables:
                    for row in table.rows:
                        cells = " | ".join(xml_escape(c.text.strip()) for c in row.cells)
                        story.append(Paragraph(cells, styles["Normal"]))
                    story.append(Spacer(1, 6))
            except Exception:
                story.append(Paragraph("(could not read report content)", styles["Normal"]))
            if not story:
                story.append(Paragraph("(empty report)", styles["Normal"]))
            rdoc.build(story)
            buf.seek(0)
            pdf_bytes = buf.read()

        return send_file(
            io.BytesIO(pdf_bytes),
            mimetype="application/pdf",
            as_attachment=True,
            download_name=pdf_name,
        )

    @app.route("/reports/<int:report_id>/download/excel")
    @login_required
    def download_report_excel(report_id):
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
        from openpyxl.utils import get_column_letter
        from docx import Document as _DocxDoc
        from docx.text.paragraph import Paragraph as _DocxPara
        from docx.table import Table as _DocxTable

        report = db.session.get(Report, report_id) or abort(404)
        project = db.session.get(Project, report.project_id) or abort(404)
        if not current_user.is_superadmin and current_user.company_id != project.company_id:
            abort(403)
        if report.status != "ready" or not report.file_path:
            flash("Report is not ready yet.", "warning")
            return redirect(url_for("drawings", project_id=project.id, tab="reports"))

        docx_path = os.path.join(app.config["REPORTS_FOLDER"], report.file_path)
        if not os.path.exists(docx_path):
            abort(404)

        doc = _DocxDoc(docx_path)

        def _iter_blocks(document):
            from docx.oxml.ns import qn
            for child in document.element.body:
                tag = child.tag.split("}")[-1]
                if tag == "p":
                    yield _DocxPara(child, document)
                elif tag == "tbl":
                    yield _DocxTable(child, document)

        # Parse docx into sections: [(sheet_name, [items])]
        sections = [("Summary", [])]
        seen_first_h1 = False
        used_names = {"Summary": 1}

        for block in _iter_blocks(doc):
            if isinstance(block, _DocxPara):
                style = (block.style.name or "") if block.style else ""
                text = block.text.strip()
                if not text:
                    continue

                if "Heading 1" in style:
                    if not seen_first_h1:
                        seen_first_h1 = True
                        sections[0][1].append({"type": "text", "text": text, "bold": True, "size": 14})
                    else:
                        raw = text[:31].strip()
                        count = used_names.get(raw, 0) + 1
                        used_names[raw] = count
                        tab_name = raw if count == 1 else f"{raw[:28]} {count}"
                        sections.append((tab_name, []))
                elif "Heading 2" in style:
                    sections[-1][1].append({"type": "text", "text": text, "bold": True, "size": 11})
                elif "Heading 3" in style:
                    sections[-1][1].append({"type": "text", "text": text, "bold": True, "size": 10})
                elif "List" in style:
                    sections[-1][1].append({"type": "text", "text": "• " + text, "bold": False, "size": 10})
                else:
                    sections[-1][1].append({"type": "text", "text": text, "bold": False, "size": 10})

            elif isinstance(block, _DocxTable):
                rows = []
                for row in block.rows:
                    rows.append([cell.text.strip() for cell in row.cells])
                if rows:
                    sections[-1][1].append({"type": "table", "rows": rows})

        # Build workbook
        BRAND_HEX = "1F4E79"
        ACCENT_HEX = "D6E4F0"

        wb = openpyxl.Workbook()
        wb.remove(wb.active)

        hdr_fill = PatternFill("solid", fgColor=BRAND_HEX)
        hdr_font = Font(name="Calibri", bold=True, color="FFFFFF", size=11)
        alt_fill = PatternFill("solid", fgColor=ACCENT_HEX)

        for section_name, items in sections:
            if not any(i.get("text") or i.get("rows") for i in items):
                continue
            ws = wb.create_sheet(title=section_name)
            row_idx = 1

            for item in items:
                if item["type"] == "text":
                    text = item.get("text", "")
                    if not text:
                        continue
                    cell = ws.cell(row=row_idx, column=1, value=text)
                    cell.font = Font(name="Calibri", bold=item.get("bold", False), size=item.get("size", 10))
                    cell.alignment = Alignment(wrap_text=True)
                    row_idx += 1

                elif item["type"] == "table":
                    rows = item["rows"]
                    if not rows:
                        continue
                    ncols = max(len(r) for r in rows)

                    # Column widths based on content
                    for c in range(ncols):
                        col_letter = get_column_letter(c + 1)
                        max_w = max((len(str(r[c])) for r in rows if c < len(r)), default=10)
                        new_w = min(max_w + 4, 55)
                        current_w = ws.column_dimensions[col_letter].width or 0
                        ws.column_dimensions[col_letter].width = max(current_w, new_w)

                    # Header row
                    for c_idx, hdr in enumerate(rows[0], 1):
                        cell = ws.cell(row=row_idx, column=c_idx, value=hdr)
                        cell.font = hdr_font
                        cell.fill = hdr_fill
                        cell.alignment = Alignment(horizontal="center", wrap_text=True)
                    row_idx += 1

                    # Data rows
                    for r_num, data_row in enumerate(rows[1:]):
                        fill = alt_fill if r_num % 2 == 0 else PatternFill()
                        for c_idx in range(1, ncols + 1):
                            val = data_row[c_idx - 1] if c_idx - 1 < len(data_row) else ""
                            cell = ws.cell(row=row_idx, column=c_idx, value=val)
                            cell.font = Font(name="Calibri", size=10)
                            cell.fill = fill
                            cell.alignment = Alignment(wrap_text=True)
                        row_idx += 1

                    row_idx += 1  # blank spacer after table

            # Ensure column A is readable for text-heavy sheets
            if ws.column_dimensions["A"].width < 60:
                ws.column_dimensions["A"].width = 80

        if not wb.sheetnames:
            ws = wb.create_sheet("Report")
            ws.cell(row=1, column=1, value="No structured content available.")

        out = io.BytesIO()
        wb.save(out)
        out.seek(0)

        xlsx_name = os.path.splitext(report.file_path)[0] + ".xlsx"
        return send_file(
            out,
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            as_attachment=True,
            download_name=xlsx_name,
        )

    @app.route("/reports/<int:report_id>/status")
    @login_required
    def report_status(report_id):
        report = db.session.get(Report, report_id) or abort(404)
        project = db.session.get(Project, report.project_id) or abort(404)
        if not current_user.is_superadmin and current_user.company_id != project.company_id:
            abort(403)
        return jsonify({
            "id": report.id,
            "status": report.status,
            "file_path": report.file_path,
            "error_message": report.error_message,
        })

    @app.route("/reports/<int:report_id>/delete", methods=["POST"])
    @login_required
    def delete_report(report_id):
        report = db.session.get(Report, report_id) or abort(404)
        project = db.session.get(Project, report.project_id) or abort(404)
        if not current_user.is_superadmin and current_user.company_id != project.company_id:
            abort(403)
        if report.file_path:
            path = os.path.join(app.config["REPORTS_FOLDER"], report.file_path)
            try:
                os.remove(path)
            except OSError:
                pass
        db.session.delete(report)
        db.session.commit()
        flash("Report deleted.", "success")
        return redirect(url_for("drawings", project_id=project.id, tab="reports"))

    @app.route("/projects/<int:project_id>/history/export")
    @login_required
    def export_search_history(project_id):
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, PageBreak, HRFlowable,
        )
        from reportlab.lib import colors
        from xml.sax.saxutils import escape as xml_escape

        project, entries = _project_history(project_id)

        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf, pagesize=letter,
            leftMargin=0.75 * inch, rightMargin=0.75 * inch,
            topMargin=0.75 * inch, bottomMargin=0.75 * inch,
            title=f"Search History — {project.name}",
        )
        styles = getSampleStyleSheet()
        h_title = styles["Title"]
        h_meta = ParagraphStyle("meta", parent=styles["Normal"], textColor=colors.grey, fontSize=9, spaceAfter=12)
        h_q = ParagraphStyle("q", parent=styles["Heading3"], textColor=colors.HexColor("#0d6efd"), spaceAfter=4)
        h_entry_meta = ParagraphStyle("em", parent=styles["Normal"], textColor=colors.grey, fontSize=8, spaceAfter=6)
        h_a = ParagraphStyle("a", parent=styles["BodyText"], leading=14, spaceAfter=10)

        def p(text, style):
            return Paragraph(xml_escape(text or "").replace("\n", "<br/>"), style)

        story = [
            Paragraph(f"Search History — {xml_escape(project.name)}", h_title),
            Paragraph(
                f"{xml_escape(project.company.name)} &middot; {len(entries)} search(es) &middot; "
                f"Generated {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}",
                h_meta,
            ),
            HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey, spaceAfter=12),
        ]

        if not entries:
            story.append(Paragraph("No searches recorded for this project.", styles["Italic"]))
        else:
            for i, e in enumerate(entries, start=1):
                who = e.user.username if e.user else "unknown"
                when = e.created_at.strftime("%Y-%m-%d %H:%M UTC")
                filt = f" &middot; filter: {xml_escape(e.doc_type_filter)}" if e.doc_type_filter else ""
                story.append(p(f"{i}. {e.query}", h_q))
                story.append(Paragraph(f"{xml_escape(who)} &middot; {when}{filt}", h_entry_meta))
                story.append(p(e.answer or "(no answer recorded)", h_a))
                if i < len(entries):
                    story.append(HRFlowable(width="100%", thickness=0.3, color=colors.whitesmoke, spaceAfter=10))

        doc.build(story)
        buf.seek(0)
        fname = f"search-history-{project.id}-{datetime.now(timezone.utc).strftime('%Y%m%d-%H%M')}.pdf"
        return send_file(
            buf,
            mimetype="application/pdf",
            as_attachment=True,
            download_name=fname,
        )

    # ── Search ──────────────────────────────────────────────────

    @app.route("/search", methods=["GET", "POST"])
    @login_required
    def search():
        if not current_user.is_superadmin:
            abort(403)
        projects_list = (
            Project.query.join(Company).order_by(Company.name, Project.name).all()
        )

        results = None
        query = ""
        selected_project_id = None
        search_doc_type = ""
        if request.method == "POST":
            query = request.form.get("query", "").strip()
            selected_project_id = request.form.get("project_id", type=int)
            search_doc_type = request.form.get("search_doc_type", "").strip()
            if search_doc_type and search_doc_type not in DOC_TYPES:
                search_doc_type = ""
            if not query:
                flash("Please enter a search query.", "danger")
            elif not selected_project_id:
                flash("Please select a project.", "danger")
            elif not app.config["ANTHROPIC_API_KEY"]:
                flash("Claude API key not configured. Set ANTHROPIC_API_KEY environment variable.", "danger")
            else:
                project = db.session.get(Project, selected_project_id)
                if not project:
                    flash("Project not found.", "danger")
                elif not current_user.is_superadmin and current_user.company_id != project.company_id:
                    abort(403)
                else:
                    results = search_drawings(
                        query,
                        selected_project_id,
                        app.config["ANTHROPIC_API_KEY"],
                        app.config["PROCESSED_FOLDER"],
                        doc_type=search_doc_type or None,
                    )
                    history = SearchHistory(
                        project_id=selected_project_id,
                        user_id=current_user.id,
                        query=query,
                        answer=results.get("answer", "") if results else "",
                        doc_type_filter=search_doc_type or None,
                    )
                    db.session.add(history)
                    db.session.commit()

        return render_template(
            "search.html",
            results=results,
            query=query,
            projects=projects_list,
            selected_project_id=selected_project_id,
            doc_types=DOC_TYPES,
            search_doc_type=search_doc_type,
        )

    # ── Admin: User Management ──────────────────────────────────

    @app.route("/admin/users")
    @login_required
    @admin_required
    def admin_users():
        if current_user.is_superadmin:
            users = User.query.order_by(User.username).all()
        else:
            users = User.query.filter_by(company_id=current_user.company_id).order_by(User.username).all()
        companies = Company.query.order_by(Company.name).all()
        return render_template("admin/users.html", users=users, companies=companies, roles=ROLES)

    @app.route("/admin/users/new", methods=["GET", "POST"])
    @login_required
    @admin_required
    def new_user():
        companies = Company.query.order_by(Company.name).all()
        if request.method == "POST":
            username = request.form.get("username", "").strip()
            email = request.form.get("email", "").strip()
            password = request.form.get("password", "")
            role = request.form.get("role", ROLE_USER)
            company_id = request.form.get("company_id", type=int)

            if not all([username, email, password]):
                flash("All fields are required.", "danger")
            elif User.query.filter_by(username=username).first():
                flash("Username already exists.", "danger")
            elif User.query.filter_by(email=email).first():
                flash("Email already exists.", "danger")
            else:
                if not current_user.is_superadmin:
                    role = ROLE_USER
                    company_id = current_user.company_id

                user = User(username=username, email=email, role=role, company_id=company_id)
                user.set_password(password)
                db.session.add(user)
                db.session.commit()
                flash(f"User '{username}' created.", "success")
                return redirect(url_for("admin_users"))

        allowed_roles = ROLES if current_user.is_superadmin else [ROLE_USER]
        return render_template("admin/user_form.html", companies=companies, roles=allowed_roles)

    @app.route("/admin/users/<int:user_id>/edit", methods=["GET", "POST"])
    @login_required
    @admin_required
    def edit_user(user_id):
        user = db.session.get(User, user_id) or abort(404)
        companies = Company.query.order_by(Company.name).all()

        if request.method == "POST":
            user.email = request.form.get("email", user.email).strip()
            new_password = request.form.get("password", "").strip()
            if new_password:
                user.set_password(new_password)
            if current_user.is_superadmin:
                user.role = request.form.get("role", user.role)
                user.company_id = request.form.get("company_id", type=int)
            db.session.commit()
            flash(f"User '{user.username}' updated.", "success")
            return redirect(url_for("admin_users"))

        allowed_roles = ROLES if current_user.is_superadmin else [ROLE_USER]
        return render_template("admin/user_form.html", user=user, companies=companies, roles=allowed_roles)

    @app.route("/admin/users/<int:user_id>/delete", methods=["POST"])
    @login_required
    @superadmin_required
    def delete_user(user_id):
        user = db.session.get(User, user_id) or abort(404)
        if user.id == current_user.id:
            flash("You cannot delete yourself.", "danger")
            return redirect(url_for("admin_users"))
        db.session.delete(user)
        db.session.commit()
        flash(f"User '{user.username}' deleted.", "success")
        return redirect(url_for("admin_users"))

    # ── Admin Rates ─────────────────────────────────────────────

    @app.route("/admin/rates")
    @login_required
    @superadmin_required
    def admin_rates():
        active_labor = LaborRate.query.filter_by(active=True).order_by(LaborRate.category, LaborRate.craft_type).all()
        active_insurance = InsuranceRate.query.filter_by(active=True).order_by(InsuranceRate.category, InsuranceRate.rate_type).all()
        history_labor = LaborRate.query.order_by(LaborRate.created_at.desc()).limit(200).all()
        history_insurance = InsuranceRate.query.order_by(InsuranceRate.created_at.desc()).limit(200).all()
        active_tab = request.args.get("tab", "labor")
        return render_template(
            "admin/rates.html",
            active_labor=active_labor,
            active_insurance=active_insurance,
            history_labor=history_labor,
            history_insurance=history_insurance,
            active_tab=active_tab,
        )

    @app.route("/admin/rates/labor/upload", methods=["POST"])
    @login_required
    @superadmin_required
    def upload_labor_rates():
        f = request.files.get("csv_file")
        if not f:
            flash("No file uploaded.", "danger")
            return redirect(url_for("admin_rates", tab="labor"))
        try:
            content = f.read().decode("utf-8-sig")
            reader = _csv.DictReader(io.StringIO(content))
            required = {"category", "craft_type", "region", "hourly_cost"}
            if not required.issubset(set(reader.fieldnames or [])):
                flash(f"CSV missing columns: {required - set(reader.fieldnames or [])}", "danger")
                return redirect(url_for("admin_rates", tab="labor"))
            rows = list(reader)
            if not rows:
                flash("CSV is empty.", "danger")
                return redirect(url_for("admin_rates", tab="labor"))
            max_ver = db.session.query(db.func.max(LaborRate.version)).scalar() or 0
            new_ver = max_ver + 1
            LaborRate.query.update({"active": False})
            count = 0
            for row in rows:
                try:
                    cost = float(row["hourly_cost"])
                except (ValueError, KeyError):
                    continue
                eff = exp = None
                try:
                    if row.get("effective_date"):
                        eff = datetime.strptime(row["effective_date"].strip(), "%Y-%m-%d").date()
                    if row.get("expiry_date"):
                        exp = datetime.strptime(row["expiry_date"].strip(), "%Y-%m-%d").date()
                except ValueError:
                    pass
                db.session.add(LaborRate(
                    category=row.get("category", "").strip(),
                    craft_type=row.get("craft_type", "").strip(),
                    region=row.get("region", "General").strip(),
                    hourly_cost=cost,
                    effective_date=eff,
                    expiry_date=exp,
                    version=new_ver,
                    uploaded_by=current_user.id,
                    active=True,
                ))
                count += 1
            db.session.commit()
            flash(f"Uploaded {count} labor rates (version {new_ver}).", "success")
        except Exception as e:
            db.session.rollback()
            flash(f"Upload failed: {e}", "danger")
        return redirect(url_for("admin_rates", tab="labor"))

    @app.route("/admin/rates/insurance/upload", methods=["POST"])
    @login_required
    @superadmin_required
    def upload_insurance_rates():
        f = request.files.get("csv_file")
        if not f:
            flash("No file uploaded.", "danger")
            return redirect(url_for("admin_rates", tab="insurance"))
        try:
            content = f.read().decode("utf-8-sig")
            reader = _csv.DictReader(io.StringIO(content))
            required = {"category", "rate_type", "rate_percent"}
            if not required.issubset(set(reader.fieldnames or [])):
                flash(f"CSV missing columns: {required - set(reader.fieldnames or [])}", "danger")
                return redirect(url_for("admin_rates", tab="insurance"))
            rows = list(reader)
            if not rows:
                flash("CSV is empty.", "danger")
                return redirect(url_for("admin_rates", tab="insurance"))
            max_ver = db.session.query(db.func.max(InsuranceRate.version)).scalar() or 0
            new_ver = max_ver + 1
            InsuranceRate.query.update({"active": False})
            count = 0
            for row in rows:
                try:
                    pct = float(row["rate_percent"])
                except (ValueError, KeyError):
                    continue
                eff = exp = None
                try:
                    if row.get("effective_date"):
                        eff = datetime.strptime(row["effective_date"].strip(), "%Y-%m-%d").date()
                    if row.get("expiry_date"):
                        exp = datetime.strptime(row["expiry_date"].strip(), "%Y-%m-%d").date()
                except ValueError:
                    pass
                db.session.add(InsuranceRate(
                    category=row.get("category", "").strip(),
                    rate_type=row.get("rate_type", "").strip(),
                    rate_percent=pct,
                    effective_date=eff,
                    expiry_date=exp,
                    version=new_ver,
                    notes=row.get("notes", "").strip() or None,
                    uploaded_by=current_user.id,
                    active=True,
                ))
                count += 1
            db.session.commit()
            flash(f"Uploaded {count} insurance rates (version {new_ver}).", "success")
        except Exception as e:
            db.session.rollback()
            flash(f"Upload failed: {e}", "danger")
        return redirect(url_for("admin_rates", tab="insurance"))

    @app.route("/admin/rates/labor/template")
    @login_required
    @superadmin_required
    def download_labor_template():
        headers = ["category", "craft_type", "region", "hourly_cost", "effective_date", "expiry_date"]
        buf = io.StringIO()
        w = _csv.DictWriter(buf, fieldnames=headers)
        w.writeheader()
        w.writerow({"category": "Painting", "craft_type": "Journeyman Painter",
                    "region": "Southern California", "hourly_cost": "75.50",
                    "effective_date": "2025-01-01", "expiry_date": "2025-12-31"})
        return send_file(
            io.BytesIO(buf.getvalue().encode()),
            mimetype="text/csv",
            as_attachment=True,
            download_name="labor_rates_template.csv",
        )

    @app.route("/admin/rates/insurance/template")
    @login_required
    @superadmin_required
    def download_insurance_template():
        headers = ["category", "rate_type", "rate_percent", "effective_date", "expiry_date", "notes"]
        buf = io.StringIO()
        w = _csv.DictWriter(buf, fieldnames=headers)
        w.writeheader()
        w.writerow({"category": "General Liability", "rate_type": "Standard GL",
                    "rate_percent": "1.25", "effective_date": "2025-01-01",
                    "expiry_date": "2025-12-31", "notes": "Per AIA A201"})
        return send_file(
            io.BytesIO(buf.getvalue().encode()),
            mimetype="text/csv",
            as_attachment=True,
            download_name="insurance_rates_template.csv",
        )

    # ── Activity History ─────────────────────────────────────────

    @app.route("/admin/history")
    @login_required
    @superadmin_required
    def admin_history():
        uid = request.args.get("user_id", type=int)
        cid = request.args.get("company_id", type=int)
        atype = request.args.get("activity_type", "")
        date_from = request.args.get("date_from", "")
        date_to = request.args.get("date_to", "")

        dt_from = dt_to = None
        if date_from:
            try: dt_from = datetime.strptime(date_from, "%Y-%m-%d")
            except ValueError: pass
        if date_to:
            try: dt_to = datetime.strptime(date_to, "%Y-%m-%d").replace(hour=23, minute=59, second=59)
            except ValueError: pass

        LIMIT = 500
        events = []

        if not atype or atype == "search":
            q = db.session.query(SearchHistory)
            if uid: q = q.filter(SearchHistory.user_id == uid)
            if cid: q = q.join(User, SearchHistory.user_id == User.id).filter(User.company_id == cid)
            if dt_from: q = q.filter(SearchHistory.created_at >= dt_from)
            if dt_to: q = q.filter(SearchHistory.created_at <= dt_to)
            for s in q.order_by(SearchHistory.created_at.desc()).limit(LIMIT).all():
                events.append({
                    "type": "search", "date": s.created_at, "user": s.user,
                    "detail": s.query, "full_answer": s.answer,
                    "context": s.project.name if s.project else "—",
                    "context2": s.doc_type_filter or "All types", "id": s.id,
                })

        if not atype or atype == "report":
            q = db.session.query(Report)
            if uid: q = q.filter(Report.user_id == uid)
            if cid: q = q.join(User, Report.user_id == User.id).filter(User.company_id == cid)
            if dt_from: q = q.filter(Report.created_at >= dt_from)
            if dt_to: q = q.filter(Report.created_at <= dt_to)
            for r in q.order_by(Report.created_at.desc()).limit(LIMIT).all():
                events.append({
                    "type": "report", "date": r.created_at, "user": r.user,
                    "detail": r.template_name,
                    "context": r.project.name if r.project else "—",
                    "context2": r.status, "id": r.id,
                    "report_status": r.status, "report_file": r.file_path,
                })

        if not atype or atype == "upload":
            q = db.session.query(Drawing)
            if uid: q = q.filter(Drawing.uploaded_by == uid)
            if cid: q = q.join(User, Drawing.uploaded_by == User.id).filter(User.company_id == cid)
            if dt_from: q = q.filter(Drawing.created_at >= dt_from)
            if dt_to: q = q.filter(Drawing.created_at <= dt_to)
            for d in q.order_by(Drawing.created_at.desc()).limit(LIMIT).all():
                events.append({
                    "type": "upload", "date": d.created_at, "user": d.uploader,
                    "detail": d.original_filename,
                    "context": d.project.name if d.project else "—",
                    "context2": d.doc_type, "id": d.id,
                })

        if not atype or atype == "login":
            q = db.session.query(LoginEvent)
            if uid: q = q.filter(LoginEvent.user_id == uid)
            if cid: q = q.join(User, LoginEvent.user_id == User.id).filter(User.company_id == cid)
            if dt_from: q = q.filter(LoginEvent.created_at >= dt_from)
            if dt_to: q = q.filter(LoginEvent.created_at <= dt_to)
            for l in q.order_by(LoginEvent.created_at.desc()).limit(LIMIT).all():
                events.append({
                    "type": "login", "date": l.created_at, "user": l.user,
                    "detail": l.ip_address or "—",
                    "context": "—", "context2": "—", "id": l.id,
                })

        events.sort(key=lambda e: e["date"] or datetime.min, reverse=True)
        total = len(events)
        events = events[:500]

        counts = {}
        for e in events:
            counts[e["type"]] = counts.get(e["type"], 0) + 1

        return render_template("admin/history.html",
            events=events, total=total, counts=counts,
            all_users=User.query.order_by(User.username).all(),
            all_companies=Company.query.order_by(Company.name).all(),
            filters={"user_id": uid, "company_id": cid, "activity_type": atype,
                     "date_from": date_from, "date_to": date_to},
        )

    @app.route("/admin/history/export")
    @login_required
    @superadmin_required
    def admin_history_export():
        from reportlab.lib.pagesizes import letter, landscape
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors as rl_colors
        from xml.sax.saxutils import escape as xml_escape

        uid = request.args.get("user_id", type=int)
        cid = request.args.get("company_id", type=int)
        atype = request.args.get("activity_type", "")
        date_from = request.args.get("date_from", "")
        date_to = request.args.get("date_to", "")

        dt_from = dt_to = None
        if date_from:
            try: dt_from = datetime.strptime(date_from, "%Y-%m-%d")
            except ValueError: pass
        if date_to:
            try: dt_to = datetime.strptime(date_to, "%Y-%m-%d").replace(hour=23, minute=59, second=59)
            except ValueError: pass

        LIMIT = 500
        events = []

        if not atype or atype == "search":
            q = db.session.query(SearchHistory)
            if uid: q = q.filter(SearchHistory.user_id == uid)
            if cid: q = q.join(User, SearchHistory.user_id == User.id).filter(User.company_id == cid)
            if dt_from: q = q.filter(SearchHistory.created_at >= dt_from)
            if dt_to: q = q.filter(SearchHistory.created_at <= dt_to)
            for s in q.order_by(SearchHistory.created_at.desc()).limit(LIMIT).all():
                events.append(("Search", s.created_at,
                    s.user.username if s.user else "—",
                    (s.user.company.name if s.user and s.user.company else "—"),
                    (s.query[:150] + "…") if len(s.query) > 150 else s.query,
                    s.project.name if s.project else "—"))

        if not atype or atype == "report":
            q = db.session.query(Report)
            if uid: q = q.filter(Report.user_id == uid)
            if cid: q = q.join(User, Report.user_id == User.id).filter(User.company_id == cid)
            if dt_from: q = q.filter(Report.created_at >= dt_from)
            if dt_to: q = q.filter(Report.created_at <= dt_to)
            for r in q.order_by(Report.created_at.desc()).limit(LIMIT).all():
                events.append(("Report", r.created_at,
                    r.user.username if r.user else "—",
                    (r.user.company.name if r.user and r.user.company else "—"),
                    r.template_name, r.project.name if r.project else "—"))

        if not atype or atype == "upload":
            q = db.session.query(Drawing)
            if uid: q = q.filter(Drawing.uploaded_by == uid)
            if cid: q = q.join(User, Drawing.uploaded_by == User.id).filter(User.company_id == cid)
            if dt_from: q = q.filter(Drawing.created_at >= dt_from)
            if dt_to: q = q.filter(Drawing.created_at <= dt_to)
            for d in q.order_by(Drawing.created_at.desc()).limit(LIMIT).all():
                events.append(("Upload", d.created_at,
                    d.uploader.username if d.uploader else "—",
                    (d.uploader.company.name if d.uploader and d.uploader.company else "—"),
                    d.original_filename, d.project.name if d.project else "—"))

        if not atype or atype == "login":
            q = db.session.query(LoginEvent)
            if uid: q = q.filter(LoginEvent.user_id == uid)
            if cid: q = q.join(User, LoginEvent.user_id == User.id).filter(User.company_id == cid)
            if dt_from: q = q.filter(LoginEvent.created_at >= dt_from)
            if dt_to: q = q.filter(LoginEvent.created_at <= dt_to)
            for l in q.order_by(LoginEvent.created_at.desc()).limit(LIMIT).all():
                events.append(("Login", l.created_at,
                    l.user.username if l.user else "—",
                    (l.user.company.name if l.user and l.user.company else "—"),
                    l.ip_address or "—", "—"))

        events.sort(key=lambda e: e[1] or datetime.min, reverse=True)

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=landscape(letter),
            leftMargin=0.5*inch, rightMargin=0.5*inch,
            topMargin=0.75*inch, bottomMargin=0.75*inch)
        styles = getSampleStyleSheet()
        BRAND = rl_colors.HexColor(0x1F4E79)

        story = []
        title_style = ParagraphStyle("Title", parent=styles["Heading1"],
            textColor=BRAND, fontSize=16, spaceAfter=6)
        story.append(Paragraph("User Activity History", title_style))

        sub_parts = []
        if date_from or date_to:
            sub_parts.append(f"{date_from or 'start'} → {date_to or 'today'}")
        if atype:
            sub_parts.append(f"Type: {atype.capitalize()}")
        if sub_parts:
            story.append(Paragraph(" | ".join(sub_parts), styles["Normal"]))
        story.append(Paragraph(
            f"Generated {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')} · {len(events)} events",
            styles["Normal"]))
        story.append(Spacer(1, 0.2*inch))

        headers = ["Type", "Date / Time", "User", "Company", "Detail", "Project"]
        table_data = [headers]
        for evtype, evdate, user, company, detail, ctx in events:
            table_data.append([
                evtype,
                evdate.strftime("%Y-%m-%d %H:%M") if evdate else "—",
                user, company,
                Paragraph(xml_escape(str(detail)), styles["Normal"]),
                ctx,
            ])

        col_widths = [0.7*inch, 1.1*inch, 1.0*inch, 1.1*inch, 4.0*inch, 1.5*inch]
        tbl = Table(table_data, colWidths=col_widths, repeatRows=1)
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), BRAND),
            ("TEXTCOLOR", (0,0), (-1,0), rl_colors.white),
            ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
            ("FONTSIZE", (0,0), (-1,-1), 8),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [rl_colors.white, rl_colors.HexColor(0xF0F4F8)]),
            ("GRID", (0,0), (-1,-1), 0.25, rl_colors.HexColor(0xCCCCCC)),
            ("VALIGN", (0,0), (-1,-1), "TOP"),
            ("TOPPADDING", (0,0), (-1,-1), 3),
            ("BOTTOMPADDING", (0,0), (-1,-1), 3),
        ]))
        story.append(tbl)

        doc.build(story)
        buf.seek(0)
        fname = f"activity-history-{datetime.now(timezone.utc).strftime('%Y%m%d-%H%M')}.pdf"
        return send_file(buf, mimetype="application/pdf", as_attachment=True, download_name=fname)

    # ── Error Handlers ──────────────────────────────────────────

    @app.errorhandler(403)
    def forbidden(e):
        return render_template("error.html", code=403, message="Access Denied"), 403

    @app.errorhandler(404)
    def not_found(e):
        return render_template("error.html", code=404, message="Page Not Found"), 404

    return app


if __name__ == "__main__":
    app = create_app()
    app.run(debug=False, host='0.0.0.0', port=5000)
