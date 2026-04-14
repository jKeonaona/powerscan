"""Document Intelligence report generator.

Takes a project's page images, runs them through Claude in batches using a
template prompt, synthesizes the batch responses into one structured payload,
and renders a branded .docx file via python-docx.
"""
import io
import json
import os
import re
import threading
import time
import traceback
from datetime import datetime, timezone

import anthropic
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from models import db, Drawing, DrawingPage, Project, Report, SearchHistory
from search import (
    CLAUDE_MODEL,
    MAX_IMAGES_PER_REQUEST,
    _build_batch_content,
)

POLL_INTERVAL = 3  # seconds
SUMMARY_MAX_WORDS = 200

# ── CCC brand constants (swap for real logo path later if desired) ──
BRAND_NAME = "CCC"
BRAND_TAGLINE = "Construction Intelligence Report"
BRAND_COLOR = RGBColor(0x1F, 0x4E, 0x79)   # deep blue
BRAND_ACCENT = RGBColor(0x8C, 0x8C, 0x8C)  # mid gray

SCHEMA_DESCRIPTION = """{
  "title": "string — report title",
  "summary": "string — 2-4 sentence executive summary",
  "sections": [
    // each entry is ONE of these section shapes:
    {"type": "heading", "text": "string", "level": 1-3},
    {"type": "paragraph", "text": "string"},
    {"type": "bullets", "items": ["string", ...]},
    {"type": "table", "headers": ["col1", "col2", ...], "rows": [["a","b"], ...]}
  ]
}"""

REPORT_TEMPLATES = {
    "submittal_list": {
        "name": "Submittal List",
        "description": "Every submittal required by the contract documents.",
        "prompt": (
            "Identify every submittal required by the project documents. For each submittal, "
            "extract: item name, specification section (if identifiable), submittal type "
            "(shop drawing / product data / sample / certificate / O&M manual / other), "
            "required timing or deadline if stated, and any special notes. "
            "Return results as a single table with columns: Item, Spec Section, Type, Timing, Notes. "
            "Start with a brief summary paragraph giving the total count and any submittals flagged "
            "as high-priority or potentially overdue."
        ),
    },
    "bid_summary": {
        "name": "Bid Summary",
        "description": "Contract overview, key dates, insurance, bonding, special provisions.",
        "prompt": (
            "Produce a bid summary covering: (1) contract/project overview, (2) key dates "
            "(bid due, award, NTP, substantial completion, final completion), (3) insurance "
            "requirements (types, limits, additional insureds), (4) bonding requirements "
            "(bid bond, performance bond, payment bond, percentages), (5) special provisions "
            "or unusual clauses the bidder must be aware of, (6) liquidated damages if any. "
            "Use headings and tables for dates, insurance, and bonding. Surface anything "
            "non-standard or unusually onerous in a clearly labeled section."
        ),
    },
    "compliance_checklist": {
        "name": "Compliance Checklist",
        "description": "Lead, environmental, and safety compliance requirements.",
        "prompt": (
            "Build a compliance checklist covering (1) lead compliance (RRP, EPA lead-safe, "
            "abatement requirements), (2) environmental requirements (hazardous materials, "
            "disposal, SWPPP, air quality, asbestos), (3) safety requirements "
            "(OSHA, site-specific safety plans, PPE, training). For each item, capture: "
            "requirement, source reference (spec section or page), and compliance action "
            "the contractor must take. Present as three tables — one per category — with "
            "columns: Requirement, Source, Action."
        ),
    },
    "risk_assessment": {
        "name": "Risk Assessment",
        "description": "Problematic clauses, non-standard terms, cost-impact risks.",
        "prompt": (
            "Produce a risk assessment flagging: (1) problematic clauses (unlimited "
            "indemnification, broad consequential damages, unusual warranty terms), "
            "(2) non-standard contract terms that deviate from AIA/ConsensusDocs norms, "
            "(3) potential cost impacts (unclear scope, coordination risks, schedule pressure, "
            "escalation exposure). For each risk, provide: description, severity "
            "(High / Medium / Low), source reference, and recommended mitigation. "
            "Present as a table with columns: Risk, Severity, Source, Mitigation. "
            "Start with an executive summary listing the top three risks."
        ),
    },
}


# ───────────────────────── Claude calls ─────────────────────────

def _parse_json_response(text):
    """Strip fences, extract first JSON object, parse. Fall back to paragraph."""
    text = (text or "").strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```\s*$", "", text)
    match = re.search(r"\{.*\}", text, re.DOTALL)
    if match:
        text = match.group(0)
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        return {
            "title": "Report",
            "summary": "",
            "sections": [{"type": "paragraph", "text": text}],
        }


def _ask_report_batch(client, project, template_prompt, batch_pages, processed_folder,
                     start_index, batch_num, total_batches):
    content, index_map = _build_batch_content(batch_pages, processed_folder, start_index)
    if not content:
        return None, {}

    batch_note = ""
    if total_batches > 1:
        batch_note = (
            f"\n\nThis is batch {batch_num} of {total_batches} from project "
            f"'{project.name}'. Report ONLY on what is visible in the images above. "
            f"A later step will merge your output with the other batches, so omit any "
            f"item you cannot confirm from this batch."
        )

    instructions = (
        f"{template_prompt}{batch_note}\n\n"
        f"Return a single JSON object matching this schema exactly:\n"
        f"{SCHEMA_DESCRIPTION}\n\n"
        f"Return ONLY the JSON. No markdown fences, no commentary, no leading or trailing text."
    )
    content.append({"type": "text", "text": instructions})

    message = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=4096,
        messages=[{"role": "user", "content": content}],
    )
    return _parse_json_response(message.content[0].text), index_map


def _summarize_batch_result(batch_report, max_words=SUMMARY_MAX_WORDS):
    """Flatten a structured batch JSON to a compact ≤max_words plain-text summary.

    Preserves the semantic skeleton (summary line, headings, a few bullets per
    list, a couple of rows per table) but drops long prose and deep tables so
    the synthesis call stays well under token limits.
    """
    parts = []
    summary_text = (batch_report.get("summary") or "").strip()
    if summary_text:
        parts.append(summary_text)

    for section in batch_report.get("sections", []) or []:
        if not isinstance(section, dict):
            continue
        stype = (section.get("type") or "").lower()

        if stype == "heading":
            text = str(section.get("text", "")).strip()
            if text:
                parts.append(f"## {text}")

        elif stype == "paragraph":
            text = str(section.get("text", "")).strip()
            if text:
                parts.append(text)

        elif stype == "bullets":
            items = [str(i).strip() for i in (section.get("items", []) or []) if str(i).strip()]
            if items:
                parts.append(" • " + " • ".join(items[:5]))

        elif stype == "table":
            headers = [str(h).strip() for h in (section.get("headers", []) or [])]
            rows = section.get("rows", []) or []
            if headers:
                parts.append("[" + " | ".join(headers) + "]")
            for row in rows[:3]:
                if isinstance(row, (list, tuple)):
                    cells = [str(c).strip() for c in row[:len(headers) or len(row)]]
                    parts.append("- " + " | ".join(cells))
            if len(rows) > 3:
                parts.append(f"(+{len(rows) - 3} more rows)")

    text = " ".join(p for p in parts if p)
    words = text.split()
    if len(words) > max_words:
        text = " ".join(words[:max_words]) + " …"
    return text


def _synthesize_reports(client, project, template_name, template_prompt, batch_summaries):
    """Combine per-batch text summaries into a single structured JSON report."""
    numbered = "\n\n".join(
        f"--- BATCH {i} SUMMARY ---\n{s}" for i, s in enumerate(batch_summaries, start=1) if s
    )

    prompt = (
        f"You previously asked {len(batch_summaries)} batches of project documents "
        f"(project: '{project.name}') for a '{template_name}' report. Each batch produced "
        f"its own structured result; those have been condensed into short summaries "
        f"(≤{SUMMARY_MAX_WORDS} words each) to fit within token limits.\n\n"
        f"TEMPLATE INTENT:\n{template_prompt}\n\n"
        f"PER-BATCH SUMMARIES:\n{numbered}\n\n"
        f"Produce ONE final report by merging the information above. Return a single JSON "
        f"object matching this schema EXACTLY:\n{SCHEMA_DESCRIPTION}\n\n"
        f"Rules:\n"
        f"1. Deduplicate items that appear in multiple batches.\n"
        f"2. Reconstruct tables where the summaries referenced them. Use your best judgment "
        f"for column structure when a batch only gave partial rows.\n"
        f"3. Reconcile contradictions or list them if unresolvable.\n"
        f"4. Write a unified executive summary covering the whole project.\n"
        f"5. Do NOT mention 'batches' or 'summaries' in the final output.\n\n"
        f"Return ONLY the final JSON. No markdown fences, no commentary."
    )
    message = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=8192,
        messages=[{"role": "user", "content": prompt}],
    )
    return _parse_json_response(message.content[0].text)


# ───────────────────────── DOCX rendering ─────────────────────────

def _set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_cover_page(doc, template_name, project_name, company_name):
    # Brand wordmark
    mark = doc.add_paragraph()
    mark.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = mark.add_run(BRAND_NAME)
    run.font.size = Pt(48)
    run.font.bold = True
    run.font.color.rgb = BRAND_COLOR

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag_run = tag.add_run(BRAND_TAGLINE.upper())
    tag_run.font.size = Pt(11)
    tag_run.font.color.rgb = BRAND_ACCENT
    tag_run.font.bold = True

    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title.add_run(template_name)
    t_run.font.size = Pt(32)
    t_run.font.bold = True
    t_run.font.color.rgb = BRAND_COLOR

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = sub.add_run(project_name)
    s_run.font.size = Pt(18)
    s_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    if company_name:
        co = doc.add_paragraph()
        co.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c_run = co.add_run(company_name)
        c_run.font.size = Pt(13)
        c_run.font.color.rgb = BRAND_ACCENT

    for _ in range(4):
        doc.add_paragraph()

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d_run = date_p.add_run(datetime.now(timezone.utc).strftime("%B %d, %Y"))
    d_run.font.size = Pt(12)
    d_run.font.color.rgb = BRAND_ACCENT


def _style_headings(doc):
    for level, size in ((1, 18), (2, 14), (3, 12)):
        try:
            style = doc.styles[f"Heading {level}"]
            style.font.color.rgb = BRAND_COLOR
            style.font.size = Pt(size)
            style.font.bold = True
        except KeyError:
            pass


def _render_section(doc, section):
    stype = (section.get("type") or "paragraph").lower()

    if stype == "heading":
        level = max(1, min(3, int(section.get("level", 2) or 2)))
        doc.add_heading(str(section.get("text", "")), level=level)

    elif stype == "paragraph":
        doc.add_paragraph(str(section.get("text", "")))

    elif stype == "bullets":
        for item in section.get("items", []) or []:
            doc.add_paragraph(str(item), style="List Bullet")

    elif stype == "table":
        headers = section.get("headers", []) or []
        rows = section.get("rows", []) or []
        if not headers:
            return
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        try:
            table.style = "Light Grid Accent 1"
        except KeyError:
            pass

        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = str(h)
            _set_cell_bg(hdr_cells[i], "1F4E79")
            for para in hdr_cells[i].paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for r_idx, row in enumerate(rows, start=1):
            cells = table.rows[r_idx].cells
            for c_idx in range(len(headers)):
                val = row[c_idx] if c_idx < len(row) else ""
                if isinstance(val, (list, dict)):
                    val = json.dumps(val)
                cells[c_idx].text = str(val) if val is not None else ""
        doc.add_paragraph()

    else:
        # Unknown type — render as paragraph so nothing is lost.
        text = section.get("text") or json.dumps(section)
        doc.add_paragraph(str(text))


def _render_report_docx(report, template_name, project):
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    _style_headings(doc)
    _add_cover_page(doc, template_name, project.name, project.company.name)
    doc.add_page_break()

    # Body header
    title = report.get("title") or template_name
    doc.add_heading(str(title), level=1)

    summary = (report.get("summary") or "").strip()
    if summary:
        doc.add_heading("Executive Summary", level=2)
        doc.add_paragraph(summary)
        doc.add_paragraph()

    sections = report.get("sections") or []
    if not sections:
        doc.add_paragraph(
            "No structured content was returned for this report.",
        )
    else:
        for section in sections:
            if isinstance(section, dict):
                _render_section(doc, section)

    # Footer / generated note
    doc.add_paragraph()
    footer = doc.add_paragraph()
    f_run = footer.add_run(
        f"Generated by {BRAND_NAME} · {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    )
    f_run.italic = True
    f_run.font.size = Pt(9)
    f_run.font.color.rgb = BRAND_ACCENT

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ───────────────────────── Public enqueue + background worker ─────────────────────────

def _resolve_template(template_id, custom_prompt):
    """Return (template_name, prompt_text) or raise ValueError."""
    if template_id == "custom":
        prompt = (custom_prompt or "").strip()
        if not prompt:
            raise ValueError("Custom report requires a prompt.")
        return "Custom Report", prompt

    tpl = REPORT_TEMPLATES.get(template_id)
    if not tpl:
        raise ValueError(f"Unknown template: {template_id}")
    return tpl["name"], tpl["prompt"]


def _build_safe_filename(project_name, template_name, report_id):
    safe_project = re.sub(r"[^A-Za-z0-9_-]+", "_", project_name).strip("_") or "project"
    safe_template = re.sub(r"[^A-Za-z0-9_-]+", "_", template_name).strip("_") or "report"
    return f"{safe_project}-{safe_template}-{report_id}.docx"


def enqueue_report(project_id, user_id, template_id, custom_prompt):
    """Validate input, create a Report row in 'generating' state, return the id.

    Actual generation runs on the background worker thread. Raises ValueError
    for invalid input so the route can flash a friendly message.
    """
    project = db.session.get(Project, project_id)
    if not project:
        raise ValueError("Project not found.")

    template_name, _ = _resolve_template(template_id, custom_prompt)

    has_pages = (
        db.session.query(DrawingPage.id)
        .join(Drawing, DrawingPage.drawing_id == Drawing.id)
        .filter(Drawing.project_id == project_id)
        .filter(Drawing.status == "ready")
        .first()
    )
    if not has_pages:
        raise ValueError("No ready documents in this project. Upload and wait for conversion first.")

    report = Report(
        project_id=project_id,
        user_id=user_id,
        template_id=template_id,
        template_name=template_name,
        custom_prompt=(custom_prompt or "").strip() or None,
        status="generating",
    )
    db.session.add(report)
    db.session.commit()
    return report.id


def _process_report(app, report_id):
    """Runs inside the background worker thread with an app context open."""
    with app.app_context():
        report = db.session.get(Report, report_id)
        if not report:
            return

        project = db.session.get(Project, report.project_id)
        if not project:
            report.status = "failed"
            report.error_message = "Project no longer exists."
            report.completed_at = datetime.now(timezone.utc)
            db.session.commit()
            return

        try:
            _, prompt = _resolve_template(report.template_id, report.custom_prompt)

            pages = (
                db.session.query(DrawingPage, Drawing)
                .join(Drawing, DrawingPage.drawing_id == Drawing.id)
                .filter(Drawing.project_id == project.id)
                .filter(Drawing.status == "ready")
                .order_by(Drawing.original_filename, DrawingPage.page_number)
                .all()
            )
            if not pages:
                raise ValueError("No ready documents in this project.")

            api_key = app.config.get("ANTHROPIC_API_KEY") or os.getenv("ANTHROPIC_API_KEY", "")
            if not api_key:
                raise ValueError("ANTHROPIC_API_KEY is not configured.")

            processed_folder = app.config["PROCESSED_FOLDER"]
            reports_folder = app.config["REPORTS_FOLDER"]
            os.makedirs(reports_folder, exist_ok=True)

            client = anthropic.Anthropic(api_key=api_key)
            batches = [
                pages[i:i + MAX_IMAGES_PER_REQUEST]
                for i in range(0, len(pages), MAX_IMAGES_PER_REQUEST)
            ]
            total_batches = len(batches)
            print(
                f"[powerscan] report {report.id} ({report.template_id}): "
                f"{len(pages)} pages in {total_batches} batch(es)",
                flush=True,
            )

            batch_reports = []
            next_index = 1
            for batch_num, batch in enumerate(batches, start=1):
                print(f"[powerscan] report {report.id} batch {batch_num}/{total_batches}", flush=True)
                result, _ = _ask_report_batch(
                    client, project, prompt, batch, processed_folder,
                    start_index=next_index, batch_num=batch_num, total_batches=total_batches,
                )
                next_index += len(batch)
                if result:
                    batch_reports.append(result)

            if not batch_reports:
                raise ValueError("No content returned from Claude for any batch.")

            if total_batches == 1:
                final_report = batch_reports[0]
            else:
                print(
                    f"[powerscan] report {report.id}: summarizing "
                    f"{len(batch_reports)} batches (≤{SUMMARY_MAX_WORDS} words each)",
                    flush=True,
                )
                summaries = [_summarize_batch_result(br) for br in batch_reports]
                total_words = sum(len(s.split()) for s in summaries)
                print(
                    f"[powerscan] report {report.id}: synthesis input ≈ {total_words} words",
                    flush=True,
                )
                final_report = _synthesize_reports(
                    client, project, report.template_name, prompt, summaries,
                )

            docx_bytes = _render_report_docx(final_report, report.template_name, project)
            filename = _build_safe_filename(project.name, report.template_name, report.id)
            out_path = os.path.join(reports_folder, filename)
            with open(out_path, "wb") as f:
                f.write(docx_bytes)

            report.filename = filename
            report.status = "ready"
            report.completed_at = datetime.now(timezone.utc)
            db.session.commit()

            # Log to search history
            log_query = f"Report: {report.template_name}"
            if report.template_id == "custom" and report.custom_prompt:
                log_query = f"Custom Report: {report.custom_prompt[:200]}"
            log_answer = f"Generated .docx report: {filename}."
            final_summary = (final_report.get("summary") or "").strip()
            if final_summary:
                log_answer += f"\n\nSummary: {final_summary}"
            history = SearchHistory(
                project_id=project.id,
                user_id=report.user_id,
                query=log_query,
                answer=log_answer,
            )
            db.session.add(history)
            db.session.commit()

            print(f"[powerscan] report {report.id} ready -> {filename}", flush=True)

        except Exception as e:
            db.session.rollback()
            print(f"[powerscan] report {report_id} failed: {e}", flush=True)
            traceback.print_exc()
            # Re-fetch in case the session was rolled back
            report = db.session.get(Report, report_id)
            if report:
                report.status = "failed"
                report.error_message = str(e)[:500]
                report.completed_at = datetime.now(timezone.utc)
                db.session.commit()


def _worker_loop(app):
    # Recover anything stuck on startup — report generation isn't resumable.
    with app.app_context():
        stuck = Report.query.filter_by(status="generating").all()
        for r in stuck:
            print(f"[powerscan] recovering stuck report {r.id} as failed", flush=True)
            r.status = "failed"
            r.error_message = "Server restarted during generation."
            r.completed_at = datetime.now(timezone.utc)
        if stuck:
            db.session.commit()

    # We use a sentinel to avoid re-picking the same row in tight loop if it
    # fails before status flips.
    while True:
        try:
            with app.app_context():
                report = (
                    Report.query
                    .filter_by(status="generating")
                    .order_by(Report.created_at)
                    .first()
                )
                if report:
                    _process_report(app, report.id)
                    continue
            time.sleep(POLL_INTERVAL)
        except Exception as e:
            print(f"[powerscan] report worker error: {e}", flush=True)
            traceback.print_exc()
            time.sleep(POLL_INTERVAL)


def start_report_worker(app):
    thread = threading.Thread(target=_worker_loop, args=(app,), daemon=True)
    thread.start()
    print("[powerscan] report worker started", flush=True)
